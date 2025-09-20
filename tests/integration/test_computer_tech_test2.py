#!/usr/bin/env python3
import pytest
pytest.skip("integration tests require manual setup (API keys, large data)", allow_module_level=True)
from tests.integration import PROJECT_ROOT
# -*- coding: utf-8 -*-
"""
测试只有目录的计算机应用技术_test2.docx文件
"""

import sys
import os
sys.path.append(str(PROJECT_ROOT))

from src.thesis_inno_eval.ai_toc_extractor import AITocExtractor
import docx

def analyze_document_structure():
    """分析文档结构"""
    doc_path = r"c:\MyProjects\thesis_Inno_Eval\data\input\计算机应用技术_test2.docx"
    print("📄 分析计算机应用技术_test2.docx文档结构")
    print("=" * 80)
    
    try:
        doc = docx.Document(doc_path)
        print(f"📊 文档总行数: {len(doc.paragraphs)}")
        print(f"\n📋 文档完整内容:")
        print("-" * 60)
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                print(f"{i+1:3d}: {text}")
            else:
                print(f"{i+1:3d}: [空行]")
                
    except Exception as e:
        print(f"❌ 分析文档失败: {str(e)}")

def test_toc_extraction():
    """测试目录提取"""
    doc_path = r"c:\MyProjects\thesis_Inno_Eval\data\input\计算机应用技术_test2.docx"
    print("\n\n🧠 测试计算机应用技术_test2.docx目录提取")
    print("=" * 80)
    
    try:
        extractor = AITocExtractor()
        result = extractor.extract_toc(doc_path)
        
        if result and result.entries:
            print(f" 成功提取目录")
            print(f"📊 提取条目数: {len(result.entries)}")
            print(f"🎯 置信度: {result.confidence_score:.2f}")
            print(f"🔧 提取方法: {result.extraction_method}")
            
            print(f"\n📋 目录条目列表:")
            print("-" * 80)
            
            for i, entry in enumerate(result.entries, 1):
                type_label = f"【{entry.section_type}】" if entry.section_type else "【unknown】"
                page_info = f"页码: {entry.page}" if entry.page else "页码: None"
                
                print(f"{i:2d}. {type_label} {entry.title}")
                print(f"     {page_info} | 级别: {entry.level} | 行号: {entry.line_number}")
            
            # 检查参考文献后章节
            ref_found = False
            ref_index = -1
            for i, entry in enumerate(result.entries):
                if entry.section_type == 'references' or '参考文献' in entry.title:
                    ref_found = True
                    ref_index = i
                    break
            
            print(f"\n🔍 参考文献后章节检查:")
            print("-" * 50)
            
            if ref_found:
                print(f" 找到参考文献: {result.entries[ref_index].title} (页码: {result.entries[ref_index].page})")
                
                post_ref_sections = result.entries[ref_index + 1:]
                if post_ref_sections:
                    print(f"\n📖 参考文献后章节 ({len(post_ref_sections)}个):")
                    for i, section in enumerate(post_ref_sections, 1):
                        print(f"   {i}. {section.title} (页码: {section.page}, 类型: {section.section_type})")
                else:
                    print("📝 没有参考文献后章节")
            else:
                print("❌ 没有找到参考文献章节")
                
        else:
            print("❌ 提取失败或结果为空")
            
    except Exception as e:
        print(f"❌ 提取过程中出错: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    analyze_document_structure()
    test_toc_extraction()
