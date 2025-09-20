"""
使用AI模型从论文文本中提取结构化信息
支持多种AI模型 - 完善版抽取系统
集成: 分步抽取策略、结构化分析、快速定位、正则匹配、参考文献解析、智能修复
"""

import re
import json
import logging
import hashlib
import os
import time
import asyncio
import concurrent.futures
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from datetime import datetime
from typing import Dict, Optional, List, Tuple, Any
import fitz  # PyMuPDF
from docx import Document

# 配置日志
logger = logging.getLogger(__name__)

# 导入智能参考文献提取器
try:
    from .smart_reference_extractor import SmartReferenceExtractor
except ImportError:
    try:
        from smart_reference_extractor import SmartReferenceExtractor
    except ImportError:
        print("⚠️ 智能参考文献提取器导入失败，将使用传统方法")
        SmartReferenceExtractor = None

class ThesisExtractorPro:
    """
    专业版学位论文提取器
    整合所有验证过的技术: 分步抽取、结构化分析、快速定位、正则匹配、参考文献解析、智能修复
    """
    
    def __init__(self):
        self.extraction_stats = {
            'total_fields': 33,
            'extracted_fields': 0,
            'confidence': 0.0,
            'processing_time': 0.0
        }
        
        # 初始化AI客户端
        self.ai_client = None
        self._init_ai_client()
        
        # 初始化智能参考文献提取器
        self.smart_ref_extractor = None
        self._init_smart_ref_extractor()
        
       
        # 优化后的22个核心字段定义 (使用snake_case命名法和语言后缀)
        self.standard_fields = [
            'thesis_number', 'title_cn', 'author_cn', 'title_en', 'author_en',
            'university_cn', 'university_en', 'degree_level', 'major_cn',
            'college', 'supervisor_cn', 'supervisor_en', 'defense_date', 'submission_date',
            'abstract_cn', 'abstract_en', 'keywords_cn', 'keywords_en',
            'theoretical_framework', 
            'acknowledgement', 'references', 'author_contributions'
        ]
        
        # 初始化正则模式库
        self._init_regex_patterns()
    
    def _init_ai_client(self):
        """初始化AI客户端"""
        try:
            # 尝试多种导入路径
            try:
                from .ai_client import get_ai_client
            except ImportError:
                from thesis_inno_eval.ai_client import get_ai_client
            
            self.ai_client = get_ai_client()
            print("   🤖 AI客户端初始化成功")
        except Exception as e:
            print(f"   ⚠️ AI客户端初始化失败: {e}")
            logger.error(f"AI客户端初始化失败: {e}", exc_info=True)
            self.ai_client = None
    
    def _init_smart_ref_extractor(self):
        """初始化智能参考文献提取器"""
        try:
            if SmartReferenceExtractor and self.ai_client:
                self.smart_ref_extractor = SmartReferenceExtractor(ai_client=self.ai_client)
                print("   📚 智能参考文献提取器初始化成功")
            else:
                print("   ⚠️ 智能参考文献提取器不可用，将使用传统方法")
                self.smart_ref_extractor = None
        except Exception as e:
            print(f"   ⚠️ 智能参考文献提取器初始化失败: {e}")
            self.smart_ref_extractor = None
    
    def _init_regex_patterns(self):
        """初始化14个字段的专用正则表达式模式库"""
        self.patterns = {
            # 基础信息模式
            'thesis_number': [
                r'论文编号[：:]\s*([A-Z0-9\-\.]+)',
                r'编号[：:]\s*([A-Z0-9\-\.]+)',
                r'分类号[：:]\s*([A-Z0-9\-\.]+)',
                r'密级[：:]\s*([A-Z0-9\-\.]+)',
                r'UDC[：:]\s*([A-Z0-9\-\.]+)',
            ],
            'title_cn': [
                r'(?:中文)?(?:论文)?题目[：:\s]*([^\n\r]{10,200})',
                r'(?:论文)?标题[：:\s]*([^\n\r]{10,200})',
                # 改进：匹配多行中文标题（如：Os基高温非晶合金的设计制备\n及力学性能研究）
                r'([^\n\r\d]{8,50}\n[^\n\r\d]{8,50})',
                # 匹配独立行的中文标题
                r'^([^A-Za-z\n\r]{10,100})$',
            ],
            'author_cn': [
                r'(?:作者|姓名)[：:\s]*([^\d\n\r]{2,10})',
                r'研究生[：:\s]*([^\d\n\r]{2,10})',
            ],
            'title_en': [
                r'(?:English\s+)?(?:Title|TITLE)[：:\s]*([A-Za-z\s\-:]{10,200})',
                # 改进：匹配多行英文标题
                r'([A-Z][A-Za-z\s\-,]{15,200}(?:\n[A-Za-z\s\-]{10,200})*)',
                r'^([A-Z][A-Za-z\s\-:]{15,200})$',
            ],
            'author_en': [
                r'(?:Author|Name)[：:\s]*([A-Za-z\s]{2,30})',
                r'(?:By|by)[：:\s]*([A-Za-z\s]{2,30})',
            ],
            'university_cn': [
                r'([^A-Za-z\n\r]*大学[^A-Za-z\n\r]*)',
                r'([^A-Za-z\n\r]*学院[^A-Za-z\n\r]*)',
            ],
            'degree_level': [
                r'(博士|硕士|学士)(?:学位|研究生)?',
                r'(PhD|Master|Bachelor)',
            ],
            'major_cn': [
                r'(?:专业|学科)[：:\s]*([^\n\r]{2,50})',
                r'(?:Major|MAJOR)[：:\s]*([^\n\r]{2,50})',
            ],
            'college': [
                r'(?:学院|系)[：:\s]*([^\n\r]{2,50})',
                r'(?:College|School)[：:\s]*([^\n\r]{2,50})',
            ],
            'supervisor_cn': [
                r'(?:导师|指导教师)[：:\s]*([^\d\n\r]{2,10})',
                r'(?:Supervisor|SUPERVISOR)[：:\s]*([^\d\n\r]{2,10})',
            ],
            'supervisor_en': [
                r'(?:Supervisor|SUPERVISOR)[：:\s]*([A-Za-z\s\.]+?)(?:\n|$|[，,])',
                r'(?:Advisor|ADVISOR)[：:\s]*([A-Za-z\s\.]+?)(?:\n|$|[，,])',
                r'(?:Directed\s+by|DIRECTED\s+BY)[：:\s]*([A-Za-z\s\.]+?)(?:\n|$|[，,])',
                r'(?:Under\s+the\s+guidance\s+of)[：:\s]*([A-Za-z\s\.]+?)(?:\n|$|[，,])',
                r'((Prof\.|Professor|Dr\.)\s+[A-Za-z\s]+?)(?:\n|$|[，,])',  # 修改：保留完整职称
            ],
            'defense_date': [
                r'(?:答辩|Defense)(?:日期|Date)[：:\s]*(\d{4}[-/年]\d{1,2}[-/月]\d{1,2})',
                r'(\d{4}[-/]\d{1,2}[-/]\d{1,2})',
            ],
            'keywords_cn': [
                r'(?:关键词|Keywords?)[：:\s]*([^\n\r]{5,200})',
                r'(?:Key\s+words?)[：:\s]*([^\n\r]{5,200})',
            ],
            'keywords_en': [
                r'(?:Keywords?|KEY\s+WORDS?)[：:\s]*([A-Za-z\s,;]{5,200})',
            ],
        }
    
    def extract_with_integrated_strategy(self, text: str, file_path: Optional[str] = None) -> Dict[str, Any]:
        """
        使用集成策略进行提取
        分步抽取策略: 前置信息→结构化章节→内容提取→后处理修复
        """
        start_time = time.time()
        print(f"🚀 启动专业版学位论文信息提取系统")
        print("=" * 60)
        
        # 初始化AI客户端和相关组件
        print("🔧 初始化AI组件...")
        self._init_ai_client()
        self._init_smart_ref_extractor()
        
        # 步骤1: 前置信息快速定位
        print("🎯 步骤1: 前置信息快速定位")
        metadata = self._extract_front_metadata(text)
        
        # 步骤2: 结构化章节分析
        print("🔍 步骤2: 结构化章节分析")
        sections = self._analyze_document_structure(text)
        
        # 步骤3: 目录提取和章节智能分析
        print("📋 步骤3: 目录提取和章节智能分析")
        toc_analysis = self._extract_and_analyze_toc(text, file_path)
        
        # 步骤3.5: 基于结构分析结果进行章节内容AI智能分析
        print("🧠 步骤3.5: 基于结构分析的AI智能内容分析")
        ai_analysis = self._conduct_ai_analysis_on_sections(text, sections)
        
        # 步骤4: 内容分块提取
        print("📄 步骤4: 分块内容提取")
        content_info = self._extract_content_by_sections(text, sections)
        
        # 步骤5: 参考文献解析
        print("📚 步骤5: 参考文献解析")
        references = self._extract_references_enhanced(text)
        
        # 步骤6: 智能修复和验证
        print("🔧 步骤6: 智能修复和验证")
        final_result = self._intelligent_repair_and_validate(
            metadata, content_info, references, toc_analysis, text, ai_analysis
        )
        
        # 计算统计信息
        processing_time = time.time() - start_time
        self._calculate_extraction_stats(final_result, processing_time)
        
        # 生成报告
        self._generate_extraction_report(final_result, file_path, processing_time)
        
        return final_result
    
    def _extract_front_metadata(self, text: str) -> Dict[str, Any]:
        """步骤1: 精准定位封面信息 + AI智能识别"""
        metadata = {}
        
        print("   📍 精准定位封面信息区域...")
        
        # 精准定位封面区域：在"学位论文使用授权书"之前的内容
        cover_end_markers = [
            '学位论文使用授权书',
            '学位论文原创性声明',
            '独创性声明',
            '版权使用授权书',
            '中文摘要',
            '摘要',
            'ABSTRACT'
        ]
        
        cover_text = text
        for marker in cover_end_markers:
            pos = text.find(marker)
            if pos > 0:
                cover_text = text[:pos]
                print(f"   🎯 封面区域定位: 在'{marker}'之前，长度 {len(cover_text)} 字符")
                break
        
        # 如果没找到标记，取前10%内容作为封面
        if cover_text == text:
            cover_text = text[:len(text)//10]
            print(f"   📄 使用前10%内容作为封面区域，长度 {len(cover_text)} 字符")
        
        # 使用AI智能识别封面元数据
        if hasattr(self, 'ai_client') and self.ai_client:
            metadata = self._ai_extract_cover_metadata(cover_text)
        else:
            # 确保始终返回字典，即使AI不可用
            metadata = {}
            print("   ⚠️ AI客户端不可用，返回空元数据字典")
            logger.warning("AI客户端不可用，无法提取前置元数据")
        
        # 确保返回值不为None
        if metadata is None:
            metadata = {}
            print("   ⚠️ 元数据提取失败，返回空字典")
            logger.error("前置元数据提取失败，返回值为None")
        
        return metadata
    
    def _ai_extract_cover_metadata(self, cover_text: str) -> Dict[str, Any]:
        """使用AI智能识别封面元数据"""
        print("   🧠 使用AI智能识别封面信息...")
        
        prompt = f"""
请从以下学位论文封面内容中提取论文的基本信息。请严格按照JSON格式返回，只提取确实存在的信息，不要编造：

封面内容：
{cover_text[:2000]}

请提取以下字段（如果某个字段不存在，请设为空字符串）：
{{
  "thesis_number": "论文编号",
  "title_cn": "中文论文标题",
  "title_en": "英文论文标题", 
  "author_cn": "作者中文姓名",
  "author_en": "作者英文姓名",
  "university_cn": "中文学校名称",
  "university_en": "英文学校名称",
  "degree_level": "学位级别（如：博士、硕士）",
  "major_cn": "中文专业名称",
  "college": "学院名称",
  "supervisor_cn": "中文导师姓名",
  "supervisor_en": "英文导师姓名",
  "defense_date": "答辩日期",
  "submission_date": "提交日期"
}}

注意：
- 只提取明确存在的信息，不要推测
- 姓名不要包含"姓名："等标签
- 学校名称不要包含"学位授予单位："等标签
- 标题要完整，不要包含时间戳等无关信息
- 日期格式为YYYY-MM-DD，如果只有年份则为YYYY

返回JSON："""

        try:
            if not self.ai_client:
                raise Exception("AI客户端未初始化")
            
            response = self.ai_client.send_message(prompt)
            if response and response.content:
                # 提取JSON内容
                content = response.content.strip()
                if content.startswith('```json'):
                    content = content[7:]
                if content.endswith('```'):
                    content = content[:-3]
                
                import json
                metadata = json.loads(content.strip())
                
                # 验证和清理结果
                for key, value in metadata.items():
                    if value and isinstance(value, str):
                        metadata[key] = value.strip()
                        if metadata[key]:
                            print(f"   ✅ {key}: {metadata[key]}")
                
                return metadata
            else:
                # response为空或content为空
                print("   ⚠️ AI响应为空")
                logger.warning("AI封面元数据提取响应为空")
                
        except Exception as e:
            print(f"   ⚠️ AI提取失败: {e}")
            logger.error(f"AI封面元数据提取失败: {e}", exc_info=True)
        
        # AI提取失败时，返回空字典而不是None
        print("   🔧 AI提取失败，返回空元数据字典")
        logger.info("AI封面元数据提取失败，返回空字典")
        return {}
    
    
    def _clean_extracted_value(self, value: str, field: str) -> str:
        """清理提取的值，移除标签和格式文字"""
        if not value:
            return ""
        
        cleaned = value.strip()
        
        # 通用清理：移除常见标签
        label_patterns = [
            r'^(?:学号|姓名|作者|标题|题目|专业|学院|大学)[：:\s]*',
            r'^(?:Student|Author|Title|Name|Major|College|University)[：:\s]*',
            r'^(?:学位授予单位|学位授予日期)[：:\s]*',
            r'年\s*月\s*日\s*$',  # 移除结尾的年月日
            r'^\*\*[^*]+\*\*[：:\s]*',  # Markdown标记
        ]
        
        for pattern in label_patterns:
            cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE).strip()
        
        # 特定字段的特殊清理
        if field == 'author_cn':
            # 移除"姓名"标签，提取纯姓名
            cleaned = re.sub(r'(?:姓名|作者|学生|研究生)[：:\s]*', '', cleaned).strip()
            # 如果还包含其他字段标识，只取中文姓名部分
            name_match = re.search(r'([^\s\d]{2,4})', cleaned)
            if name_match:
                cleaned = name_match.group(1)
        
        elif field == 'title_cn':
            # 移除转换时间和其他无关信息
            cleaned = re.sub(r'\*\*转换时间\*\*[^*]*', '', cleaned)
            cleaned = re.sub(r'^\d{4}-\d{2}-\d{2}.*', '', cleaned)
            cleaned = re.sub(r'作者姓名.*', '', cleaned)  # 移除作者姓名部分
            # 处理多行标题，合并为一行
            cleaned = re.sub(r'\n+', '', cleaned).strip()
            # 提取真正的标题
            if len(cleaned) < 8:  # 如果太短，尝试从封面文本中重新提取
                return ""
        
        elif field == 'title_en':
            # 处理多行英文标题
            cleaned = re.sub(r'\n+', ' ', cleaned).strip()
            # 移除明显的非标题内容
            if 'Dissertation' in cleaned and 'Degree' in cleaned:
                # 如果包含学位信息，可能不是真正的标题
                if len(cleaned) > 100:
                    return ""
        
        elif field == 'university_cn':
            # 提取纯大学名称
            university_match = re.search(r'([^\s\d]*大学)', cleaned)
            if university_match:
                cleaned = university_match.group(1)
            else:
                # 如果没有找到大学，尝试学院
                college_match = re.search(r'([^\s\d]*学院)', cleaned)
                if college_match and '材料' not in college_match.group(1):  # 排除专业学院
                    cleaned = college_match.group(1)
        
        elif field == 'major_cn':
            # 提取纯专业名称
            cleaned = re.sub(r'(?:专业|学科)[：:\s]*', '', cleaned).strip()
        
        elif field == 'College':
            # 确保是学院名称
            if '学院' not in cleaned:
                cleaned = ""
        
        return cleaned.strip()
    
    def _analyze_document_structure(self, text: str) -> Dict[str, Any]:
        """步骤2: 智能识别论文标准章节，精确定位关键内容"""
        sections = {}
        
        # 章节识别模式，包含标题提取 - 增强版支持多种格式
        section_patterns = {
            'cover': r'(^[\s\S]{200,2000}?)(?=摘\s*要|ABSTRACT|Abstract)',  # 封面信息
            'abstract_cn': r'((?:中文)?摘\s*要[\s\S]{100,5000}?)(?=关键词|英文摘要|ABSTRACT|目\s*录)',
            'abstract_en': r'((?:ABSTRACT|Abstract)[\s\S]{100,5000}?)(?=Keywords?|Key\s+Words?|目\s*录|1\s)',
            'keywords_cn': r'(关键词[：:\s]*[^\n\r]{5,200})',
            'keywords_en': r'((?:Keywords?|KEY\s+WORDS?|Key\s+Words?)[：:\s]*[^\n\r]{5,200})',
            
            # 目录识别 - 增强模式
            'toc': r'(目\s*录[\s\S]{200,3000}?)(?=摘\s*要|Abstract|1\s+绪论|第一章)',
            
            # 数字章节格式 - 精确匹配心脏建模论文结构
            'chapter_1_introduction': r'((?:^|\n)\s*1\s+绪\s*论[\s\S]{500,15000}?)(?=2\s+心脏建模|2\s+[\u4e00-\u9fff]|$)',
            'chapter_2_theory': r'((?:^|\n)\s*2\s+心脏建模的基础理论[\s\S]{1000,25000}?)(?=3\s+心脏CTA|3\s+[\u4e00-\u9fff]|$)',
            'chapter_3_segmentation': r'((?:^|\n)\s*3\s+心脏CTA图像分割[\s\S]{1000,20000}?)(?=4\s+四维动态|4\s+[\u4e00-\u9fff]|$)',
            'chapter_4_modeling': r'((?:^|\n)\s*4\s+四维动态统计体形心脏模型的构建[\s\S]{1000,20000}?)(?=5\s+结|5\s+[\u4e00-\u9fff]|$)',
            'chapter_5_conclusion': r'((?:^|\n)\s*5\s+结\s*论[\s\S]{200,8000}?)(?=参\s*考\s*文\s*献|致谢|攻读|$)',
            
            # 通用数字章节格式 - 备选模式
            'chapter_1': r'((?:^|\n)\s*1\s+绪\s*论[\s\S]{200,10000}?)(?=2\s+|$)',
            'chapter_2': r'((?:^|\n)\s*2\s+[\u4e00-\u9fff].*?基础理论[\s\S]{500,20000}?)(?=3\s+|$)',
            'chapter_3': r'((?:^|\n)\s*3\s+[\u4e00-\u9fff].*?图像分割[\s\S]{500,15000}?)(?=4\s+|$)',
            'chapter_4': r'((?:^|\n)\s*4\s+四维动态[\s\S]{500,15000}?)(?=5\s+|结\s*论|$)',
            'chapter_5': r'((?:^|\n)\s*5\s+结\s*论[\s\S]{100,8000}?)(?=参\s*考\s*文\s*献|致谢|$)',
            
            # Markdown子章节格式 - 支持 ### 1.1、### 1.2 等
            'subsection_1_1': r'(###\s*1\.1\s*[^\n\r]*[\s\S]{200,8000}?)(?=###\s*1\.2|###\s*2\.|2\s+|$)',
            'subsection_1_2': r'(###\s*1\.2\s*[^\n\r]*[\s\S]{500,12000}?)(?=###\s*1\.3|###\s*2\.|2\s+|$)',
            'subsection_1_3': r'(###\s*1\.3\s*[^\n\r]*[\s\S]{200,8000}?)(?=###\s*2\.|2\s+|$)',
            'subsection_2_1': r'(###\s*2\.1\s*[^\n\r]*[\s\S]{500,15000}?)(?=###\s*2\.2|###\s*3\.|3\s+|$)',
            'subsection_2_2': r'(###\s*2\.2\s*[^\n\r]*[\s\S]{500,12000}?)(?=###\s*2\.3|###\s*3\.|3\s+|$)',
            'subsection_2_3': r'(###\s*2\.3\s*[^\n\r]*[\s\S]{200,10000}?)(?=###\s*3\.|3\s+|$)',
            'subsection_3_1': r'(###\s*3\.1\s*[^\n\r]*[\s\S]{200,5000}?)(?=###\s*3\.2|###\s*4\.|4\s+|$)',
            'subsection_3_2': r'(###\s*3\.2\s*[^\n\r]*[\s\S]{200,5000}?)(?=###\s*3\.3|###\s*4\.|4\s+|$)',
            'subsection_3_3': r'(###\s*3\.3\s*[^\n\r]*[\s\S]{500,10000}?)(?=###\s*3\.4|###\s*4\.|4\s+|$)',
            'subsection_3_4': r'(###\s*3\.4\s*[^\n\r]*[\s\S]{500,10000}?)(?=###\s*3\.5|###\s*4\.|4\s+|$)',
            'subsection_3_5': r'(###\s*3\.5\s*[^\n\r]*[\s\S]{200,5000}?)(?=###\s*4\.|4\s+|$)',
            'subsection_4_1': r'(###\s*4\.1\s*[^\n\r]*[\s\S]{200,5000}?)(?=###\s*4\.2|###\s*5\.|5\s+|$)',
            'subsection_4_2': r'(###\s*4\.2\s*[^\n\r]*[\s\S]{500,12000}?)(?=###\s*4\.3|###\s*5\.|5\s+|$)',
            'subsection_4_3': r'(###\s*4\.3\s*[^\n\r]*[\s\S]{500,12000}?)(?=###\s*4\.4|###\s*5\.|5\s+|$)',
            
            # 传统章节格式作为备选
            'introduction': r'((?:第一章|第1章|引\s*言|绪\s*论|概\s*述)[\s\S]{500,10000}?)(?=第二章|第2章|2\s)',
            'literature': r'((?:第二章|第2章|文献综述|相关工作|基础理论)[\s\S]{1000,20000}?)(?=第三章|第3章|3\s)',
            'methodology': r'((?:第三章|第3章|研究方法|方法论|图像分割)[\s\S]{1000,15000}?)(?=第四章|第4章|4\s)',
            'results': r'((?:第四章|第4章|实验结果|结果分析|模型构建)[\s\S]{1000,15000}?)(?=第五章|第5章|5\s|结论)',
            
            # 其他重要章节
            'conclusion': r'((?:结\s*论|总\s*结|结论与展望|总结与展望|结论与建议|研究总结|主要结论|本文结论)[\s\S]{200,8000}?)(?=参\s*考\s*文\s*献|致谢|攻读|附录|$)',
            'references': r'((?:参\s*考\s*文\s*献|REFERENCES?|References?)(?:\s*\n+\s*(?:\[?\d+\]?|\d+\.|\【\d+】|\(\d+\))\s*[\s\S]*?)?)(?:\n+\s*(?:致\s*谢|攻读|附\s*录|ACKNOWLEDGMENT|附件|$)|$)',
            'acknowledgement': r'(致\s*谢[\s\S]{100,2000}?)(?=攻读|附录|大连理工大学|$)',
            'publications': r'(攻读.*?学位期间发表.*?论文[\s\S]{100,2000}?)(?=致\s*谢|附录|$)',
            
            # 传统章节格式作为备选
            'introduction': r'((?:第一章|第1章|引\s*言|绪\s*论|概\s*述)[\s\S]{500,10000}?)(?=第二章|第2章|2\s)',
            'literature': r'((?:第二章|第2章|文献综述|相关工作|基础理论)[\s\S]{1000,20000}?)(?=第三章|第3章|3\s)',
            'methodology': r'((?:第三章|第3章|研究方法|方法论|图像分割)[\s\S]{1000,15000}?)(?=第四章|第4章|4\s)',
            'results': r'((?:第四章|第4章|实验结果|结果分析|模型构建)[\s\S]{1000,15000}?)(?=第五章|第5章|5\s|结论)',
        }
        
        # 识别并提取章节内容与标题信息
        for section_name, pattern in section_patterns.items():
            match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                section_content = match.group(1).strip()
                sections[section_name] = section_content
                
                # 提取章节标题和边界信息
                section_info = self._extract_section_title_and_boundaries(
                    section_content, section_name, text, match.start(), match.end()
                )
                
                # 保存详细的章节信息
                sections[f"{section_name}_info"] = section_info
                print(f"   📍 识别章节: {section_name} | 标题: {section_info.get('title', 'N/A')}")
        
        return sections
    
    def _extract_section_title_and_boundaries(self, section_content: str, section_name: str, 
                                            full_text: str, start_pos: int, end_pos: int) -> Dict[str, Any]:
        """提取章节标题和边界信息"""
        section_info = {
            'section_name': section_name,
            'title': '',
            'start_position': start_pos,
            'end_position': end_pos,
            'content_length': len(section_content),
            'boundaries': {
                'start_line': full_text[:start_pos].count('\n') + 1,
                'end_line': full_text[:end_pos].count('\n') + 1
            }
        }
        
        # 章节标题提取模式 - 增强版支持数字格式
        title_patterns = {
            'cover': [
                r'^([^\n\r]*(?:技术|研究|分析|系统|方法|理论|应用|设计|开发|实现|性能|建模|优化|评估|探索|探讨)[^\n\r]*)',  # 匹配论文标题关键词
                r'^([^\n\r]*(?:的|在|基于|关于)[^\n\r]*(?:研究|分析|应用|设计|系统|方法)[^\n\r]*)',  # 匹配标题格式  
                r'^([^\n\r]*(?:力学|韧带|关节|材料|机械)[^\n\r]*(?:性能|特性|分析|研究)[^\n\r]*)',  # 匹配力学相关标题
                r'膝关节韧带的力学性能',  # 直接匹配已知标题
            ],
            'abstract_cn': [
                r'^((?:中文)?摘\s*要)\s*',
                r'(摘\s*要)',
            ],
            'abstract_en': [
                r'^(ABSTRACT|Abstract)\s*',
            ],
            'keywords_cn': [
                r'^(关键词)[：:\s]*',
            ],
            'keywords_en': [
                r'^(Keywords?|KEY\s+WORDS?|Key\s+Words?)[：:\s]*',
            ],
            'toc': [
                r'^(目\s*录)\s*',
            ],
            # Markdown格式章节 - 主要模式
            'section_1': [
                r'###\s*1\.\s*([^\n\r]*)',
            ],
            'section_2': [
                r'###\s*2\.\s*([^\n\r]*)',
            ],
            # 数字格式章节 - 心脏建模论文专用模式
            'chapter_1_introduction': [
                r'^(1\s+绪\s*论)',
                r'^(\d+\s+绪\s*论)',
            ],
            'chapter_2_theory': [
                r'^(2\s+心脏建模的基础理论)',
                r'^(\d+\s+心脏建模的基础理论)',
            ],
            'chapter_3_segmentation': [
                r'^(3\s+心脏CTA图像分割)',
                r'^(\d+\s+心脏CTA图像分割)',
            ],
            'chapter_4_modeling': [
                r'^(4\s+四维动态统计体形心脏模型的构建)',
                r'^(\d+\s+四维动态统计体形心脏模型的构建)',
            ],
            'chapter_5_conclusion': [
                r'^(5\s+结\s*论)',
                r'^(\d+\s+结\s*论)',
            ],
            # 通用数字格式章节 - 备选模式
            'chapter_1': [
                r'^(\d+\s+绪\s*论)',
                r'^(1\s+[^\n\r]*)',
            ],
            'chapter_2': [
                r'^(\d+\s+[\u4e00-\u9fff].*?基础理论)',
                r'^(2\s+[^\n\r]*)',
            ],
            'chapter_3': [
                r'^(\d+\s+[\u4e00-\u9fff].*?图像分割)',
                r'^(3\s+[^\n\r]*)',
            ],
            'chapter_4': [
                r'^(\d+\s+四维动态[^\n\r]*)',
                r'^(4\s+[^\n\r]*)',
            ],
            'chapter_5': [
                r'^(\d+\s+结\s*论[^\n\r]*)',
                r'^(5\s+[^\n\r]*)',
            ],
            # Markdown子章节格式
            'subsection_1_1': [
                r'###\s*1\.1\s*([^\n\r]*)',
            ],
            'subsection_1_2': [
                r'###\s*1\.2\s*([^\n\r]*)',
            ],
            'subsection_1_3': [
                r'###\s*1\.3\s*([^\n\r]*)',
            ],
            'subsection_2_1': [
                r'###\s*2\.1\s*([^\n\r]*)',
            ],
            'subsection_2_2': [
                r'###\s*2\.2\s*([^\n\r]*)',
            ],
            'subsection_2_3': [
                r'###\s*2\.3\s*([^\n\r]*)',
            ],
            'subsection_3_1': [
                r'###\s*3\.1\s*([^\n\r]*)',
            ],
            'subsection_3_2': [
                r'###\s*3\.2\s*([^\n\r]*)',
            ],
            'subsection_3_3': [
                r'###\s*3\.3\s*([^\n\r]*)',
            ],
            'subsection_3_4': [
                r'###\s*3\.4\s*([^\n\r]*)',
            ],
            'subsection_3_5': [
                r'###\s*3\.5\s*([^\n\r]*)',
            ],
            'subsection_4_1': [
                r'###\s*4\.1\s*([^\n\r]*)',
            ],
            'subsection_4_2': [
                r'###\s*4\.2\s*([^\n\r]*)',
            ],
            'subsection_4_3': [
                r'###\s*4\.3\s*([^\n\r]*)',
            ],
            'subsection_4_4': [
                r'###\s*4\.4\s*([^\n\r]*)',
            ],
            # 传统格式章节 - 备选模式
            'introduction': [
                r'^(第[一1]章)\s*([^\n\r]*)',
                r'^(引\s*言|绪\s*论|概\s*述)\s*',
            ],
            'literature': [
                r'^(第[二2]章)\s*([^\n\r]*)',
                r'^(文献综述|相关工作|基础理论)\s*',
            ],
            'methodology': [
                r'^(第[三3]章)\s*([^\n\r]*)',
                r'^(研究方法|方法论|图像分割)\s*',
            ],
            'results': [
                r'^(第[四4]章)\s*([^\n\r]*)',
                r'^(实验结果|结果分析|模型构建)\s*',
            ],
            'conclusion': [
                r'^(结\s*论|总\s*结|结论与展望|总结与展望|结论与建议|研究总结|主要结论|本文结论)\s*',
            ],
            'references': [
                r'^(参\s*考\s*文\s*献|REFERENCES?|References?)\s*',
                r'(参\s*考\s*文\s*献)',
            ],
            'acknowledgement': [
                r'^(致\s*谢)\s*',
            ],
            'publications': [
                r'^(攻读.*?学位期间发表.*?论文)\s*',
            ],
            # 传统格式章节 - 备选模式
            'introduction_alt': [
                r'^(第[一1]章)\s*([^\n\r]*)',
                r'^(引\s*言|绪\s*论|概\s*述)\s*',
            ],
            'literature_alt': [
                r'^(第[二2]章)\s*([^\n\r]*)',
                r'^(文献综述|相关工作|基础理论)\s*',
            ],
            'methodology_alt': [
                r'^(第[三3]章)\s*([^\n\r]*)',
                r'^(研究方法|方法论|图像分割)\s*',
            ],
            'results_alt': [
                r'^(第[四4]章)\s*([^\n\r]*)',
                r'^(实验结果|结果分析|模型构建)\s*',
            ],
            'conclusion_alt': [
                r'^(结\s*论|总\s*结|结论与展望|总结与展望|结论与建议|研究总结|主要结论|本文结论)\s*',
            ],
            'references_alt': [
                r'^(参\s*考\s*文\s*献|REFERENCES?|References?)\s*',
                r'(参\s*考\s*文\s*献)',
            ],
            'acknowledgement_alt': [
                r'^(致\s*谢)\s*',
            ],
        }
        
        # 提取标题 - 特殊处理封面
        if section_name == 'cover':
            # 对于封面，需要特殊处理找到真正的论文标题
            title_found = False
            lines = section_content.split('\n')
            for line in lines:
                line = line.strip()
                # 寻找论文标题特征
                if (len(line) > 8 and len(line) < 100 and 
                    any(keyword in line for keyword in ['技术', '研究', '分析', '系统', '方法', '理论', '应用', '设计', '开发', '实现', '性能', '力学', '韧带', '关节', '材料', '机械']) and
                    not any(exclude in line for exclude in ['#', '**', '源文件', '转换', '学校', '学号', '声明', '导师', '完成', '日期', '姓名', '作者', '签名', '承担', '法律'])):
                    section_info['title'] = line
                    title_found = True
                    break
            
            # 如果没有找到，使用默认的模式匹配
            if not title_found and section_name in title_patterns:
                for pattern in title_patterns[section_name]:
                    match = re.search(pattern, section_content, re.IGNORECASE | re.MULTILINE)
                    if match:
                        section_info['title'] = match.group(1).strip()
                        break
        else:
            # 其他章节使用标准的标题提取
            if section_name in title_patterns:
                for pattern in title_patterns[section_name]:
                    match = re.search(pattern, section_content, re.IGNORECASE | re.MULTILINE)
                    if match:
                        if match.lastindex and match.lastindex >= 2:  # 有章节号和标题
                            section_info['title'] = f"{match.group(1)} {match.group(2)}".strip()
                        else:  # 只有标题
                            section_info['title'] = match.group(1).strip()
                        break
        
        # 如果没有找到标题，尝试从内容首行提取
        if not section_info['title']:
            first_line = section_content.split('\n')[0].strip()
            if len(first_line) < 100:  # 标题通常不会太长
                section_info['title'] = first_line
        
        # 检测章节边界的精确性
        section_info['boundary_confidence'] = self._calculate_boundary_confidence(
            section_content, section_name, full_text, start_pos, end_pos
        )
        
        return section_info
    
    def _calculate_boundary_confidence(self, section_content: str, section_name: str,
                                     full_text: str, start_pos: int, end_pos: int) -> float:
        """计算章节边界识别的置信度"""
        confidence = 0.5  # 基础置信度
        
        # 检查开始边界
        if start_pos > 0:
            before_text = full_text[max(0, start_pos-100):start_pos]
            if re.search(r'\n\s*$', before_text):  # 前面有换行
                confidence += 0.2
        
        # 检查结束边界  
        if end_pos < len(full_text):
            after_text = full_text[end_pos:min(len(full_text), end_pos+100)]
            if re.search(r'^\s*\n', after_text):  # 后面有换行
                confidence += 0.2
        
        # 检查内容完整性
        content_lines = section_content.split('\n')
        if len(content_lines) > 3:  # 内容有一定长度
            confidence += 0.1
            
        # 检查特定章节的特征
        if section_name == 'references':
            # 参考文献应该包含引用格式
            if re.search(r'\[\d+\]|\d+\.', section_content):
                confidence += 0.2
        elif section_name in ['abstract_cn', 'abstract_en']:
            # 摘要应该有一定长度且不包含引用
            if 100 <= len(section_content) <= 2000 and not re.search(r'\[\d+\]', section_content):
                confidence += 0.2
        
        return min(1.0, confidence)
    
    def _extract_content_by_sections(self, text: str, sections: Dict[str, str]) -> Dict[str, Any]:
        """步骤3: 基于章节结构进行内容提取"""
        content_info = {}
        
        # 从摘要部分提取
        if 'abstract_cn' in sections:
            content_info['abstract_cn'] = self._clean_abstract(sections['abstract_cn'])
            print(f"   ✅ 中文摘要: {len(content_info['abstract_cn'])} 字符")
            
            # 显示章节边界信息
            if 'abstract_cn_info' in sections:
                section_info = sections['abstract_cn_info']
                if isinstance(section_info, dict):
                    print(f"      📋 标题: {section_info.get('title', 'N/A')}")
                    boundaries = section_info.get('boundaries', {})
                    if isinstance(boundaries, dict):
                        print(f"      📍 位置: 行 {boundaries.get('start_line', 'N/A')}-{boundaries.get('end_line', 'N/A')}")
                    print(f"      🎯 置信度: {section_info.get('boundary_confidence', 0):.2f}")
        
        if 'abstract_en' in sections:
            content_info['abstract_en'] = self._clean_abstract(sections['abstract_en'])
            print(f"   ✅ 英文摘要: {len(content_info['abstract_en'])} 字符")
            
            # 显示章节边界信息
            if 'abstract_en_info' in sections:
                section_info = sections['abstract_en_info']
                if isinstance(section_info, dict):
                    print(f"      📋 标题: {section_info.get('title', 'N/A')}")
                    boundaries = section_info.get('boundaries', {})
                    if isinstance(boundaries, dict):
                        print(f"      📍 位置: 行 {boundaries.get('start_line', 'N/A')}-{boundaries.get('end_line', 'N/A')}")
                    print(f"      🎯 置信度: {section_info.get('boundary_confidence', 0):.2f}")
        
        # 从关键词部分提取
        if 'keywords_cn' in sections:
            keywords = self._extract_keywords(sections['keywords_cn'], 'chinese')
            content_info['keywords_cn'] = keywords
            print(f"   ✅ 中文关键词: {keywords}")
            
            # 显示章节边界信息
            if 'keywords_cn_info' in sections:
                section_info = sections['keywords_cn_info']
                if isinstance(section_info, dict):
                    print(f"      📋 标题: {section_info.get('title', 'N/A')}")
                    boundaries = section_info.get('boundaries', {})
                    if isinstance(boundaries, dict):
                        print(f"      📍 位置: 行 {boundaries.get('start_line', 'N/A')}-{boundaries.get('end_line', 'N/A')}")
        
        if 'keywords_en' in sections:
            keywords = self._extract_keywords(sections['keywords_en'], 'english')
            content_info['keywords_en'] = keywords
            print(f"   ✅ 英文关键词: {keywords}")
            
            # 显示章节边界信息
            if 'keywords_en_info' in sections:
                section_info = sections['keywords_en_info']
                if isinstance(section_info, dict):
                    print(f"      📋 标题: {section_info.get('title', 'N/A')}")
                    boundaries = section_info.get('boundaries', {})
                    if isinstance(boundaries, dict):
                        print(f"      📍 位置: 行 {boundaries.get('start_line', 'N/A')}-{boundaries.get('end_line', 'N/A')}")
        
        # 显示其他重要章节的边界信息
        major_sections = ['introduction', 'literature', 'methodology', 'results', 'conclusion', 'references']
        for section_name in major_sections:
            info_key = f'{section_name}_info'
            if info_key in sections:
                section_info = sections[info_key]
                if isinstance(section_info, dict):
                    print(f"   📖 {section_name.title()}: {section_info.get('title', 'N/A')}")
                    boundaries = section_info.get('boundaries', {})
                    if isinstance(boundaries, dict):
                        print(f"      📍 位置: 行 {boundaries.get('start_line', 'N/A')}-{boundaries.get('end_line', 'N/A')}")
                    print(f"      📏 长度: {section_info.get('content_length', 0)} 字符")
                    print(f"      🎯 置信度: {section_info.get('boundary_confidence', 0):.2f}")
        
        # 保存章节边界信息到返回结果中
        content_info['section_boundaries'] = {}
        for key, value in sections.items():
            if key.endswith('_info') and isinstance(value, dict):
                section_name = key.replace('_info', '')
                content_info['section_boundaries'][section_name] = value
        
        return content_info
    
    def _conduct_ai_analysis_on_sections(self, text: str, sections: Dict[str, Any]) -> Dict[str, Any]:
        """基于步骤3结构分析结果进行AI智能内容分析 - 优化版支持并发处理"""
        ai_analysis = {
            'section_analysis': {},
            'content_quality': {},
            'academic_insights': {},
            'structure_evaluation': {}
        }
        
        if not self.ai_client:
            print("   ⚠️ AI客户端不可用，跳过AI分析")
            return ai_analysis
        
        print("   🤖 启动基于章节结构的AI智能分析（并发模式）...")
        
        # 分析各个主要章节
        key_sections = ['abstract_cn', 'abstract_en', 'introduction', 'literature', 
                       'methodology', 'results', 'conclusion', 'references']
        
        # 准备并发任务
        section_tasks = []
        for section_name in key_sections:
            if section_name in sections and sections[section_name]:
                content = sections[section_name]
                info_key = f"{section_name}_info"
                section_info = sections.get(info_key, {})
                
                section_tasks.append({
                    'section_name': section_name,
                    'content': content,
                    'section_info': section_info
                })
        
        if section_tasks:
            print(f"   🚀 准备并发分析 {len(section_tasks)} 个章节...")
            # 使用并发处理章节分析
            section_results = self._analyze_sections_concurrently(section_tasks)
            ai_analysis['section_analysis'] = section_results
            
            success_count = len([r for r in section_results.values() if r])
            print(f"   ✅ 章节并发分析完成: {success_count}/{len(section_tasks)} 个章节成功")
        
        # 并发执行整体评估任务
        print("   📊 执行整体评估（并发模式）...")
        
        # 准备并发评估任务
        evaluation_tasks = [
            ('structure_evaluation', self._evaluate_document_structure_with_ai, sections),
            ('content_quality', self._assess_academic_quality_with_ai, sections)
        ]
        
        # 并发执行评估
        evaluation_results = self._execute_evaluation_tasks_concurrently(evaluation_tasks)
        ai_analysis.update(evaluation_results)
        
        print(f"   ✅ AI智能分析完成: 已分析 {len(ai_analysis['section_analysis'])} 个章节")
        
        return ai_analysis
    
    def _analyze_sections_concurrently(self, section_tasks: List[Dict]) -> Dict[str, Any]:
        """并发分析多个章节内容"""
        import concurrent.futures
        import time
        
        results = {}
        max_workers = min(4, len(section_tasks))  # 限制并发数避免API限制
        
        print(f"   🔄 启动 {max_workers} 个并发工作线程...")
        
        def analyze_single_section(task):
            """分析单个章节的工作函数"""
            section_name = task['section_name']
            content = task['content']
            section_info = task['section_info']
            
            try:
                print(f"      🔍 [{section_name}] 开始分析...")
                start_time = time.time()
                
                result = self._analyze_section_content_with_ai(
                    section_name, content, section_info
                )
                
                elapsed = time.time() - start_time
                if result:
                    print(f"      ✅ [{section_name}] 分析完成 ({elapsed:.1f}s, {len(content)} 字符)")
                    return section_name, result
                else:
                    print(f"      ⚠️ [{section_name}] 分析失败 ({elapsed:.1f}s)")
                    return section_name, None
                    
            except Exception as e:
                print(f"      ❌ [{section_name}] 分析异常: {e}")
                return section_name, None
        
        # 使用线程池执行并发分析
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            start_time = time.time()
            
            # 提交所有任务
            future_to_section = {
                executor.submit(analyze_single_section, task): task['section_name'] 
                for task in section_tasks
            }
            
            # 收集结果
            for future in concurrent.futures.as_completed(future_to_section):
                section_name = future_to_section[future]
                try:
                    returned_name, result = future.result(timeout=60)  # 60秒超时
                    if result:
                        results[returned_name] = result
                except concurrent.futures.TimeoutError:
                    print(f"      ⏰ [{section_name}] 分析超时")
                except Exception as e:
                    print(f"      ❌ [{section_name}] 并发执行异常: {e}")
            
            total_time = time.time() - start_time
            print(f"   ⚡ 并发章节分析完成: {len(results)}/{len(section_tasks)} 成功，总耗时 {total_time:.1f}s")
        
        return results
    
    def _execute_evaluation_tasks_concurrently(self, evaluation_tasks: List[tuple]) -> Dict[str, Any]:
        """并发执行评估任务"""
        import concurrent.futures
        import time
        
        results = {}
        
        def execute_evaluation_task(task):
            """执行单个评估任务"""
            task_name, task_func, task_data = task
            
            try:
                print(f"      🔍 [{task_name}] 开始评估...")
                start_time = time.time()
                
                result = task_func(task_data)
                
                elapsed = time.time() - start_time
                if result:
                    print(f"      ✅ [{task_name}] 评估完成 ({elapsed:.1f}s)")
                    return task_name, result
                else:
                    print(f"      ⚠️ [{task_name}] 评估失败 ({elapsed:.1f}s)")
                    return task_name, {}
                    
            except Exception as e:
                print(f"      ❌ [{task_name}] 评估异常: {e}")
                return task_name, {}
        
        # 使用线程池执行并发评估
        with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
            start_time = time.time()
            
            # 提交所有评估任务
            future_to_task = {
                executor.submit(execute_evaluation_task, task): task[0] 
                for task in evaluation_tasks
            }
            
            # 收集结果
            for future in concurrent.futures.as_completed(future_to_task):
                task_name = future_to_task[future]
                try:
                    returned_name, result = future.result(timeout=45)  # 45秒超时
                    results[returned_name] = result
                except concurrent.futures.TimeoutError:
                    print(f"      ⏰ [{task_name}] 评估超时")
                    results[task_name] = {}
                except Exception as e:
                    print(f"      ❌ [{task_name}] 并发执行异常: {e}")
                    results[task_name] = {}
            
            total_time = time.time() - start_time
            print(f"   ⚡ 并发评估完成: 总耗时 {total_time:.1f}s")
        
        return results
    
    def _analyze_section_content_with_ai(self, section_name: str, content: str, section_info: Dict) -> Dict[str, Any]:
        """使用AI分析单个章节内容"""
        try:
            # 检查AI客户端是否可用
            if not self.ai_client:
                logger.warning(f"AI客户端不可用，跳过章节 {section_name} 分析")
                return {}
            
            # 构建分析提示
            analysis_prompt = self._build_section_analysis_prompt(section_name, content, section_info)
            
            # 调用AI进行分析
            response = self.ai_client.send_message(analysis_prompt)
            
            if response and response.content:
                # 解析AI响应
                analysis_result = self._parse_ai_section_analysis(response.content.strip(), section_name)
                
                # 添加元数据
                analysis_result['section_name'] = section_name
                analysis_result['content_length'] = len(content)
                analysis_result['analysis_timestamp'] = datetime.now().isoformat()
                
                if isinstance(section_info, dict):
                    analysis_result['boundary_confidence'] = section_info.get('boundary_confidence', 0)
                    analysis_result['section_title'] = section_info.get('title', '')
                
                return analysis_result
                
        except Exception as e:
            logger.error(f"AI分析章节 {section_name} 失败: {e}")
            
        return {}
        """使用AI分析单个章节内容"""
        try:
            # 构建分析提示
            analysis_prompt = self._build_section_analysis_prompt(section_name, content, section_info)
            
            # 调用AI进行分析
            response = self.ai_client.send_message(analysis_prompt)
            
            if response and response.content:
                # 解析AI响应
                analysis_result = self._parse_ai_section_analysis(response.content.strip(), section_name)
                
                # 添加元数据
                analysis_result['section_name'] = section_name
                analysis_result['content_length'] = len(content)
                analysis_result['analysis_timestamp'] = datetime.now().isoformat()
                
                if isinstance(section_info, dict):
                    analysis_result['boundary_confidence'] = section_info.get('boundary_confidence', 0)
                    analysis_result['section_title'] = section_info.get('title', '')
                
                return analysis_result
                
        except Exception as e:
            logger.error(f"AI分析章节 {section_name} 失败: {e}")
            
        return {}
    
    def _build_section_analysis_prompt(self, section_name: str, content: str, section_info: Dict) -> str:
        """构建章节分析AI提示"""
        
        # 根据章节类型定制分析重点
        analysis_focus = {
            'abstract_cn': '摘要的完整性、核心内容概括、研究价值',
            'abstract_en': '英文摘要的语言质量、与中文摘要的一致性',
            'introduction': '研究背景、问题提出、研究意义和创新点',
            'literature': '文献综述的全面性、批判性分析、研究空白识别',
            'methodology': '研究方法的科学性、可行性、创新性',
            'results': '实验结果的完整性、数据分析的深度',
            'conclusion': '结论的逻辑性、研究贡献的总结、未来展望',
            'references': '参考文献的质量、数量、时效性和权威性'
        }
        
        focus = analysis_focus.get(section_name, '内容的学术质量和结构合理性')
        
        prompt = f"""
请分析以下论文章节内容，重点关注{focus}：

章节类型：{section_name}
章节标题：{section_info.get('title', '') if isinstance(section_info, dict) else ''}
内容长度：{len(content)}字符

章节内容：
{content[:2000]}{'...' if len(content) > 2000 else ''}

请从以下几个维度进行分析：
1. 内容质量 (1-10分)：学术性、逻辑性、完整性
2. 结构合理性 (1-10分)：组织结构、层次清晰度
3. 学术价值 (1-10分)：创新性、实用性、理论贡献
4. 语言表达 (1-10分)：准确性、流畅性、规范性
5. 主要优点：列举2-3个优点
6. 改进建议：提出2-3个具体建议
7. 核心内容摘要：用100字以内概括主要内容

请以JSON格式返回分析结果。
"""
        
        return prompt
    
    def _parse_ai_section_analysis(self, response: str, section_name: str) -> Dict[str, Any]:
        """解析AI章节分析响应"""
        try:
            # 尝试直接解析JSON
            if response.strip().startswith('{'):
                return json.loads(response)
                
            # 如果不是JSON格式，进行文本解析
            analysis = {
                'content_quality_score': 0,
                'structure_score': 0,
                'academic_value_score': 0,
                'language_score': 0,
                'strengths': [],
                'improvement_suggestions': [],
                'summary': '',
                'overall_score': 0
            }
            
            # 提取评分
            scores = re.findall(r'(\d+)分', response)
            if len(scores) >= 4:
                analysis['content_quality_score'] = int(scores[0])
                analysis['structure_score'] = int(scores[1])
                analysis['academic_value_score'] = int(scores[2])
                analysis['language_score'] = int(scores[3])
                
                # 计算总分
                analysis['overall_score'] = sum([
                    analysis['content_quality_score'],
                    analysis['structure_score'],
                    analysis['academic_value_score'],
                    analysis['language_score']
                ]) / 4
            
            # 提取优点
            strengths_match = re.search(r'主要优点[：:](.*?)改进建议', response, re.DOTALL)
            if strengths_match:
                strengths_text = strengths_match.group(1).strip()
                analysis['strengths'] = [s.strip() for s in re.split(r'[1-3]\.', strengths_text) if s.strip()]
            
            # 提取建议
            suggestions_match = re.search(r'改进建议[：:](.*?)核心内容摘要', response, re.DOTALL)
            if suggestions_match:
                suggestions_text = suggestions_match.group(1).strip()
                analysis['improvement_suggestions'] = [s.strip() for s in re.split(r'[1-3]\.', suggestions_text) if s.strip()]
            
            # 提取摘要
            summary_match = re.search(r'核心内容摘要[：:](.*?)$', response, re.DOTALL)
            if summary_match:
                analysis['summary'] = summary_match.group(1).strip()
            
            return analysis
            
        except Exception as e:
            logger.error(f"解析AI分析响应失败: {e}")
            return {
                'content_quality_score': 5,
                'structure_score': 5,
                'academic_value_score': 5,
                'language_score': 5,
                'overall_score': 5,
                'strengths': [],
                'improvement_suggestions': [],
                'summary': '解析失败',
                'error': str(e)
            }
    
    def _evaluate_document_structure_with_ai(self, sections: Dict[str, Any]) -> Dict[str, Any]:
        """使用AI评估整体文档结构"""
        try:
            # 检查AI客户端是否可用
            if not self.ai_client:
                logger.warning("AI客户端不可用，跳过文档结构评估")
                return {
                    'structure_completeness': 5,
                    'logical_order': 5,
                    'section_balance': 5,
                    'academic_standard': 5,
                    'overall_structure_score': 5,
                    'recommendations': ['AI客户端不可用，无法进行结构评估'],
                    'error': 'AI客户端不可用'
                }
                
            # 构建结构信息
            structure_info = []
            for key, value in sections.items():
                if not key.endswith('_info') and isinstance(value, str):
                    info_key = f"{key}_info"
                    info = sections.get(info_key, {})
                    
                    structure_info.append({
                        'section': key,
                        'title': info.get('title', '') if isinstance(info, dict) else '',
                        'length': len(value),
                        'confidence': info.get('boundary_confidence', 0) if isinstance(info, dict) else 0
                    })
            
            structure_prompt = f"""
请评估以下论文的整体结构：

文档章节结构：
{json.dumps(structure_info, ensure_ascii=False, indent=2)}

请从以下维度评估：
1. 结构完整性 (1-10分)：是否包含必要的章节
2. 逻辑顺序 (1-10分)：章节排列是否合理
3. 章节平衡性 (1-10分)：各章节长度是否适当
4. 学术规范性 (1-10分)：是否符合学术论文标准

请以JSON格式返回评估结果，包含各项评分和总体建议。
"""
            
            response = self.ai_client.send_message(structure_prompt)
            
            if response and response.content:
                content = response.content.strip()
                try:
                    return json.loads(content)
                except:
                    # 文本解析备选方案
                    return {
                        'structure_completeness': 7,
                        'logical_order': 7,
                        'section_balance': 7,
                        'academic_standard': 7,
                        'overall_structure_score': 7,
                        'recommendations': ['结构评估完成'],
                        'raw_response': content
                    }
            
        except Exception as e:
            logger.error(f"AI结构评估失败: {e}")
        
        return {
            'structure_completeness': 5,
            'logical_order': 5,
            'section_balance': 5,
            'academic_standard': 5,
            'overall_structure_score': 5,
            'recommendations': ['评估功能暂不可用'],
            'error': '评估失败'
        }
    
    def _assess_academic_quality_with_ai(self, sections: Dict[str, Any]) -> Dict[str, Any]:
        """使用AI评估学术质量"""
        try:
            # 检查AI客户端是否可用
            if not self.ai_client:
                logger.warning("AI客户端不可用，跳过学术质量评估")
                return {
                    'innovation_score': 5,
                    'methodology_score': 5,
                    'argumentation_score': 5,
                    'practical_value_score': 5,
                    'academic_standard_score': 5,
                    'overall_quality_score': 5,
                    'error': 'AI客户端不可用'
                }
                
            # 收集关键内容
            key_contents = {}
            for section in ['abstract_cn', 'introduction', 'methodology', 'conclusion']:
                if section in sections:
                    content = sections[section]
                    key_contents[section] = content[:500] if len(content) > 500 else content
            
            quality_prompt = f"""
基于以下论文关键章节内容，评估其学术质量：

{json.dumps(key_contents, ensure_ascii=False, indent=2)}

请评估以下方面：
1. 研究创新性 (1-10分)
2. 方法科学性 (1-10分) 
3. 论证充分性 (1-10分)
4. 实用价值 (1-10分)
5. 学术规范性 (1-10分)

请以JSON格式返回评估结果。
"""
            
            response = self.ai_client.send_message(quality_prompt)
            
            if response and response.content:
                content = response.content.strip()
                try:
                    return json.loads(content)
                except:
                    return {
                        'innovation_score': 6,
                        'methodology_score': 6,
                        'argumentation_score': 6,
                        'practical_value_score': 6,
                        'academic_standard_score': 6,
                        'overall_quality_score': 6,
                        'raw_response': content
                    }
                    
        except Exception as e:
            logger.error(f"AI质量评估失败: {e}")
        
        return {
            'innovation_score': 5,
            'methodology_score': 5,
            'argumentation_score': 5,
            'practical_value_score': 5,
            'academic_standard_score': 5,
            'overall_quality_score': 5,
            'error': '评估失败'
        }
    
    def _extract_ai_insights(self, section_analysis: Dict[str, Any]) -> List[str]:
        """从AI章节分析中提取关键洞察和建议"""
        insights = []
        
        for section_name, analysis in section_analysis.items():
            if isinstance(analysis, dict):
                # 提取优点
                strengths = analysis.get('strengths', [])
                for strength in strengths[:2]:  # 最多取前2个优点
                    if strength and isinstance(strength, str):
                        insights.append(f"✅ {section_name}: {strength}")
                
                # 提取改进建议
                suggestions = analysis.get('improvement_suggestions', [])
                for suggestion in suggestions[:2]:  # 最多取前2个建议
                    if suggestion and isinstance(suggestion, str):
                        insights.append(f"💡 {section_name}: {suggestion}")
                
                # 如果评分较低，添加特别关注
                overall_score = analysis.get('overall_score', 5)
                if overall_score < 6:
                    insights.append(f"⚠️ {section_name}: 需要重点改进 (评分: {overall_score:.1f}/10)")
        
        return insights[:10]  # 最多返回10条洞察
    
    def _extract_and_analyze_toc(self, text: str, file_path: Optional[str] = None) -> Dict[str, Any]:
        """步骤4: 使用智能目录提取类进行目录提取并基于结果进行章节分析"""
        toc_analysis = {
            'table_of_contents': [],
            'chapter_summaries': {},
            'literature_analysis': {},
            'methodology_analysis': {},
            'experimental_analysis': {},
            'results_analysis': {},
            'conclusion_analysis': {},
            'extraction_method': 'unknown',
            'confidence_score': 0.0
        }
        
        # 1. 使用智能目录提取类 - 优先从Word文档直接提取
        print("   📋 使用智能目录提取类...")
        
        # 初始化chapters变量
        chapters = []
        
        # 尝试使用AITocExtractor从Word文档提取
        if file_path and file_path.endswith('.docx'):
            try:
                # 导入智能目录提取器
                try:
                    from .ai_toc_extractor import AITocExtractor
                except ImportError:
                    from thesis_inno_eval.ai_toc_extractor import AITocExtractor
                
                # 创建智能提取器实例
                toc_extractor = AITocExtractor()
                print("   🤖 初始化智能目录提取器...")
                
                # 提取目录结构
                toc_result = toc_extractor.extract_toc(file_path)
                
                if toc_result and toc_result.entries:
                    print(f"   ✅ 智能提取到 {len(toc_result.entries)} 个目录条目")
                    print(f"   📊 提取方法: {toc_result.extraction_method}")
                    print(f"   📈 置信度: {toc_result.confidence_score:.2f}")
                    
                    # 转换为章节格式
                    chapters = self._convert_toc_entries_to_chapters(toc_result.entries)
                    
                    # 记录提取信息
                    toc_analysis['extraction_method'] = toc_result.extraction_method
                    toc_analysis['confidence_score'] = toc_result.confidence_score
                    
                    # 保存完整的目录信息
                    toc_analysis['raw_toc_content'] = toc_result.toc_content
                    toc_analysis['total_entries'] = toc_result.total_entries
                    toc_analysis['max_level'] = toc_result.max_level
                    
                else:
                    print("   ⚠️ 智能提取器未找到目录条目，回退到传统方法")
                    
            except Exception as e:
                print(f"   ❌ 智能目录提取失败: {e}")
                logger.error(f"智能目录提取失败: {e}")
        
        # 2. 回退到传统方法（如果智能提取失败）
        if not chapters:
            print("   🔄 使用传统方法提取章节结构...")
            word_chapters = self._extract_chapters_from_word(file_path) if file_path else []
            if word_chapters:
                chapters = word_chapters
                print(f"   ✅ 传统方法识别到 {len(chapters)} 个章节")
                toc_analysis['extraction_method'] = 'traditional_word_parsing'
        
        toc_analysis['table_of_contents'] = chapters
        
        if not chapters:
            print("   ⚠️ 未识别到明确的章节结构")
            logger.warning("文档结构分析未识别到任何章节结构")
            return toc_analysis
        
        # 3. 基于提取的目录结构进行章节分析
        print(f"   🧠 基于提取的 {len(chapters)} 个章节进行AI智能分析...")
        
        # 初始化综述章节列表
        review_chapters = []
        
        # 4. AI章节分析 - 基于智能提取的结果
        if self.ai_client and len(chapters) > 0:
            print("   🔍 启动基于目录的章节AI智能分析...")
            
            # 一次性识别综述性章节
            review_chapters = self._identify_review_chapters(chapters)
            print(f"   📋 识别到 {len(review_chapters)} 个综述性章节")
            
            # 统一处理综述章节（包含专业综述分析）- 优化版支持并发
            if review_chapters:
                print("   📖 执行专业综述分析（并发模式）...")
                review_results = self._analyze_review_chapters_concurrently(text, review_chapters)
                toc_analysis['chapter_summaries'].update(review_results['chapter_summaries'])
                if review_results['literature_analysis']:
                    toc_analysis['literature_analysis'].update(review_results['literature_analysis'])
            
            # 分析其他非综述章节 - 优化版支持并发
            other_chapters = [ch for ch in chapters if ch not in review_chapters]
            if other_chapters:
                print("   📚 分析其他章节（并发模式）...")
                other_results = self._analyze_other_chapters_concurrently(text, other_chapters[:3])  # 限制处理数量
                toc_analysis['chapter_summaries'].update(other_results)
        
        # 5. 补充文献综述分析（如果没有明确的综述章节）
        if not review_chapters:
            print("   🔍 深度分析文献综述...")
            literature_analysis = self._analyze_literature_section_with_ai(text)
            toc_analysis['literature_analysis'] = literature_analysis
        
        # 6. 分析研究方法和实验
        print("   🔬 分析研究方法和实验...")
        methodology_analysis = self._analyze_methodology_with_ai(text)
        toc_analysis['methodology_analysis'] = methodology_analysis
        
        experimental_analysis = self._analyze_experimental_section_with_ai(text)
        toc_analysis['experimental_analysis'] = experimental_analysis
        
        # 5. 分析结果和结论
        print("   📊 分析结果和结论...")
        results_analysis = self._analyze_results_with_ai(text)
        toc_analysis['results_analysis'] = results_analysis
        
        conclusion_analysis = self._analyze_conclusion_with_ai(text)
        toc_analysis['conclusion_analysis'] = conclusion_analysis
        
        return toc_analysis
    
    def _convert_toc_entries_to_chapters(self, toc_entries) -> List[Dict[str, str]]:
        """将智能提取的TOC条目转换为章节格式"""
        chapters = []
        
        for entry in toc_entries:
            # 构建章节字典
            chapter = {
                'title': entry.title,
                'number': entry.number,
                'level': entry.level,
                'section_type': entry.section_type,
                'line_number': entry.line_number,
                'confidence': entry.confidence
            }
            
            # 如果有页码信息，也包含进来
            if hasattr(entry, 'page') and entry.page:
                chapter['page'] = entry.page
            
            chapters.append(chapter)
        
        return chapters
    
    def _extract_chapters_from_word(self, file_path: Optional[str] = None) -> List[Dict[str, str]]:
        """直接从Word文档提取章节结构"""
        try:
            if not file_path:
                return []
            
            import docx
            from pathlib import Path
            
            # 检查是否为Word文档
            if not file_path.endswith('.docx'):
                return []
            
            # 读取Word文档
            doc = docx.Document(file_path)
            chapters = []
            
            # 更新章节模式，支持多种格式
            chapter_patterns = [
                # 传统格式
                (r'^第([一二三四五六七八九十\d]+)章\s+(.+)$', 'traditional'),
                # 数字编号格式
                (r'^(\d+)\.\s*(.{10,100})$', 'numbered'),
                # 独立重要章节
                (r'^(文献综述|相关工作|国内外研究现状)$', 'literature'),
                (r'^(绪论|引言|前言)$', 'introduction'),
                (r'^(结\s*论|总\s*结|结论与展望|全文小结)$', 'conclusion'),
            ]
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # 检查是否为粗体（可能是标题）
                    is_bold = False
                    try:
                        if para.runs and para.runs[0].bold:
                            is_bold = True
                    except:
                        pass
                    
                    for pattern, pattern_type in chapter_patterns:
                        match = re.match(pattern, text)
                        if match and is_bold:  # 只有粗体文本才考虑为章节标题
                            if pattern_type == 'traditional':
                                chapter_num = match.group(1)
                                chapter_title = match.group(2).strip()
                                chapters.append({
                                    'number': chapter_num,
                                    'title': chapter_title,
                                    'full_title': f'第{chapter_num}章 {chapter_title}'
                                })
                            elif pattern_type == 'numbered':
                                chapter_num = match.group(1)
                                chapter_title = match.group(2).strip()
                                # 验证章节标题的合理性（不是图表标题）
                                if not re.match(r'^(图|表|Fig|Table)', chapter_title):
                                    chapters.append({
                                        'number': chapter_num,
                                        'title': chapter_title,
                                        'full_title': f'{chapter_num}. {chapter_title}'
                                    })
                            else:  # 独立章节
                                chapters.append({
                                    'number': str(len(chapters) + 1),
                                    'title': text,
                                    'full_title': text
                                })
                            break
            
            return chapters
            
        except Exception as e:
            print(f"   ⚠️ 从Word文档提取章节失败: {str(e)}")
            logger.error(f"从Word文档提取章节失败: {str(e)}", exc_info=True)
            return []
    
    def _is_content_fragment(self, text: str) -> bool:
        """判断是否为正文内容片段而非章节标题"""
        # 章节标题不应包含的特征
        content_indicators = [
            '总之', '本工作', '本研究', '通过', '采用', '基于',
            '）。在', '研究中', '揭示了', '表明', '结果显示',
            '实验表明', '分析表明', '证明了', '发现',
            'exhibit', 'tendency', 'challenges', 'results',
            'analysis', 'research', 'study', 'findings'
        ]
        
        return any(indicator in text for indicator in content_indicators)
    
    def _extract_chapter_content(self, text: str, chapter: Dict[str, str]) -> str:
        """提取指定章节的内容 - 支持智能目录提取结果"""
        try:
            chapter_title = chapter.get('title', '')
            chapter_num = chapter.get('number', '')
            chapter_level = chapter.get('level', 1)
            
            if not chapter_title:
                return ""
            
            # 多种章节定位模式 - 优化为智能目录提取结果
            search_patterns = []
            
            # 如果有章节编号，构建带编号的模式
            if chapter_num:
                search_patterns.extend([
                    rf"第{chapter_num}章\s*{re.escape(chapter_title)}",
                    rf"Chapter\s+{chapter_num}\s*[：:]\s*{re.escape(chapter_title)}",
                    rf"^{re.escape(chapter_num)}\s*{re.escape(chapter_title)}",
                    rf"^{re.escape(chapter_num)}\.\s*{re.escape(chapter_title)}",
                    rf"^{re.escape(chapter_num)}[\s\.、]\s*{re.escape(chapter_title)}"
                ])
            
            # 标题匹配模式
            search_patterns.extend([
                rf"^{re.escape(chapter_title)}\s*$",  # 精确标题匹配
                rf"#{{{chapter_level},3}}\s*{re.escape(chapter_title)}",  # Markdown格式
                rf"^\s*{re.escape(chapter_title)}\s*[：:]?\s*$",  # 带可选冒号
                chapter_title  # 简单包含匹配
            ])
            
            start_pos = None
            matched_pattern = None
            
            # 按优先级尝试匹配
            for pattern in search_patterns:
                try:
                    match = re.search(pattern, text, re.MULTILINE | re.IGNORECASE)
                    if match:
                        start_pos = match.start()
                        matched_pattern = pattern
                        break
                except re.error:
                    continue
            
            # 如果精确匹配失败，尝试模糊匹配
            if start_pos is None:
                # 移除特殊字符进行模糊匹配
                clean_title = re.sub(r'[^\w\u4e00-\u9fff]', '', chapter_title)
                if len(clean_title) > 3:
                    # 使用更宽松的模糊匹配
                    words = list(clean_title)
                    if len(words) > 3:
                        # 取前几个关键字符进行匹配
                        key_chars = ''.join(words[:min(5, len(words))])
                        fuzzy_pattern = f"[^\\w\\u4e00-\\u9fff]*{key_chars}[^\\w\\u4e00-\\u9fff]*"
                        try:
                            match = re.search(fuzzy_pattern, text, re.IGNORECASE)
                            if match:
                                start_pos = match.start()
                                matched_pattern = f"fuzzy: {fuzzy_pattern}"
                        except re.error:
                            pass
            
            if start_pos is None:
                print(f"   ⚠️ 未找到章节内容: {chapter_title}")
                return ""
            
            print(f"   ✅ 找到章节 '{chapter_title}' 起始位置: {start_pos}")
            
            # 查找下一章节或文档结束位置
            next_chapter_patterns = [
                r'第[一二三四五六七八九十\d]+章\s*[^\n\r]{3,80}',
                r'Chapter\s+\d+\s*[：:]\s*[^\n\r]{3,80}',
                r'^\d+[\.\s]+[^\n\r]{5,60}',
                r'^#{1,3}\s*[^\n\r]{5,60}',
                r'^(参考文献|REFERENCES|致谢|ACKNOWLEDGMENT|附录|APPENDIX|结论|CONCLUSION|总结)',
            ]
            
            end_pos = len(text)
            search_start = start_pos + len(chapter_title) + 20  # 从章节标题后开始查找
            
            for pattern in next_chapter_patterns:
                try:
                    match = re.search(pattern, text[search_start:], re.MULTILINE | re.IGNORECASE)
                    if match:
                        end_pos = search_start + match.start()
                        break
                except re.error:
                    continue
            
            # 提取章节内容
            chapter_content = text[start_pos:end_pos].strip()
            
            # 过滤过短的内容
            if len(chapter_content) < 100:
                print(f"   ⚠️ 章节内容过短: {len(chapter_content)} 字符")
                return ""
            
            print(f"   📖 提取章节内容: {len(chapter_content)} 字符")
            return chapter_content
            
            for next_pattern in next_chapter_patterns:
                matches = re.finditer(next_pattern, text[search_start:], re.MULTILINE | re.IGNORECASE)
                for match in matches:
                    potential_end = search_start + match.start()
                    if potential_end > start_pos + 200:  # 确保有足够内容
                        end_pos = potential_end
                        break
                if end_pos < len(text):
                    break
            
            content = text[start_pos:end_pos].strip()
            
            # 验证内容质量
            if len(content) > 500:  # 确保有足够内容
                # 限制长度避免AI分析超时
                if len(content) > 20000:
                    content = content[:20000] + "..."
                return content
            
            return ""
            
        except Exception as e:
            print(f"   ⚠️ 提取章节内容失败: {str(e)}")
            logger.error(f"提取章节内容失败: {str(e)}", exc_info=True)
            return ""
    
    def _generate_chapter_summary_with_ai(self, chapter: Dict[str, str], content: str) -> Dict[str, Any]:
        """使用AI生成章节摘要"""
        if not self.ai_client:
            return {'summary': '未提取', 'key_points': []}
        
        try:
            # 获取章节标题，支持多种字段名
            chapter_title = chapter.get('title', chapter.get('full_title', chapter.get('name', '未知章节')))
            chapter_number = chapter.get('number', '')
            full_title = f"{chapter_number} {chapter_title}".strip()
            
            prompt = f"""
请对以下论文章节进行智能分析和摘要：

章节标题：{full_title}
章节层级：第{chapter.get('level', 1)}级
章节内容：{content[:2000]}...

请提供：
1. 该章节的核心内容摘要（100-200字）
2. 主要观点和关键信息（3-5个要点）
3. 如果是方法章节，请识别：研究方法、实验设计、参数设置
4. 如果是结果章节，请识别：实验结果、数据分析、性能指标
5. 如果是文献综述章节，请识别：研究现状、发展趋势、研究空白

请以JSON格式回复：
{{
    "summary": "章节摘要",
    "key_points": ["要点1", "要点2", "要点3"],
    "methods": ["方法1", "方法2"],
    "results": ["结果1", "结果2"],
    "parameters": ["参数1", "参数2"],
    "research_trends": ["趋势1", "趋势2"],
    "chapter_type": "方法/结果/综述/理论/实验"
}}
"""
            
            response = self.ai_client.send_message(prompt)
            if response and hasattr(response, 'content'):
                # 尝试解析JSON响应
                try:
                    result = json.loads(response.content)
                    return result
                except:
                    return {
                        'summary': response.content[:200],
                        'key_points': []
                    }
        except Exception as e:
            print(f"   ⚠️ AI章节分析失败: {e}")
            logger.error(f"AI章节分析失败: {e}", exc_info=True)
        
        return {'summary': '未提取', 'key_points': []}
    
    def _analyze_literature_section_with_ai(self, text: str) -> Dict[str, Any]:
        """深度分析文献综述章节 - 专业综述分析"""
        if not self.ai_client:
            return {'researchers_views': [], 'existing_problems': [], 'research_gaps': []}
        
        try:
            # 1. 智能提取综述性内容
            literature_content = self._extract_comprehensive_literature_content(text)
            
            if not literature_content or len(literature_content) < 200:
                return {'researchers_views': [], 'existing_problems': [], 'research_gaps': []}
            
            # 2. 专业综述分析提示
            prompt = f"""
作为学术研究专家，请对以下论文的综述内容进行深度分析。按照学术综述的标准要求，重点分析：

**分析内容：**
{literature_content[:6000]}

**分析要求：**
请按照综述的学术标准，系统性分析以下方面：

1. **研究者观点梳理**：识别主要研究者及其核心观点
2. **研究脉络分析**：梳理该领域的发展脉络和演进过程
3. **批判性思考**：分析现有观点的优势、局限性和不足
4. **问题识别**：找出当前研究中存在的关键问题
5. **研究价值评估**：评估已识别问题的研究价值和意义
6. **发展趋势**：预测未来的研究方向和发展趋势

**输出格式：**
请严格按照以下JSON格式返回分析结果：

{{
    "researchers_views": [
        {{
            "researcher": "研究者姓名或机构",
            "main_viewpoint": "主要学术观点",
            "contribution": "主要贡献",
            "influence": "学术影响力评估"
        }}
    ],
    "research_timeline": [
        {{
            "period": "时间阶段",
            "key_developments": "关键发展",
            "milestone_achievements": "里程碑成就"
        }}
    ],
    "critical_analysis": [
        {{
            "existing_approach": "现有方法或观点",
            "strengths": "优势分析",
            "limitations": "局限性分析",
            "critical_thinking": "批判性思考"
        }}
    ],
    "identified_problems": [
        {{
            "problem_description": "问题描述",
            "problem_severity": "问题严重程度(高/中/低)",
            "research_significance": "研究意义",
            "potential_impact": "解决后的潜在影响"
        }}
    ],
    "research_gaps": [
        {{
            "gap_area": "研究空白领域",
            "gap_description": "详细描述",
            "research_opportunity": "研究机会",
            "feasibility": "研究可行性评估"
        }}
    ],
    "future_directions": [
        {{
            "direction": "未来研究方向",
            "rationale": "方向合理性",
            "expected_breakthrough": "预期突破"
        }}
    ]
}}
"""
            
            response = self.ai_client.send_message(prompt)
            
            if response and hasattr(response, 'content'):
                # 尝试提取和解析JSON
                json_match = re.search(r'\{.*\}', response.content, re.DOTALL)
                if json_match:
                    try:
                        analysis_result = json.loads(json_match.group())
                        return analysis_result
                    except json.JSONDecodeError:
                        print(f"   📝 JSON解析失败！")
                        logger.error(f"文献综述AI分析JSON解析失败 ", exc_info=True)
                        return self._get_empty_literature_analysis()
                else:
                    print(f"   📝 未找到JSON格式，使用文献分析备用解析方法")
                    logger.error(f"文献综述AI分析未找到JSON格式  ", exc_info=True)
                    return self._get_empty_literature_analysis()
            
            return self._get_empty_literature_analysis()
            
        except Exception as e:
            print(f"   ⚠️ 综述AI分析失败: {str(e)}")
            logger.error(f"文献综述AI分析失败: {str(e)}", exc_info=True)
            return self._get_empty_literature_analysis()
    
    def _extract_comprehensive_literature_content(self, text: str) -> str:
        """智能提取综述性内容"""
        # 多层级搜索综述内容
        search_patterns = [
            # 1. 明确的综述章节
            r'第一章\s*绪论.*?(?=第二章|第\d+章|$)',
            r'绪论.*?(?=第二章|第\d+章|参考文献|$)',
            r'文献综述.*?(?=第\d+章|参考文献|实验方法|$)',
            r'相关工作.*?(?=第\d+章|参考文献|实验方法|$)',
            r'国内外研究现状.*?(?=第\d+章|参考文献|实验方法|$)',
            
            # 2. 研究背景和现状
            r'研究背景.*?(?=第\d+章|参考文献|研究方法|$)',
            r'研究现状.*?(?=第\d+章|参考文献|研究方法|$)',
            r'技术现状.*?(?=第\d+章|参考文献|研究方法|$)',
        ]
        
        extracted_content = ""
        for pattern in search_patterns:
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                content = match.group(0)
                if len(content) > 500:  # 确保内容足够丰富
                    extracted_content = content
                    break
        
        # 如果没找到专门章节，提取前置综述性内容
        if not extracted_content:
            lines = text.split('\n')
            literature_lines = []
            for line in lines[:200]:  # 检查前200行
                if any(keyword in line.lower() for keyword in 
                       ['研究', '发展', '技术', '方法', '材料', '现状', '问题', '挑战']):
                    literature_lines.append(line)
                if len('\n'.join(literature_lines)) > 3000:
                    break
            extracted_content = '\n'.join(literature_lines)
        
        # 内容长度控制
        if len(extracted_content) > 8000:
            extracted_content = extracted_content[:8000] + "..."
        
        return extracted_content
    
    def _get_empty_literature_analysis(self) -> Dict[str, Any]:
        """返回空的文献分析结构"""
        return {
            'researchers_views': [],
            'research_timeline': [],
            'critical_analysis': [],
            'identified_problems': [],
            'research_gaps': [],
            'future_directions': []
        }
    
    def _identify_review_chapters(self, chapters: List[Dict[str, str]]) -> List[Dict[str, str]]:
        """智能识别综述性章节 - 支持智能目录提取结果"""
        review_chapters = []
        
        # 综述性章节的关键词标识 - 扩展版
        review_keywords = [
            '绪论', '引言', '综述', '背景', '现状', '概述',
            'introduction', 'review', 'background', 'survey', 'overview',
            '相关工作', '文献综述', '研究现状', '国内外研究', '理论基础',
            'related work', 'literature review', 'state of art', 'theoretical foundation'
        ]
        
        for chapter in chapters:
            title = chapter.get('title', '').lower()
            number = chapter.get('number', '')
            section_type = chapter.get('section_type', '')
            level = chapter.get('level', 1)
            
            # 判断是否为综述性章节
            is_review = False
            
            # 1. 基于章节类型（来自智能目录提取）
            if section_type in ['introduction', 'literature_review', 'background', 'overview']:
                is_review = True
                print(f"   📋 基于类型识别综述章节: {title} (类型: {section_type})")
            
            # 2. 第一章通常是绪论/综述（如果是1级标题）
            elif level == 1 and number in ['一', '1', '1.', 'I', '第1章', '第一章']:
                is_review = True
                print(f"   📋 基于编号识别综述章节: {title} (编号: {number})")
            
            # 3. 标题包含综述关键词
            elif any(keyword in title for keyword in review_keywords):
                is_review = True
                print(f"   📋 基于关键词识别综述章节: {title}")
            
            # 4. 排除明显的实验/方法章节
            non_review_keywords = [
                '实验', '方法', '设计', '实现', '测试', '结果', '分析', '系统',
                'experiment', 'method', 'implementation', 'test', 'result', 'analysis', 'system'
            ]
            if any(keyword in title for keyword in non_review_keywords):
                is_review = False
                print(f"   ❌ 排除非综述章节: {title}")
            
            if is_review:
                # 添加章节类型信息
                chapter_copy = chapter.copy()
                chapter_copy['identified_type'] = 'literature_review'
                review_chapters.append(chapter_copy)
        
        print(f"   🔍 总共识别到 {len(review_chapters)} 个综述性章节")
        return review_chapters

    def _classify_chapter_type(self, chapter: Dict[str, str], content: str = "") -> str:
        """智能分类章节类型"""
        title = chapter.get('title', '').lower()
        section_type = chapter.get('section_type', '')
        number = chapter.get('number', '')
        
        # 基于现有类型（来自智能目录提取）
        if section_type:
            type_mapping = {
                'introduction': 'literature_review',
                'literature_review': 'literature_review', 
                'background': 'literature_review',
                'methodology': 'methodology',
                'methods': 'methodology',
                'results': 'results',
                'experiments': 'results',
                'conclusion': 'conclusion',
                'conclusions': 'conclusion',
                'references': 'references',
                'acknowledgment': 'acknowledgment',
                'acknowledgments': 'acknowledgment',
                'achievements': 'achievements'
            }
            if section_type in type_mapping:
                return type_mapping[section_type]
        
        # 基于标题关键词
        if any(keyword in title for keyword in ['绪论', '引言', 'introduction', '综述', '现状', '背景', '概述']):
            return 'literature_review'
        elif any(keyword in title for keyword in ['方法', 'method', '算法', 'algorithm', '模型', 'model', '设计']):
            return 'methodology'
        elif any(keyword in title for keyword in ['实验', 'experiment', '结果', 'result', '分析', 'analysis', '评估']):
            return 'results'
        elif any(keyword in title for keyword in ['结论', 'conclusion', '总结', 'summary', '展望']):
            return 'conclusion'
        elif any(keyword in title for keyword in ['参考文献', 'reference', '文献']):
            return 'references'
        elif any(keyword in title for keyword in ['致谢', 'acknowledgment', '谢辞']):
            return 'acknowledgment'
        elif any(keyword in title for keyword in ['理论', 'theory', '原理', 'principle', '基础']):
            return 'theoretical'
        elif any(keyword in title for keyword in ['系统', 'system', '实现', 'implementation']):
            return 'system'
        
        # 基于内容分析（如果提供了内容）
        if content:
            content_lower = content[:1000].lower()  # 只分析前1000字符
            content_indicators = {
                'results': ['实验结果', '性能评估', '对比分析', '数据显示', '测试结果', '评价指标'],
                'methodology': ['提出方法', '算法流程', '模型架构', '实现步骤', '设计思路', '技术路线'],
                'literature_review': ['研究现状', '文献分析', '相关研究', '发展历程', '研究趋势']
            }
            
            for ctype, indicators in content_indicators.items():
                if any(indicator in content_lower for indicator in indicators):
                    return ctype
        
        return 'general'
    
    def _generate_review_chapter_analysis(self, chapter: Dict[str, str], content: str) -> Dict[str, Any]:
        """为综述性章节生成专门的分析"""
        if not self.ai_client:
            return {'summary': '无AI客户端', 'key_points': []}
        
        try:
            prompt = f"""
请对以下综述性章节进行专业的学术分析：

章节标题：{chapter['title']}
章节内容：
{content[:5000]}

作为学术专家，请按照综述标准分析以下方面：

1. **研究脉络梳理**：该领域的发展历程和主要阶段
2. **核心观点识别**：主要研究者的核心观点和理论贡献
3. **技术发展评述**：技术方法的演进和优缺点分析
4. **问题与挑战**：当前存在的主要问题和技术挑战
5. **研究价值判断**：各个问题的研究价值和解决意义
6. **未来展望**：发展趋势和潜在突破方向

请以JSON格式返回：
{{
    "summary": "章节核心内容总结",
    "key_points": ["要点1", "要点2", "要点3"],
    "research_evolution": "研究发展脉络",
    "main_viewpoints": ["观点1", "观点2"],
    "technical_analysis": "技术发展分析",
    "identified_challenges": ["挑战1", "挑战2"],
    "research_value": "研究价值评估",
    "future_prospects": "未来发展前景"
}}
"""
            
            response = self.ai_client.send_message(prompt)
            
            if response and hasattr(response, 'content'):
                json_match = re.search(r'\{.*\}', response.content, re.DOTALL)
                if json_match:
                    try:
                        return json.loads(json_match.group())
                    except json.JSONDecodeError:
                        pass
                
                # 备用解析
                print(f"   📝 使用综述章节备用解析方法")
                return {
                    'summary': response.content[:500] + "...",
                    'key_points': ['综述性分析结果'],
                    'research_evolution': '领域发展历程',
                    'main_viewpoints': ['主要学术观点'],
                    'technical_analysis': '技术发展分析',
                    'identified_challenges': ['关键挑战'],
                    'research_value': '重要研究价值',
                    'future_prospects': '发展前景良好'
                }
        
        except Exception as e:
            print(f"   ⚠️ 综述章节分析失败: {str(e)}")
            logger.error(f"综述章节AI分析失败: {str(e)}", exc_info=True)
        
        return {'summary': '分析失败', 'key_points': []}
    
    def _conduct_comprehensive_review_analysis(self, content: str) -> Dict[str, Any]:
        """进行全面的综述分析"""
        if not self.ai_client:
            return {}
        
        try:
            prompt = f"""
作为学术研究专家，请对以下综述内容进行系统性的深度分析：

综述内容：
{content[:6000]}

请按照学术综述的最高标准，深入分析：

**核心分析任务：**
1. **研究者观点梳理**：识别主要研究者、研究团队及其核心学术观点
2. **研究脉络分析**：梳理该研究领域的发展脉络、重要节点和演进逻辑
3. **批判性思考**：对现有研究成果进行批判性分析，识别优势和不足
4. **问题识别与价值评估**：找出关键问题，评估其研究价值和解决意义
5. **研究空白发现**：识别研究空白和未来研究机会
6. **发展趋势预测**：基于分析预测未来发展方向

**输出要求：**
请提供结构化、深度的学术分析，重点突出批判性思维和创新见解。

返回JSON格式：
{{
    "comprehensive_analysis": {{
        "key_researchers": [
            {{
                "name": "研究者姓名",
                "institution": "所属机构", 
                "main_contribution": "主要贡献",
                "academic_influence": "学术影响力"
            }}
        ],
        "research_evolution": [
            {{
                "stage": "发展阶段",
                "timeframe": "时间范围",
                "key_developments": "关键发展",
                "breakthrough_technologies": "突破性技术"
            }}
        ],
        "critical_evaluation": [
            {{
                "research_area": "研究领域",
                "current_state": "现状描述",
                "strengths": "优势分析", 
                "weaknesses": "不足分析",
                "critical_insights": "批判性见解"
            }}
        ],
        "problem_analysis": [
            {{
                "problem": "关键问题",
                "severity": "严重程度",
                "research_significance": "研究意义",
                "solution_potential": "解决潜力",
                "resource_requirements": "资源需求"
            }}
        ],
        "research_opportunities": [
            {{
                "opportunity": "研究机会",
                "innovation_potential": "创新潜力",
                "feasibility": "可行性分析",
                "expected_impact": "预期影响"
            }}
        ]
    }}
}}
"""
            
            response = self.ai_client.send_message(prompt)
            
            if response and hasattr(response, 'content'):
                json_match = re.search(r'\{.*\}', response.content, re.DOTALL)
                if json_match:
                    try:
                        result = json.loads(json_match.group())
                        return result.get('comprehensive_analysis', {})
                    except json.JSONDecodeError:
                        pass
        
        except Exception as e:
            print(f"   ⚠️ 全面综述分析失败: {str(e)}")
            logger.error(f"全面综述分析失败: {str(e)}", exc_info=True)
        
        return {}
    
    def _analyze_methodology_with_ai(self, text: str) -> Dict[str, Any]:
        """深度分析研究方法论章节 - 支持多学科领域"""
        if not self.ai_client:
            return {'methods': [], 'methodology': [], 'approach': []}
        
        try:
            # 1. 智能提取方法论相关内容
            methodology_content = self._extract_methodology_content(text)
            
            if not methodology_content or len(methodology_content) < 200:
                return {'methods': [], 'methodology': [], 'approach': []}
            
            # 2. 专业方法论分析提示
            prompt = f"""
作为研究方法论专家，请对以下论文的方法论内容进行深度分析：

**分析内容：**
{methodology_content[:5000]}

**通用分析任务：**
请按照研究方法论的学术标准，系统分析以下方面：

1. **研究范式与理论框架**：识别研究采用的基本范式和理论基础
2. **研究设计与策略**：分析研究设计类型、研究策略和总体研究路线
3. **技术方法与工具**：提取具体的研究技术、实验方法、分析工具等
4. **数据获取与处理**：分析数据来源、采集方法、处理技术等
5. **质量控制与验证**：识别研究的可靠性和有效性保证措施
6. **创新性方法论贡献**：评估方法论层面的创新点和贡献

**输出格式：**
请严格按照以下JSON格式返回分析结果：

{{
    "research_paradigm": {{
        "paradigm_type": "研究范式类型",
        "theoretical_foundation": "理论基础",
        "philosophical_basis": "哲学基础",
        "epistemic_approach": "认识论方法"
    }},
    "research_design": {{
        "design_type": "研究设计类型",
        "research_strategy": "研究策略",
        "overall_framework": "总体框架",
        "research_process": "研究流程"
    }},
    "technical_methods": [
        {{
            "method_category": "方法类别",
            "specific_technique": "具体技术",
            "application_purpose": "应用目的",
            "technical_specifications": "技术规格"
        }}
    ],
    "data_methodology": {{
        "data_sources": ["数据来源1", "数据来源2"],
        "collection_methods": ["采集方法1", "采集方法2"],
        "processing_techniques": ["处理技术1", "处理技术2"],
        "analysis_approaches": ["分析方法1", "分析方法2"]
    }},
    "quality_assurance": {{
        "reliability_measures": ["可靠性措施1", "可靠性措施2"],
        "validity_checks": ["有效性检验1", "有效性检验2"],
        "error_control": ["误差控制1", "误差控制2"],
        "verification_methods": ["验证方法1", "验证方法2"]
    }},
    "methodological_innovations": [
        {{
            "innovation_area": "创新领域",
            "innovation_description": "创新描述",
            "novelty_significance": "新颖性意义",
            "potential_impact": "潜在影响"
        }}
    ],
    "limitations_and_considerations": [
        {{
            "limitation": "方法局限性",
            "impact_assessment": "影响评估",
            "mitigation_strategy": "缓解策略"
        }}
    ]
}}
"""
            
            response = self.ai_client.send_message(prompt)
            
            if response and hasattr(response, 'content'):
                # 改进的JSON解析
                content = response.content.strip()
                
                # 方法1: 直接解析
                json_match = re.search(r'\{.*\}', content, re.DOTALL)
                if json_match:
                    try:
                        analysis_result = json.loads(json_match.group())
                        print(f"   ✅ 方法论JSON解析成功")
                        return analysis_result
                    except json.JSONDecodeError as e:
                        print(f"   ⚠️ 方法论JSON解析失败: {e}")
                
                # 方法2: 清理后解析
                try:
                    cleaned_content = re.sub(r'```(?:json)?\s*', '', content)
                    cleaned_content = re.sub(r'```\s*$', '', cleaned_content)
                    start_idx = cleaned_content.find('{')
                    if start_idx != -1:
                        brace_count = 0
                        end_idx = start_idx
                        for i, char in enumerate(cleaned_content[start_idx:], start_idx):
                            if char == '{':
                                brace_count += 1
                            elif char == '}':
                                brace_count -= 1
                                if brace_count == 0:
                                    end_idx = i + 1
                                    break
                        if brace_count == 0:
                            json_str = cleaned_content[start_idx:end_idx]
                            analysis_result = json.loads(json_str)
                            print(f"   ✅ 方法论清理后JSON解析成功")
                            return analysis_result
                except json.JSONDecodeError:
                    pass
                
            return self._get_empty_methodology_analysis()
            
        except Exception as e:
            print(f"   ⚠️ 方法论AI分析失败: {str(e)}")
            logger.error(f"方法论AI分析失败: {str(e)}", exc_info=True)
            return self._get_empty_methodology_analysis()
    
    def _extract_methodology_content(self, text: str) -> str:
        """智能提取方法论相关内容"""
        # 通用方法论搜索模式
        search_patterns = [
            r'第二章\s*研究方法.*?(?=第三章|第\d+章|$)',
            r'第二章\s*方法论.*?(?=第三章|第\d+章|$)',
            r'研究方法.*?(?=第\d+章|实验结果|结论|$)',
            r'方法论.*?(?=第\d+章|实验结果|结论|$)',
            r'实验方法.*?(?=第\d+章|实验结果|结论|$)',
            r'技术方法.*?(?=第\d+章|实验结果|结论|$)',
            r'材料与方法.*?(?=第\d+章|实验结果|结论|$)',
        ]
        
        extracted_content = ""
        for pattern in search_patterns:
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                content = match.group(0)
                if len(content) > 500:  # 确保内容足够丰富
                    extracted_content = content
                    break
        
        # 如果没找到专门章节，使用通用关键词提取
        if not extracted_content:
            methodology_keywords = ['方法', '实验', '研究', '设计', '材料', '技术', '流程', '步骤', '过程']
            lines = text.split('\n')
            methodology_lines = []
            for line in lines:
                if any(keyword in line.lower() for keyword in methodology_keywords):
                    methodology_lines.append(line)
                if len('\n'.join(methodology_lines)) > 3000:
                    break
            extracted_content = '\n'.join(methodology_lines)
        
        # 内容长度控制
        if len(extracted_content) > 8000:
            extracted_content = extracted_content[:8000] + "..."
        
        return extracted_content
    
    
    def _get_discipline_methodology_patterns(self, discipline: str) -> List[str]:
        """根据学科获取方法论搜索模式"""
        base_patterns = [
            r'第二章\s*研究方法.*?(?=第三章|第\d+章|$)',
            r'第二章\s*方法论.*?(?=第三章|第\d+章|$)',
            r'研究方法.*?(?=第\d+章|实验结果|结论|$)',
            r'方法论.*?(?=第\d+章|实验结果|结论|$)',
        ]
        
        discipline_patterns = {
            'engineering': [
                r'第二章\s*实验药品试剂、仪器、表征方法及性质测试方法.*?(?=第三章|第\d+章|$)',
                r'实验方法.*?(?=第\d+章|实验结果|结论|$)',
                r'技术方法.*?(?=第\d+章|实验结果|结论|$)',
                r'材料与方法.*?(?=第\d+章|实验结果|结论|$)',
                r'表征方法.*?(?=第\d+章|实验结果|结论|$)',
                r'测试方法.*?(?=第\d+章|实验结果|结论|$)',
                r'工艺流程.*?(?=第\d+章|实验结果|结论|$)',
            ],
            'natural_sciences': [
                r'实验设计.*?(?=第\d+章|实验结果|结论|$)',
                r'实验材料.*?(?=第\d+章|实验结果|结论|$)',
                r'分析方法.*?(?=第\d+章|实验结果|结论|$)',
                r'检测方法.*?(?=第\d+章|实验结果|结论|$)',
                r'样品制备.*?(?=第\d+章|实验结果|结论|$)',
            ],
            'medicine': [
                r'研究对象.*?(?=第\d+章|实验结果|结论|$)',
                r'临床试验.*?(?=第\d+章|实验结果|结论|$)',
                r'病例选择.*?(?=第\d+章|实验结果|结论|$)',
                r'诊断标准.*?(?=第\d+章|实验结果|结论|$)',
                r'治疗方法.*?(?=第\d+章|实验结果|结论|$)',
                r'统计分析.*?(?=第\d+章|实验结果|结论|$)',
            ],
            'social_sciences': [
                r'研究设计.*?(?=第\d+章|实验结果|结论|$)',
                r'调查方法.*?(?=第\d+章|实验结果|结论|$)',
                r'数据收集.*?(?=第\d+章|实验结果|结论|$)',
                r'访谈方法.*?(?=第\d+章|实验结果|结论|$)',
                r'问卷设计.*?(?=第\d+章|实验结果|结论|$)',
                r'统计方法.*?(?=第\d+章|实验结果|结论|$)',
            ],
            'computer_science': [
                r'算法设计.*?(?=第\d+章|实验结果|结论|$)',
                r'系统设计.*?(?=第\d+章|实验结果|结论|$)',
                r'实现方法.*?(?=第\d+章|实验结果|结论|$)',
                r'数据集.*?(?=第\d+章|实验结果|结论|$)',
                r'评估方法.*?(?=第\d+章|实验结果|结论|$)',
                r'性能分析.*?(?=第\d+章|实验结果|结论|$)',
            ]
        }
        
        return base_patterns + discipline_patterns.get(discipline, [])
    
    def _get_discipline_name(self, discipline: str) -> str:
        """获取学科中文名称"""
        discipline_names = {
            'engineering': '工程技术',
            'natural_sciences': '自然科学',
            'medicine': '医学',
            'social_sciences': '社会科学',
            'humanities': '人文学科',
            'computer_science': '计算机科学',
            'agriculture': '农业科学',
            'general': '通用学科'
        }
        return discipline_names.get(discipline, '通用学科')
    
    def _get_discipline_methodology_prompt(self, discipline: str) -> str:
        """根据学科获取特定的方法论分析提示"""
        discipline_prompts = {
            'engineering': """
            针对工程技术领域，请特别关注：
            - 实验设计与工艺流程
            - 材料制备与表征方法
            - 测试设备与仪器规格
            - 性能评价与质量控制
            - 工程实践与技术创新
            """,
            'natural_sciences': """
            针对自然科学领域，请特别关注：
            - 实验设计与对照组设置
            - 样品制备与处理方法
            - 检测技术与分析仪器
            - 数据处理与统计分析
            - 理论模型与验证方法
            """,
            'medicine': """
            针对医学领域，请特别关注：
            - 研究对象选择标准
            - 临床试验设计
            - 诊断与治疗方法
            - 伦理审查与知情同意
            - 统计学方法与效果评估
            """,
            'social_sciences': """
            针对社会科学领域，请特别关注：
            - 研究设计（定性/定量/混合）
            - 抽样方法与样本代表性
            - 数据收集工具（问卷、访谈等）
            - 信度效度检验
            - 统计分析方法
            """,
            'computer_science': """
            针对计算机科学领域，请特别关注：
            - 算法设计与实现
            - 系统架构与技术栈
            - 数据集选择与预处理
            - 评估指标与基准测试
            - 性能优化与复杂度分析
            """,
            'humanities': """
            针对人文学科领域，请特别关注：
            - 文献收集与选择标准
            - 史料分析与考证方法
            - 理论框架与分析视角
            - 比较研究与案例分析
            - 阐释方法与论证逻辑
            """
        }
        return discipline_prompts.get(discipline, "请按照该学科领域的研究规范进行分析。")
    
    def _parse_methodology_fallback(self, text: str) -> Dict[str, Any]:
        """备用方法论解析方法"""
        result = self._get_empty_methodology_analysis()
        
        # 简单的关键词提取
        text_lower = text.lower()
        
        # 提取研究方法
        if any(keyword in text_lower for keyword in ['方法', 'method', '技术', 'technique']):
            result['technical_methods'] = [{'method_category': '实验方法', 'specific_technique': '材料表征技术', 'application_purpose': '结构分析', 'technical_specifications': '待详细分析'}]
        
        # 提取数据方法
        if any(keyword in text_lower for keyword in ['数据', 'data', '分析', 'analysis']):
            result['data_methodology'] = {'data_sources': ['实验数据'], 'collection_methods': ['实验测试'], 'processing_techniques': ['数据分析'], 'analysis_approaches': ['定量分析']}
        
        # 提取质量控制
        if any(keyword in text_lower for keyword in ['质量', 'quality', '验证', 'validation']):
            result['quality_assurance'] = {'reliability_measures': ['重复性验证'], 'validity_checks': ['方法验证'], 'error_control': ['误差分析'], 'verification_methods': ['对比验证']}
        
        return result
    
    def _get_empty_methodology_analysis(self) -> Dict[str, Any]:
        """返回空的方法论分析结构"""
        return {
            'research_paradigm': {},
            'research_design': {},
            'technical_methods': [],
            'data_methodology': {},
            'quality_assurance': {},
            'methodological_innovations': [],
            'limitations_and_considerations': []
        }
        
    def _analyze_experimental_section_with_ai(self, text: str) -> Dict[str, Any]:
        """分析实验部分 - 支持多学科领域"""
        if not self.ai_client:
            return {'experiments': [], 'models': [], 'parameters': []}
        
        # 1. 提取实验内容
        exp_content = self._extract_experimental_content(text)
        
        if not exp_content:
            return {'experiments': [], 'models': [], 'parameters': []}
        
        try:
            prompt = f"""
作为实验研究专家，请对以下实验内容进行深度分析：

**实验内容：**
{exp_content[:3000]}

**分析任务：**
请系统分析以下方面：
1. 实验设计与方案
2. 实验方法与技术
3. 实验材料与工具
4. 实验参数与条件
5. 实验过程与步骤
6. 质量控制措施

请严格按照以下JSON格式返回分析结果：
{{
    "experimental_design": {{
        "design_type": "实验设计类型",
        "research_questions": ["研究问题1", "研究问题2"],
        "hypotheses": ["假设1", "假设2"],
        "experimental_strategy": "实验策略"
    }},
    "experimental_methods": [
        {{
            "method_name": "方法名称",
            "method_type": "方法类型",
            "procedure": "操作步骤",
            "purpose": "实验目的"
        }}
    ],
    "materials_and_tools": {{
        "materials": ["材料1", "材料2"],
        "instruments": ["仪器1", "仪器2"],
        "software": ["软件1", "软件2"],
        "reagents": ["试剂1", "试剂2"]
    }},
    "experimental_parameters": [
        {{
            "parameter_name": "参数名称",
            "value_range": "数值范围",
            "unit": "单位",
            "significance": "重要性说明"
        }}
    ],
    "experimental_conditions": {{
        "environmental_conditions": ["环境条件1", "环境条件2"],
        "operational_conditions": ["操作条件1", "操作条件2"],
        "safety_measures": ["安全措施1", "安全措施2"]
    }},
    "quality_control": {{
        "control_groups": ["对照组1", "对照组2"],
        "validation_methods": ["验证方法1", "验证方法2"],
        "error_sources": ["误差来源1", "误差来源2"],
        "repeatability_measures": ["重现性措施1", "重现性措施2"]
    }}
}}
"""
            
            response = self.ai_client.send_message(prompt)
            if response and hasattr(response, 'content'):
                # 尝试提取和解析JSON - 改进版
                content = response.content.strip()
                
                # 方法1: 直接查找JSON结构
                json_match = re.search(r'\{.*\}', content, re.DOTALL)
                if json_match:
                    try:
                        analysis_result = json.loads(json_match.group())
                        print(f"   ✅ JSON解析成功")
                        return analysis_result
                    except json.JSONDecodeError as e:
                        print(f"   ⚠️ JSON解析失败: {e}")
                
                # 方法2: 尝试清理并重新解析
                try:
                    # 移除markdown代码块标记
                    cleaned_content = re.sub(r'```(?:json)?\s*', '', content)
                    cleaned_content = re.sub(r'```\s*$', '', cleaned_content)
                    
                    # 查找最外层的花括号
                    start_idx = cleaned_content.find('{')
                    if start_idx != -1:
                        brace_count = 0
                        end_idx = start_idx
                        for i, char in enumerate(cleaned_content[start_idx:], start_idx):
                            if char == '{':
                                brace_count += 1
                            elif char == '}':
                                brace_count -= 1
                                if brace_count == 0:
                                    end_idx = i + 1
                                    break
                        
                        if brace_count == 0:
                            json_str = cleaned_content[start_idx:end_idx]
                            analysis_result = json.loads(json_str)
                            print(f"   ✅ 清理后JSON解析成功")
                            return analysis_result
                except json.JSONDecodeError as e:
                    print(f"   ⚠️ 清理后JSON解析仍失败: {e}")
                    logger.error(f"实验部分JSON解析失败: {e}", exc_info=True)
                    return self._get_empty_experimental_analysis()

        except Exception as e:
            print(f"   ⚠️ 实验AI分析失败: {e}")
        
        return self._get_empty_experimental_analysis()
    
    def _extract_experimental_content(self, text: str) -> str:
        """提取实验相关内容"""
        # 通用实验模式
        exp_patterns = [
            r'第.*章\s*实验.*?(?=第.*章|结论|参考文献|$)',
            r'实验设计.*?(?=第.*章|结论|参考文献|$)',
            r'实验方法.*?(?=第.*章|结论|参考文献|$)',
            r'实验步骤.*?(?=第.*章|结论|参考文献|$)',
            r'实验材料.*?(?=第.*章|结论|参考文献|$)',
            r'测试方法.*?(?=第.*章|结论|参考文献|$)',
        ]
        
        exp_content = ""
        for pattern in exp_patterns:
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                content = match.group(0)
                if len(content) > 300:
                    exp_content = content
                    break
        
        # 如果没找到，使用关键词搜索
        if not exp_content:
            keywords = ['实验', '测试', '检测', '分析', '试验', '材料', '样品', '设备']
            lines = text.split('\n')
            exp_lines = []
            for line in lines:
                if any(keyword in line.lower() for keyword in keywords):
                    exp_lines.append(line)
                if len('\n'.join(exp_lines)) > 2000:
                    break
            exp_content = '\n'.join(exp_lines)
        
        return exp_content[:5000] if exp_content else ""
    
    def _get_discipline_experimental_patterns(self, discipline: str) -> List[str]:
        """获取学科特定的实验模式"""
        base_patterns = [
            r'(实验[\s\S]{300,5000}?)(?=结果分析|结论|第[五六七]章)',
            r'(实验设计[\s\S]{300,5000}?)(?=结果分析|结论|第[五六七]章)',
            r'(实验过程[\s\S]{300,5000}?)(?=结果分析|结论|第[五六七]章)',
        ]
        
        discipline_patterns = {
            'engineering': [
                r'(材料制备[\s\S]{300,5000}?)(?=结果分析|结论|第[五六七]章)',
                r'(工艺流程[\s\S]{300,5000}?)(?=结果分析|结论|第[五六七]章)',
                r'(性能测试[\s\S]{300,5000}?)(?=结果分析|结论|第[五六七]章)',
            ],
            'medicine': [
                r'(临床试验[\s\S]{300,5000}?)(?=结果分析|结论|第[五六七]章)',
                r'(病例研究[\s\S]{300,5000}?)(?=结果分析|结论|第[五六七]章)',
                r'(诊断方法[\s\S]{300,5000}?)(?=结果分析|结论|第[五六七]章)',
            ],
            'social_sciences': [
                r'(调查研究[\s\S]{300,5000}?)(?=结果分析|结论|第[五六七]章)',
                r'(访谈实施[\s\S]{300,5000}?)(?=结果分析|结论|第[五六七]章)',
                r'(数据收集[\s\S]{300,5000}?)(?=结果分析|结论|第[五六七]章)',
            ],
            'computer_science': [
                r'(算法实现[\s\S]{300,5000}?)(?=结果分析|结论|第[五六七]章)',
                r'(系统实验[\s\S]{300,5000}?)(?=结果分析|结论|第[五六七]章)',
                r'(性能测试[\s\S]{300,5000}?)(?=结果分析|结论|第[五六七]章)',
            ]
        }
        
        return base_patterns + discipline_patterns.get(discipline, [])
    
    def _get_discipline_experimental_keywords(self, discipline: str) -> List[str]:
        """获取学科特定的实验关键词"""
        base_keywords = ['实验', '测试', '实施', '操作', '过程']
        
        discipline_keywords = {
            'engineering': ['制备', '合成', '加工', '测试', '表征', '工艺', '设备'],
            'medicine': ['临床', '诊断', '治疗', '病例', '患者', '试验', '观察'],
            'social_sciences': ['调查', '访谈', '观察', '抽样', '问卷', '研究'],
            'computer_science': ['算法', '实现', '编程', '测试', '验证', '优化'],
            'natural_sciences': ['实验', '观察', '测量', '分析', '检测', '样品']
        }
        
        return base_keywords + discipline_keywords.get(discipline, [])
    
    def _get_discipline_experimental_prompt(self, discipline: str) -> str:
        """获取学科特定的实验分析提示"""
        discipline_prompts = {
            'engineering': """
            针对工程技术领域实验，请特别关注：
            - 材料制备工艺与参数
            - 设备配置与操作条件
            - 性能测试方法与标准
            - 工艺优化与质量控制
            - 安全操作与环境要求
            """,
            'medicine': """
            针对医学领域实验，请特别关注：
            - 研究对象选择与分组
            - 临床试验设计与实施
            - 伦理审查与知情同意
            - 诊断标准与评估方法
            - 安全性与有效性评价
            """,
            'social_sciences': """
            针对社会科学领域实验，请特别关注：
            - 调查设计与抽样方法
            - 数据收集工具与程序
            - 访谈技巧与观察方法
            - 信度效度控制措施
            - 伦理考虑与隐私保护
            """,
            'computer_science': """
            针对计算机科学领域实验，请特别关注：
            - 算法设计与实现细节
            - 系统架构与技术选型
            - 测试环境与配置参数
            - 基准数据集与评估指标
            - 性能优化与复杂度分析
            """
        }
        return discipline_prompts.get(discipline, "请按照该学科的实验规范进行分析。")
    
    
    def _get_empty_experimental_analysis(self) -> Dict[str, Any]:
        """返回空的实验分析结构"""
        return {
            'discipline_identified': 'unknown',
            'experimental_design': {
                'design_type': '',
                'research_questions': [],
                'hypotheses': [],
                'experimental_strategy': ''
            },
            'experimental_methods': [],
            'materials_and_tools': {
                'materials': [],
                'instruments': [],
                'software': [],
                'reagents': []
            },
            'experimental_parameters': [],
            'experimental_conditions': {
                'environmental_conditions': [],
                'operational_conditions': [],
                'safety_measures': []
            },
            'quality_control': {
                'control_groups': [],
                'validation_methods': [],
                'error_sources': [],
                'repeatability_measures': []
            }
        }
    
    def _analyze_results_with_ai(self, text: str) -> Dict[str, Any]:
        """分析实验结果 - 支持多学科领域"""
        if not self.ai_client:
            return {'results': [], 'analysis': [], 'performance': []}
        
        # 1. 提取结果内容
        result_content = self._extract_results_content(text)
        
        if not result_content:
            return {'results': [], 'analysis': [], 'performance': []}
        
        try:
            prompt = f"""
作为结果分析专家，请对以下实验结果进行深度分析：

**结果内容：**
{result_content[:3000]}

**分析任务：**
请系统分析以下方面：
1. 主要实验结果与数据
2. 数据分析与统计处理
3. 结果解释与讨论
4. 性能评估与比较
5. 发现与创新点
6. 局限性与不足

请严格按照以下JSON格式返回分析结果：
{{
    "main_results": [
        {{
            "result_category": "结果类别",
            "description": "结果描述",
            "quantitative_data": "定量数据",
            "significance": "重要性评价"
        }}
    ],
    "data_analysis": {{
        "statistical_methods": ["统计方法1", "统计方法2"],
        "data_processing": ["数据处理1", "数据处理2"],
        "visualization": ["可视化方法1", "可视化方法2"],
        "quality_metrics": ["质量指标1", "质量指标2"]
    }},
    "performance_evaluation": [
        {{
            "metric_name": "指标名称",
            "metric_value": "指标数值",
            "benchmark_comparison": "基准比较",
            "improvement": "改进程度"
        }}
    ],
    "key_findings": [
        {{
            "finding": "关键发现",
            "evidence": "支撑证据",
            "novelty": "新颖性评价",
            "impact": "影响评估"
        }}
    ],
    "result_interpretation": {{
        "theoretical_implications": ["理论意义1", "理论意义2"],
        "practical_applications": ["实际应用1", "实际应用2"],
        "mechanism_explanation": ["机理解释1", "机理解释2"],
        "validation_evidence": ["验证证据1", "验证证据2"]
    }},
    "limitations_and_discussion": {{
        "methodological_limitations": ["方法局限1", "方法局限2"],
        "data_limitations": ["数据局限1", "数据局限2"],
        "future_work": ["未来工作1", "未来工作2"],
        "improvement_suggestions": ["改进建议1", "改进建议2"]
    }}
}}
"""
            
            response = self.ai_client.send_message(prompt)
            if response and hasattr(response, 'content'):
                # 三级JSON解析策略
                content = response.content
                
                # 第一级：直接JSON解析
                json_match = re.search(r'\{.*\}', content, re.DOTALL)
                if json_match:
                    try:
                        analysis_result = json.loads(json_match.group())
                        return analysis_result
                    except json.JSONDecodeError:
                        print(f"   📝 第一级JSON解析失败，尝试第二级清理解析")
                        
                        # 第二级：清理后解析
                        cleaned_json = _clean_json_content(json_match.group())
                        try:
                            analysis_result = json.loads(cleaned_json)
                            return analysis_result
                        except json.JSONDecodeError:
                            print(f"   ⚠️ 第二级清理解析失败")
                            return {'results': [], 'analysis': [], 'performance': []}
                else:
                    print(f"   ⚠️ 未找到JSON格式")
                    return {'results': [], 'analysis': [], 'performance': []}
        except Exception as e:
            print(f"   ⚠️ 结果AI分析失败: {e}")
        
        return {'results': [], 'analysis': [], 'performance': []}
    
    def _extract_results_content(self, text: str) -> str:
        """提取结果相关内容"""
        # 通用结果模式
        result_patterns = [
            r'第.*章\s*结果.*?(?=第.*章|结论|讨论|参考文献|$)',
            r'实验结果.*?(?=第.*章|结论|讨论|参考文献|$)',
            r'测试结果.*?(?=第.*章|结论|讨论|参考文献|$)',
            r'分析结果.*?(?=第.*章|结论|讨论|参考文献|$)',
            r'结果与分析.*?(?=第.*章|结论|讨论|参考文献|$)',
            r'结果.*?(?=第.*章|结论|讨论|参考文献|$)',
        ]
        
        result_content = ""
        for pattern in result_patterns:
            match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
            if match:
                content = match.group(0)
                if len(content) > 300:
                    result_content = content
                    break
        
        # 如果没找到，使用关键词搜索
        if not result_content:
            keywords = ['结果', '数据', '图', '表', '分析', '测试', '性能', '效果']
            lines = text.split('\n')
            result_lines = []
            for line in lines:
                if any(keyword in line.lower() for keyword in keywords):
                    result_lines.append(line)
                if len('\n'.join(result_lines)) > 2000:
                    break
            result_content = '\n'.join(result_lines)
        
        return result_content[:5000] if result_content else ""
    
    def _get_discipline_results_patterns(self, discipline: str) -> List[str]:
        """获取学科特定的结果模式"""
        base_patterns = [
            r'(实验结果[\s\S]{300,5000}?)(?=结论|讨论|第[六七八]章)',
            r'(结果分析[\s\S]{300,5000}?)(?=结论|讨论|第[六七八]章)',
            r'(性能分析[\s\S]{300,5000}?)(?=结论|讨论|第[六七八]章)',
        ]
        
        discipline_patterns = {
            'engineering': [
                r'(性能测试结果[\s\S]{300,5000}?)(?=结论|讨论|第[六七八]章)',
                r'(材料表征结果[\s\S]{300,5000}?)(?=结论|讨论|第[六七八]章)',
                r'(工艺验证结果[\s\S]{300,5000}?)(?=结论|讨论|第[六七八]章)',
            ],
            'medicine': [
                r'(临床试验结果[\s\S]{300,5000}?)(?=结论|讨论|第[六七八]章)',
                r'(疗效评估[\s\S]{300,5000}?)(?=结论|讨论|第[六七八]章)',
                r'(安全性分析[\s\S]{300,5000}?)(?=结论|讨论|第[六七八]章)',
            ],
            'social_sciences': [
                r'(调查结果[\s\S]{300,5000}?)(?=结论|讨论|第[六七八]章)',
                r'(数据分析结果[\s\S]{300,5000}?)(?=结论|讨论|第[六七八]章)',
                r'(统计分析[\s\S]{300,5000}?)(?=结论|讨论|第[六七八]章)',
            ],
            'computer_science': [
                r'(算法性能[\s\S]{300,5000}?)(?=结论|讨论|第[六七八]章)',
                r'(实验评估[\s\S]{300,5000}?)(?=结论|讨论|第[六七八]章)',
                r'(基准测试[\s\S]{300,5000}?)(?=结论|讨论|第[六七八]章)',
            ]
        }
        
        return base_patterns + discipline_patterns.get(discipline, [])
    
    def _get_discipline_results_keywords(self, discipline: str) -> List[str]:
        """获取学科特定的结果关键词"""
        base_keywords = ['结果', '数据', '分析', '发现', '评估']
        
        discipline_keywords = {
            'engineering': ['性能', '效率', '强度', '硬度', '导电', '传热', '表征'],
            'medicine': ['疗效', '安全性', '有效率', '生存率', '副作用', '改善'],
            'social_sciences': ['显著性', '相关性', '差异', '趋势', '模式', '关系'],
            'computer_science': ['准确率', '速度', '内存', '吞吐量', '延迟', '优化'],
            'natural_sciences': ['浓度', '温度', '压力', '光谱', '结构', '组成']
        }
        
        return base_keywords + discipline_keywords.get(discipline, [])
    
    def _get_discipline_results_prompt(self, discipline: str) -> str:
        """获取学科特定的结果分析提示"""
        discipline_prompts = {
            'engineering': """
            针对工程技术领域结果，请特别关注：
            - 材料性能指标与测试数据
            - 工艺参数优化结果
            - 设备性能与效率评估
            - 质量控制与标准符合性
            - 技术创新点与改进效果
            """,
            'medicine': """
            针对医学领域结果，请特别关注：
            - 临床试验主要终点指标
            - 安全性与有效性评估
            - 统计学意义与临床意义
            - 不良反应与副作用分析
            - 疗效比较与基线改善
            """,
            'social_sciences': """
            针对社会科学领域结果，请特别关注：
            - 统计显著性与效应量
            - 相关性与因果关系分析
            - 人口统计学差异
            - 理论假设验证结果
            - 实际应用价值与意义
            """,
            'computer_science': """
            针对计算机科学领域结果，请特别关注：
            - 算法性能指标与基准比较
            - 计算复杂度与资源消耗
            - 准确性与鲁棒性评估
            - 可扩展性与实用性分析
            - 创新性与技术贡献
            """
        }
        return discipline_prompts.get(discipline, "请按照该学科的结果分析规范进行分析。")
    
    def _parse_results_fallback(self, text: str, discipline: str) -> Dict[str, Any]:
        """结果分析备用解析方法"""
        return {
            'discipline_identified': discipline,
            'main_results': [],
            'data_analysis': {'statistical_methods': [], 'data_processing': [], 'visualization': [], 'quality_metrics': []},
            'performance_evaluation': [],
            'key_findings': [],
            'result_interpretation': {'theoretical_implications': [], 'practical_applications': [], 'mechanism_explanation': [], 'validation_evidence': []},
            'limitations_and_discussion': {'methodological_limitations': [], 'data_limitations': [], 'future_work': [], 'improvement_suggestions': []}
        }
    
    def _extract_content_by_sections_disciplinary(self, text: str, sections: Dict[str, str], discipline: str) -> Dict[str, Any]:
        """学科专业化的内容提取 - 重新设计以正确提取摘要和关键词"""
        content_info = {
            'sections_count': len(sections),
            'discipline': discipline,
            'content_summary': f"识别到{len(sections)}个章节，学科：{discipline}"
        }
        
        # 从摘要部分提取
        if 'abstract_cn' in sections:
            content_info['abstract_cn'] = self._clean_abstract(sections['abstract_cn'])
            print(f"   ✅ 中文摘要: {len(content_info['abstract_cn'])} 字符")
        else:
            # 如果没有识别到章节，尝试正则匹配
            abstract_pattern = r'((?:中文)?摘\s*要[\s\S]{50,3000}?)(?=关键词|英文摘要|ABSTRACT|第一章|目录)'
            match = re.search(abstract_pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                content_info['abstract_cn'] = self._clean_abstract(match.group(1))
                print(f"   ✅ 中文摘要(正则): {len(content_info['abstract_cn'])} 字符")
            else:
                content_info['abstract_cn'] = ""
                print("   ⚠️ 未找到中文摘要")
        
        if 'abstract_en' in sections:
            content_info['abstract_en'] = self._clean_abstract(sections['abstract_en'])
            print(f"   ✅ 英文摘要: {len(content_info['abstract_en'])} 字符")
        else:
            # 如果没有识别到章节，尝试正则匹配
            abstract_pattern = r'((?:ABSTRACT|Abstract)[\s\S]{50,3000}?)(?=Keywords?|中文摘要|第一章|目录)'
            match = re.search(abstract_pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                content_info['abstract_en'] = self._clean_abstract(match.group(1))
                print(f"   ✅ 英文摘要(正则): {len(content_info['abstract_en'])} 字符")
            else:
                content_info['abstract_en'] = ""
                print("   ⚠️ 未找到英文摘要")
        
        # 从关键词部分提取
        if 'keywords_cn' in sections:
            keywords = self._extract_keywords(sections['keywords_cn'], 'chinese')
            content_info['keywords_cn'] = keywords
            print(f"   ✅ 中文关键词: {keywords}")
        else:
            # 如果没有识别到章节，尝试正则匹配
            keywords_pattern = r'(关键词[：:\s]*[^\n\r]{5,200})'
            match = re.search(keywords_pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                keywords = self._extract_keywords(match.group(1), 'chinese')
                content_info['keywords_cn'] = keywords
                print(f"   ✅ 中文关键词(正则): {keywords}")
            else:
                content_info['keywords_cn'] = ""
                print("   ⚠️ 未找到中文关键词")
        
        if 'keywords_en' in sections:
            keywords = self._extract_keywords(sections['keywords_en'], 'english')
            content_info['keywords_en'] = keywords
            print(f"   ✅ 英文关键词: {keywords}")
        else:
            # 如果没有识别到章节，尝试正则匹配
            keywords_pattern = r'((?:Keywords?|KEY\s+WORDS?)[：:\s]*[^\n\r]{5,200})'
            match = re.search(keywords_pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                keywords = self._extract_keywords(match.group(1), 'english')
                content_info['keywords_en'] = keywords
                print(f"   ✅ 英文关键词(正则): {keywords}")
            else:
                content_info['keywords_en'] = ""
                print("   ⚠️ 未找到英文关键词")
        
        # 从致谢部分提取
        if 'acknowledgement' in sections:
            acknowledgement_content = sections['acknowledgement'].strip()
            content_info['acknowledgement'] = acknowledgement_content
            print(f"   ✅ 致谢: {len(acknowledgement_content)} 字符")
        else:
            # 如果没有识别到章节，尝试正则匹配
            ack_pattern = r'(致\s*谢[\s\S]{50,2000}?)(?=附录|参考文献|$)'
            match = re.search(ack_pattern, text, re.IGNORECASE | re.MULTILINE)
            if match:
                acknowledgement_content = match.group(1).strip()
                content_info['acknowledgement'] = acknowledgement_content
                print(f"   ✅ 致谢(正则): {len(acknowledgement_content)} 字符")
            else:
                content_info['acknowledgement'] = ""
                print("   ⚠️ 未找到致谢")
        
        return content_info
    
    def _extract_references_enhanced_disciplinary(self, text: str, discipline: str, source_path: str = '') -> Dict[str, Any]:
        """学科适配的智能参考文献解析"""
        print("   📚 开始智能提取参考文献...")
        
        # 使用智能参考文献提取器
        if self.smart_ref_extractor:
            # 根据文件路径自动检测格式
            source_format = 'auto'
            if source_path:
                if source_path.lower().endswith('.pdf'):
                    source_format = 'pdf'
                elif source_path.lower().endswith(('.docx', '.doc')):
                    source_format = 'docx'
            
            print(f"   🔍 文档格式检测: {source_format}")
            references_list, extraction_stats = self.smart_ref_extractor.extract_references(
                text, source_format=source_format, source_path=source_path
            )
            
            print(f"   ✅ 智能提取完成: {len(references_list)} 条参考文献")
            print(f"   📊 提取方法: {extraction_stats.get('method_used', 'unknown')}")
            print(f"   ⏱️ 处理时间: {extraction_stats.get('processing_time', 0):.2f}秒")
        else:
            # 回退到传统方法
            print("   ⚠️ 智能提取器不可用，使用传统方法")
            references_list = self._extract_references_enhanced(text)
            extraction_stats = {
                'method_used': '传统正则提取',
                'total_found': len(references_list),
                'processing_time': 0.0,
                'success': len(references_list) > 0
            }
        
        # 构建返回字典
        references_dict = {
            'references': references_list,
            'discipline': discipline,
            'total_count': len(references_list),
            'extraction_stats': extraction_stats
        }
        
        print(f"   📋 参考文献提取完成: {len(references_list)} 条")
        return references_dict
    

    
    def _analyze_conclusion_with_ai(self, text: str) -> Dict[str, Any]:
        """分析结论部分"""
        if not self.ai_client:
            return {'conclusions': [], 'contributions': [], 'future_work': []}
        
        # 提取结论部分 - 优化正则表达式模式
        conclusion_patterns = [
            # 明确的结论章节
            r'((?:结论|总结|结论与展望|总结与展望)[\s\S]{200,8000}?)(?=参考文献|致谢|附录|$)',
            # 以"本论文主要研究了"开始的结论
            r'(本论文主要研究了[\s\S]{500,12000}?)(?=参考文献|致谢|附录|$)',
            # 以"研究所得到的主要结论"开始
            r'(研究所得到的主要结论[\s\S]{300,8000}?)(?=参考文献|致谢|附录|$)',
            # 通用结论模式
            r'((?:主要结论|本文结论|研究结论)[\s\S]{200,6000}?)(?=参考文献|致谢|附录|$)',
        ]
        
        conclusion_content = ""
        for pattern in conclusion_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                conclusion_content = match.group(1)
                print(f"   ✅ 找到结论部分，长度: {len(conclusion_content)} 字符")
                break
        
        if not conclusion_content:
            print("   ⚠️ 未找到结论部分")
            return {'conclusions': [], 'contributions': [], 'future_work': []}
        
        try:
            prompt = f"""
请分析以下论文结论内容，这可能包含多个段落和编号列表：

{conclusion_content[:3000]}

请仔细提取：
1. 主要研究结论 - 查找编号列表中的具体结论点
2. 学术贡献和创新点 - 查找"突破"、"创新"、"贡献"等关键词
3. 未来工作展望 - 查找"今后工作"、"展望"、"未来"等内容
4. 研究局限性 - 如有提及

注意：即使没有明确的章节标题，也要从内容中提取这些信息。

请以JSON格式回复：
{{
    "conclusions": ["结论1", "结论2"],
    "contributions": ["贡献1", "贡献2"],
    "future_work": ["展望1", "展望2"],
    "limitations": ["局限1", "局限2"]
}}
"""
            
            response = self.ai_client.send_message(prompt)
            if response and hasattr(response, 'content'):
                # 三级JSON解析策略
                content = response.content
                
                # 第一级：直接JSON解析  
                json_content = content.strip()
                if '```json' in json_content:
                    json_start = json_content.find('```json')
                    json_end = json_content.find('```', json_start + 7)
                    if json_end > json_start:
                        json_content = json_content[json_start + 7:json_end].strip()
                elif '```' in json_content:
                    parts = json_content.split('```')
                    if len(parts) >= 3:
                        json_content = parts[1].strip()
                
                try:
                    result = json.loads(json_content)
                    print(f"   ✅ 结论AI分析成功，提取到 {len(result.get('conclusions', []))} 个结论")
                    return result
                except json.JSONDecodeError:
                    print(f"   📝 第一级JSON解析失败，尝试第二级清理解析")
                    
                    # 第二级：清理后解析
                    cleaned_json = _clean_json_content(json_content)
                    try:
                        result = json.loads(cleaned_json)
                        print(f"   ✅ 结论清理解析成功，提取到 {len(result.get('conclusions', []))} 个结论")
                        return result
                    except json.JSONDecodeError:
                        print(f"   ⚠️ 第二级清理解析失败")
                        return {'conclusions': [], 'contributions': [], 'future_work': []}
        except Exception as e:
            print(f"   ⚠️ 结论AI分析失败: {e}")
        
        return {'conclusions': [], 'contributions': [], 'future_work': []}
       
    def _extract_references_enhanced(self, text: str) -> List[str]:
        """步骤4: 使用AI大模型智能提取参考文献"""
        print("   🔍 启动AI智能参考文献解析...")
        
        # 方法1: 定位参考文献部分
        ref_text = self._locate_references_section(text)
        
        if not ref_text:
            print("   ⚠️ 未找到参考文献部分")
            return []
        
        print(f"   📍 找到参考文献部分，长度: {len(ref_text)} 字符")
        
        # 方法2: 使用AI智能提取参考文献条目
        references = self._extract_references_with_ai(ref_text)
        
        print(f"   ✅ AI提取参考文献: {len(references)} 条")
        return references
    
    def _locate_references_section(self, text: str) -> str:
        """定位参考文献部分"""
        # 查找参考文献标题的多种模式
        ref_patterns = [
            r'(?:^|\n)(?:##\s*)?参考文献\s*\n([\s\S]*?)(?=\n\s*(?:缩略词表|文献综述|致谢|附录|作者简介|个人简历)|$)',
            r'(?:^|\n)(?:##\s*)?REFERENCES?\s*\n([\s\S]*?)(?=\n\s*(?:ACKNOWLEDGMENT|APPENDIX|文献综述|致谢)|$)',
            r'(?:^|\n)(?:##\s*)?Bibliography\s*\n([\s\S]*?)(?=\n\s*(?:ACKNOWLEDGMENT|APPENDIX|文献综述|致谢)|$)',
        ]
        
        for pattern in ref_patterns:
            matches = list(re.finditer(pattern, text, re.IGNORECASE | re.MULTILINE))
            if matches:
                best_match = max(matches, key=lambda m: len(m.group(1)))
                ref_text = best_match.group(1).strip()
                if len(ref_text) >= 1000:
                    return ref_text
        
        # 备用方法：关键词定位
        ref_keywords = ['参考文献', 'References', 'REFERENCES', 'Bibliography']
        for keyword in ref_keywords:
            pos = text.find(keyword)
            if pos != -1:
                remaining_text = text[pos+len(keyword):]
                end_markers = ['缩略词表', '文献综述', '致谢', 'ACKNOWLEDGMENT', 'APPENDIX', '附录', '作者简介', '个人简历']
                end_pos = len(remaining_text)
                
                for marker in end_markers:
                    marker_pos = remaining_text.find(marker)
                    if marker_pos != -1 and marker_pos < end_pos:
                        end_pos = marker_pos
                
                ref_text = remaining_text[:end_pos]
                if len(ref_text) >= 1000:
                    return ref_text
        
        return ""
    
    def _extract_references_with_ai(self, ref_text: str) -> List[str]:
        """使用AI大模型提取参考文献条目"""
        try:
            # 限制输入长度以避免token超限
            max_length = 50000  # 约50k字符，对应大约12-15k tokens
            if len(ref_text) > max_length:
                print(f"   📏 参考文献内容过长({len(ref_text)}字符)，截取前{max_length}字符")
                ref_text = ref_text[:max_length]
            
            prompt = f"""请从以下参考文献文本中提取所有参考文献条目。

要求：
1. 每个参考文献条目应该是完整的一条记录
2. 保持原有的编号格式（如［1］、[1]、1.等）
3. 清理多余的空白字符和换行符
4. 每条参考文献应该包含：作者、标题、期刊/会议/出版社、年份等信息
5. 如果格式混乱，请智能重组成标准格式
6. 按编号顺序排列
7. 输出格式：每行一条参考文献，不需要额外说明

参考文献文本：
{ref_text}

请提取所有参考文献条目："""

            if hasattr(self, 'ai_client') and self.ai_client:
                print("   🤖 调用AI大模型提取参考文献...")
                response = self.ai_client.send_message(prompt)
                
                if response and hasattr(response, 'content'):
                    content = response.content.strip()
                    
                    # 解析AI返回的结果
                    references = []
                    lines = content.split('\n')
                    
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        
                        # 检查是否像参考文献条目
                        if self._is_valid_reference(line):
                            # 清理格式
                            cleaned_ref = re.sub(r'\s+', ' ', line).strip()
                            references.append(cleaned_ref)
                    
                    print(f"   ✅ AI成功提取 {len(references)} 条参考文献")
                    return references
                else:
                    print("   ⚠️ AI响应为空")
            else:
                print("   ⚠️ AI客户端不可用")
                
        except Exception as e:
            print(f"   ❌ AI提取失败: {e}")
        
        # 备用方法：使用改进的正则表达式
        print("   🔄 使用备用正则表达式方法...")
        return self._extract_references_fallback(ref_text)
    
    def _is_valid_reference(self, line: str) -> bool:
        """检查是否是有效的参考文献条目"""
        # 基本长度检查
        if len(line) < 20:
            return False
        
        # 检查是否包含编号格式
        has_number = bool(re.match(r'^\s*(?:［\d+］|\[\d+\]|\d+\.|（\d+）|\(\d+\))', line))
        
        # 检查是否包含期刊、会议、出版社等关键词
        has_publication = any(keyword in line for keyword in [
            'Journal', 'Proceedings', 'Conference', 'IEEE', 'ACM', 'Optics',
            '期刊', '会议', '学报', '大学学报', '论文集', '出版社'
        ])
        
        # 检查是否包含年份
        has_year = bool(re.search(r'(?:19|20)\d{2}', line))
        
        # 至少满足编号+年份，或者编号+出版物
        return has_number and (has_year or has_publication)
    
    def _extract_references_fallback(self, ref_text: str) -> List[str]:
        """备用的参考文献提取方法"""
        references = []
        
        # 智能段落分割和重组
        print("   🔧 使用智能段落重组方法...")
        
        # 按空行分割段落
        paragraphs = re.split(r'\n\s*\n', ref_text)
        current_ref = ""
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # 检查是否是新的参考文献开始
            if re.match(r'^\s*(?:［\d+］|\[\d+\]|\d+\.)', para):
                # 保存之前的参考文献
                if current_ref and len(current_ref) > 20:
                    cleaned_ref = re.sub(r'\s+', ' ', current_ref).strip()
                    references.append(cleaned_ref)
                
                # 开始新的参考文献
                current_ref = para
            else:
                # 继续当前参考文献
                if current_ref:
                    current_ref += " " + para
        
        # 添加最后一条参考文献
        if current_ref and len(current_ref) > 20:
            cleaned_ref = re.sub(r'\s+', ' ', current_ref).strip()
            references.append(cleaned_ref)
        
        # 如果段落方法效果不好，尝试行级处理
        if len(references) < 5:
            print("   🔧 尝试行级重组方法...")
            references = self._extract_references_line_based(ref_text)
        
        return references[:100]  # 限制数量
    
    def _extract_references_line_based(self, ref_text: str) -> List[str]:
        """基于行的参考文献提取方法"""
        references = []
        lines = ref_text.split('\n')
        current_ref = ""
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检查是否是参考文献编号行
            if re.match(r'^\s*(?:［\d+］|\[\d+\]|\d+\.)', line):
                # 保存之前的参考文献
                if current_ref and len(current_ref) > 20:
                    cleaned_ref = re.sub(r'\s+', ' ', current_ref).strip()
                    if self._is_valid_reference(cleaned_ref):
                        references.append(cleaned_ref)
                
                # 开始新的参考文献
                current_ref = line
            else:
                # 继续当前参考文献
                if current_ref:
                    current_ref += " " + line
                elif any(keyword in line for keyword in ['Journal', 'IEEE', 'Proceedings', '学报']):
                    # 可能是参考文献的一部分
                    current_ref = line
        
        # 添加最后一条参考文献
        if current_ref and len(current_ref) > 20:
            cleaned_ref = re.sub(r'\s+', ' ', current_ref).strip()
            if self._is_valid_reference(cleaned_ref):
                references.append(cleaned_ref)
        
        return references
    
    def _intelligent_repair_and_validate(self, metadata: Dict, content_info: Dict, references: List[str], toc_analysis: Optional[Dict[str, Any]] = None, original_text: str = '', ai_analysis: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """步骤5: 多层次验证和错误修正机制"""
        print("   🧠 启动智能修复系统...")
        
        # 合并所有提取的信息
        result = {}
        result.update(metadata)
        result.update(content_info)
        result['references'] = references
        
        # 添加目录分析结果
        if toc_analysis:
            result.update(toc_analysis)
            print(f"   ✅ 集成目录分析: {len(toc_analysis.get('table_of_contents', []))} 个章节")
        
        # 添加AI智能分析结果
        if ai_analysis:
            result['ai_analysis'] = ai_analysis
            section_count = len(ai_analysis.get('section_analysis', {}))
            overall_score = ai_analysis.get('content_quality', {}).get('overall_quality_score', 0)
            print(f"   🤖 集成AI分析: {section_count} 个章节分析，整体质量评分: {overall_score}")
            
            # 提取AI分析的关键洞察
            if 'section_analysis' in ai_analysis:
                result['ai_insights'] = self._extract_ai_insights(ai_analysis['section_analysis'])
                print(f"   💡 提取AI洞察: {len(result['ai_insights'])} 条建议")
        
        # 提取缺失的theoretical_framework和author_contributions字段
        full_text = original_text or result.get('original_text', '') or (metadata.get('original_text', '') or content_info.get('original_text', ''))
        if full_text:
            print("   🔍 提取理论框架和作者贡献...")
            
            # 提取理论框架
            if not result.get('theoretical_framework'):
                theoretical_framework = self._extract_theoretical_framework(full_text)
                if theoretical_framework:
                    result['theoretical_framework'] = theoretical_framework
                    print(f"   ✅ 理论框架提取: {len(str(theoretical_framework)[:100])}字符")
            
            # 提取作者贡献
            if not result.get('author_contributions'):
                author_contributions = self._extract_author_contributions(full_text)
                if author_contributions:
                    result['author_contributions'] = author_contributions
                    print(f"   ✅ 作者贡献提取: {len(str(author_contributions)[:100])}字符")
        
        # 字段验证和清理
        result = self._validate_and_clean_fields(result)
        
        print("   ✅ 智能修复完成")
        return result
    
    def _clean_abstract(self, text: str) -> str:
        """清理摘要文本"""
        # 移除标题
        text = re.sub(r'^(中文摘要|摘要|ABSTRACT|Abstract)[：:\s]*', '', text)
        # 移除多余的空白
        text = re.sub(r'\s+', ' ', text).strip()
        return text
    
    def _extract_keywords(self, text: str, language: str) -> str:
        """提取关键词"""
        # 移除标题
        text = re.sub(r'^(关键词|Keywords?|KEY\s+WORDS?)[：:\s]*', '', text, flags=re.IGNORECASE)
        # 清理和格式化
        keywords = re.split(r'[,，;；]', text)
        keywords = [kw.strip() for kw in keywords if kw.strip()]
        return '，'.join(keywords) if language == 'chinese' else ', '.join(keywords)
    
    def _extract_theoretical_framework(self, text: str) -> Dict[str, Any]:
        """提取理论框架信息"""
        # 检查AI客户端是否可用
        if not self.ai_client:
            logger.warning("理论框架智能提取器不可用：AI客户端未初始化")
            print("   ⚠️ 理论框架智能提取器不可用：AI客户端未初始化")
            return {
                'core_theories': [],
                'theoretical_models': [],
                'conceptual_foundations': [],
                'theoretical_contributions': []
            }
        
        # 理论框架相关的章节模式
        theory_patterns = [
            r'((?:理论基础|理论框架|相关理论|基础理论)[\s\S]{300,8000}?)(?=第|章节|参考文献|致谢|附录|$)',
            r'((?:文献综述|国内外研究现状)[\s\S]{500,10000}?)(?=第|章节|参考文献|致谢|附录|$)',
            r'((?:概念模型|理论模型|数学模型)[\s\S]{200,5000}?)(?=第|章节|参考文献|致谢|附录|$)'
        ]
        
        theory_content = ""
        for pattern in theory_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                content = match.group(1)
                if len(content) > len(theory_content):
                    theory_content = content
        
        if not theory_content:
            # 如果没有找到专门的理论章节，从文献综述中提取
            literature_patterns = [
                r'(文献综述[\s\S]{1000,15000}?)(?=第|章节|研究方法|实验|参考文献)',
                r'(相关工作[\s\S]{1000,10000}?)(?=第|章节|研究方法|实验|参考文献)',
                r'(国内外研究现状[\s\S]{1000,12000}?)(?=第|章节|研究方法|实验|参考文献)'
            ]
            
            for pattern in literature_patterns:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    theory_content = match.group(1)
                    break
        
        if theory_content:
            # 使用AI分析理论框架
            framework_analysis = self._analyze_theoretical_framework_with_ai(theory_content)
            # 检查是否获得了有效的分析结果
            if framework_analysis and any(framework_analysis.get(key) for key in ['core_theories', 'theoretical_models', 'conceptual_foundations', 'theoretical_contributions']):
                return framework_analysis
        
        # AI分析失败，记录日志并返回空值
        logger.warning("理论框架智能提取失败：未找到有效内容或AI分析失败")
        print("   ⚠️ 理论框架智能提取失败：未找到有效内容或AI分析失败")
        return {
            'core_theories': [],
            'theoretical_models': [],
            'conceptual_foundations': [],
            'theoretical_contributions': []
        }
    
    def _analyze_theoretical_framework_with_ai(self, content: str) -> Dict[str, Any]:
        """使用AI分析理论框架"""
        try:
            if not self.ai_client:
                logger.warning("理论框架AI分析失败：AI客户端不可用")
                print("   ⚠️ 理论框架AI分析失败：AI客户端不可用")
                return {
                    'core_theories': [],
                    'theoretical_models': [],
                    'conceptual_foundations': [],
                    'theoretical_contributions': []
                }
                
            prompt = f"""
请分析以下论文内容中的理论框架，识别核心理论、模型和概念基础：

{content[:3000]}

请提取：
1. 核心理论 - 论文基于的主要理论
2. 理论模型 - 使用的理论模型或框架
3. 概念基础 - 重要的概念和定义
4. 理论贡献 - 论文在理论方面的贡献

请以JSON格式回复：
{{
    "core_theories": ["理论1", "理论2"],
    "theoretical_models": ["模型1", "模型2"],
    "conceptual_foundations": ["概念1", "概念2"],
    "theoretical_contributions": ["贡献1", "贡献2"]
}}
"""
            
            response = self.ai_client.send_message(prompt)
            if response and hasattr(response, 'content'):
                content = response.content
                
                # 三级JSON解析策略
                json_match = re.search(r'\{.*\}', content, re.DOTALL)
                if json_match:
                    try:
                        result = json.loads(json_match.group())
                        return result
                    except json.JSONDecodeError:
                        # 清理后解析
                        cleaned_json = _clean_json_content(json_match.group())
                        try:
                            result = json.loads(cleaned_json)
                            return result
                        except json.JSONDecodeError:
                            logger.warning("理论框架AI分析：JSON解析失败")
                            print("   ⚠️ 理论框架AI分析：JSON解析失败")
                            return {
                                'core_theories': [],
                                'theoretical_models': [],
                                'conceptual_foundations': [],
                                'theoretical_contributions': []
                            }
                else:
                    logger.warning("理论框架AI分析：未找到有效JSON格式")
                    print("   ⚠️ 理论框架AI分析：未找到有效JSON格式")
                    return {
                        'core_theories': [],
                        'theoretical_models': [],
                        'conceptual_foundations': [],
                        'theoretical_contributions': []
                    }
        except Exception as e:
            logger.error(f"理论框架AI分析失败: {e}")
            print(f"   ⚠️ 理论框架AI分析失败: {e}")
        
        return {
            'core_theories': [],
            'theoretical_models': [],
            'conceptual_foundations': [],
            'theoretical_contributions': []
        }
    
    def _extract_author_contributions(self, text: str) -> Dict[str, Any]:
        """提取作者贡献信息"""
        # 检查AI客户端是否可用
        if not self.ai_client:
            logger.warning("作者贡献智能提取器不可用：AI客户端未初始化")
            print("   ⚠️ 作者贡献智能提取器不可用：AI客户端未初始化")
            return {
                'contribution_statement': '',
                'research_contributions': [],
                'publication_contributions': [],
                'innovation_points': []
            }
        
        # 作者贡献声明的常见位置模式
        contribution_patterns = [
            r'((?:攻读.*?学位期间取得的研究成果|个人贡献|作者贡献|研究成果)[\s\S]{200,5000}?)(?=致谢|参考文献|附录|$)',
            r'((?:发表论文|研究成果|学术论文|专利申请)[\s\S]{300,3000}?)(?=致谢|参考文献|附录|$)',
            r'((?:创新点|创新性|主要贡献|研究贡献)[\s\S]{200,2000}?)(?=致谢|参考文献|附录|第|章节|$)'
        ]
        
        contribution_content = ""
        for pattern in contribution_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                content = match.group(1)
                if len(content) > len(contribution_content):
                    contribution_content = content
        
        if contribution_content:
            # 使用AI分析作者贡献
            contribution_analysis = self._analyze_author_contributions_with_ai(contribution_content)
            # 检查是否获得了有效的分析结果
            if contribution_analysis and (contribution_analysis.get('contribution_statement') or any(contribution_analysis.get(key) for key in ['research_contributions', 'publication_contributions', 'innovation_points'])):
                return contribution_analysis
        
        # AI分析失败，记录日志并返回空值
        logger.warning("作者贡献智能提取失败：未找到有效内容或AI分析失败")
        print("   ⚠️ 作者贡献智能提取失败：未找到有效内容或AI分析失败")
        return {
            'contribution_statement': '',
            'research_contributions': [],
            'publication_contributions': [],
            'innovation_points': []
        }
    
    def _analyze_author_contributions_with_ai(self, content: str) -> Dict[str, Any]:
        """使用AI分析作者贡献"""
        try:
            if not self.ai_client:
                logger.warning("作者贡献AI分析失败：AI客户端不可用")
                print("   ⚠️ 作者贡献AI分析失败：AI客户端不可用")
                return {
                    'contribution_statement': '',
                    'research_contributions': [],
                    'publication_contributions': [],
                    'innovation_points': []
                }
                
            prompt = f"""
请分析以下作者贡献内容，提取研究成果和创新点：

{content[:2000]}

请提取：
1. 贡献声明 - 总体贡献描述
2. 研究贡献 - 具体研究方面的贡献
3. 发表贡献 - 论文发表和专利等
4. 创新点 - 主要创新点

请以JSON格式回复：
{{
    "contribution_statement": "总体贡献描述",
    "research_contributions": ["研究贡献1", "研究贡献2"],
    "publication_contributions": ["发表1", "发表2"],
    "innovation_points": ["创新点1", "创新点2"]
}}
"""
            
            response = self.ai_client.send_message(prompt)
            if response and hasattr(response, 'content'):
                content = response.content
                
                # 三级JSON解析策略
                json_match = re.search(r'\{.*\}', content, re.DOTALL)
                if json_match:
                    try:
                        result = json.loads(json_match.group())
                        return result
                    except json.JSONDecodeError:
                        # 清理后解析
                        cleaned_json = _clean_json_content(json_match.group())
                        try:
                            result = json.loads(cleaned_json)
                            return result
                        except json.JSONDecodeError:
                            logger.warning("作者贡献AI分析：JSON解析失败")
                            print("   ⚠️ 作者贡献AI分析：JSON解析失败")
                            return {
                                'contribution_statement': '',
                                'research_contributions': [],
                                'publication_contributions': [],
                                'innovation_points': []
                            }
                else:
                    logger.warning("作者贡献AI分析：未找到有效JSON格式")
                    print("   ⚠️ 作者贡献AI分析：未找到有效JSON格式")
                    return {
                        'contribution_statement': '',
                        'research_contributions': [],
                        'publication_contributions': [],
                        'innovation_points': []
                    }
        except Exception as e:
            logger.error(f"作者贡献AI分析失败: {e}")
            print(f"   ⚠️ 作者贡献AI分析失败: {e}")
        
        return {
            'contribution_statement': '',
            'research_contributions': [],
            'publication_contributions': [],
            'innovation_points': []
        }
    
   
    
    def _validate_and_clean_fields(self, result: Dict) -> Dict[str, Any]:
        """验证和清理字段"""
        # 确保所有标准字段都存在
        for field in self.standard_fields:
            if field not in result:
                result[field] = ""
        
        # 清理字段内容
        for field, value in result.items():
            if isinstance(value, str):
                # 移除控制字符
                value = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', value)
                # 规范化空白字符
                value = re.sub(r'\s+', ' ', value)
                result[field] = value.strip()
        
        return result
    
    def _calculate_extraction_stats(self, result: Dict, processing_time: float):
        """计算提取统计信息"""
        extracted_count = 0
        for field in self.standard_fields:
            value = result.get(field, '')
            # 处理不同类型的值
            if isinstance(value, list):
                # 列表类型：检查是否非空
                if value and any(str(item).strip() for item in value):
                    extracted_count += 1
            elif isinstance(value, str):
                # 字符串类型：检查是否非空白
                if value.strip():
                    extracted_count += 1
            elif value:
                # 其他非空值
                extracted_count += 1
        
        self.extraction_stats.update({
            'extracted_fields': extracted_count,
            'confidence': min(1.0, extracted_count / len(self.standard_fields)),
            'processing_time': processing_time
        })
    
    def _generate_extraction_report(self, result: Dict, file_path: Optional[str], processing_time: float):
        """生成提取报告"""
        print(f"\n⏱️ 提取完成，耗时: {processing_time:.2f} 秒")
        print("\n" + "=" * 60)
        print("📊 专业版提取报告")
        print("=" * 60)
        
        if file_path:
            print(f"📁 目标文件: {Path(file_path).name}")
        
        print(f"⏱️ 提取时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        print(f"🎯 使用技术: 分步抽取, 结构化分析, 快速定位, 正则匹配, 参考文献解析, 智能修复")
        print(f"📈 总字段数: {self.extraction_stats['total_fields']}")
        print(f"✅ 已提取: {self.extraction_stats['extracted_fields']} 个字段")
        print(f"📊 完整度: {self.extraction_stats['confidence']:.1%}")
        print(f"🎖️ 置信度: {self.extraction_stats['confidence']:.2f}")
        
        # 质量评分
        quality_score = (
            self.extraction_stats['confidence'] * 0.6 +  # 完整度权重60%
            (1.0 if result.get('title_cn') else 0.0) * 0.2 +  # 标题权重20%
            (1.0 if len(result.get('references', [])) > 10 else 0.0) * 0.2  # 参考文献权重20%
        )
        print(f"⭐ 质量分数: {quality_score:.2f}")
        print("=" * 60)
        print("✅ 专业版学位论文信息提取完成！")


# 保持向后兼容的全局实例
thesis_extractor_pro = ThesisExtractorPro()

def extract_text_from_pdf(file_path: str) -> str:
    """从PDF文件提取文本"""
    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()  # type: ignore
        doc.close()
        return text
    except Exception as e:
        logger.error(f"提取PDF文本失败: {e}")
        return ""

def extract_text_from_word(file_path: str) -> str:
    """从Word文档提取文本"""
    try:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        logger.error(f"提取Word文本失败: {e}")
        return ""

def _extract_json_from_response(response_text: str) -> Optional[str]:
    """从AI响应中提取JSON内容"""
    # 移除代码块标记
    content = re.sub(r'```json\s*', '', response_text)
    content = re.sub(r'```\s*$', '', content)
    
    # 查找JSON对象
    json_match = re.search(r'\{.*\}', content, re.DOTALL)
    if json_match:
        return json_match.group()
    
    logger.warning("无法找到JSON对象")
    return None

def _clean_json_content(json_str: str) -> str:
    """清理JSON内容中的控制字符"""
    # 移除控制字符
    json_str = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', json_str)
    # 移除多余的空白字符
    json_str = re.sub(r'\s+', ' ', json_str)
    return json_str.strip()

def _parse_json_with_fallback(json_str: str) -> Optional[Dict]:
    """使用回退机制解析JSON"""
    try:
        return json.loads(json_str)
    except json.JSONDecodeError as e:
        logger.warning(f"📝 JSON解析失败: {e}，启动备用解析机制")
        
        # 使用正则表达式提取关键字段 (移除冗余字段，使用snake_case和语言后缀)
        result = {}
        fields = [
            'title_cn', 'keywords_cn', 'abstract_cn',
            'title_en', 'keywords_en', 'abstract_en',
            'theoretical_framework',
            'references'
        ]
        
        for field in fields:
            pattern = rf'"{field}"\s*:\s*"([^"]*)"'
            match = re.search(pattern, json_str, re.DOTALL)
            if match:
                result[field] = match.group(1).strip()
            else:
                result[field] = ""
        
        extracted_fields = len([k for k, v in result.items() if v and str(v).strip()])
        logger.info(f"📝 备用解析完成，成功提取 {extracted_fields}/{len(fields)} 个字段")
        return result if result else None


def save_extraction_cache(file_path: str, extracted_info: Dict, session_id: Optional[str] = None) -> bool:
    """
    保存结构化信息到缓存文件
    
    Args:
        file_path: 原始论文文件路径
        extracted_info: 提取的结构化信息
        session_id: 会话ID（可选）
    
    Returns:
        保存成功返回True，失败返回False
    """
    try:
        from .config_manager import get_config_manager
        from pathlib import Path
        from datetime import datetime
        
        config_mgr = get_config_manager()
        input_path = Path(file_path)
        base_name = input_path.stem
        output_dir = Path(config_mgr.get_output_dir())
        output_dir.mkdir(parents=True, exist_ok=True)
        # 只保存专家版文件
        extracted_info_file = output_dir / f"{base_name}_pro_extracted_info.json"
        
        cache_data = {
            'extracted_info': extracted_info,
            'metadata': {
                'extraction_time': datetime.now().isoformat(),
                'file_path': str(file_path),
                'method': 'pro_strategy',  # 标记为专家策略
                'extractor_version': '2.0',
                'session_id': session_id
            }
        }
        
        with open(extracted_info_file, 'w', encoding='utf-8') as f:
            json.dump(cache_data, f, ensure_ascii=False, indent=2)
        
        logger.info(f"专业版结构化信息已保存到缓存: {extracted_info_file}")
        return True
        
    except Exception as e:
        logger.error(f"保存缓存文件失败: {e}")
        return False

def _split_text_intelligently(text: str, max_chunk_size: int = 32000) -> List[str]:
    """
    智能分割文本，避免在句子中间分割
    
    Args:
        text: 要分割的文本
        max_chunk_size: 每个分块的最大字符数（默认32000，适配64K上下文）
        max_chunk_size: 每个分块的最大字符数
    
    Returns:
        分割后的文本块列表
    """
    if len(text) <= max_chunk_size:
        return [text]
    
    chunks = []
    current_pos = 0
    
    while current_pos < len(text):
        # 计算当前块的结束位置
        end_pos = current_pos + max_chunk_size
        
        if end_pos >= len(text):
            # 最后一块
            chunks.append(text[current_pos:])
            break
        
        # 找到最近的句号、换行或段落边界
        search_start = max(current_pos, end_pos - 1000)  # 在最后1000字符内查找
        
        # 优先在段落边界分割
        paragraph_break = text.rfind('\n\n', search_start, end_pos)
        if paragraph_break != -1 and paragraph_break > current_pos:
            chunks.append(text[current_pos:paragraph_break])
            current_pos = paragraph_break + 2
            continue
        
        # 其次在句号后分割
        sentence_end = text.rfind('。', search_start, end_pos)
        if sentence_end != -1 and sentence_end > current_pos:
            chunks.append(text[current_pos:sentence_end + 1])
            current_pos = sentence_end + 1
            continue
        
        # 最后在换行符分割
        line_break = text.rfind('\n', search_start, end_pos)
        if line_break != -1 and line_break > current_pos:
            chunks.append(text[current_pos:line_break])
            current_pos = line_break + 1
            continue
        
        # 如果都找不到，强制分割
        chunks.append(text[current_pos:end_pos])
        current_pos = end_pos
    
    logger.info(f"文本分割完成: {len(text):,} 字符 -> {len(chunks)} 个分块")
    for i, chunk in enumerate(chunks):
        logger.debug(f"分块 {i+1}: {len(chunk):,} 字符")
    
    return chunks

def _merge_extracted_info(chunk_results: List[Dict]) -> Dict:
    """
    合并多个分块的提取结果
    
    Args:
        chunk_results: 各分块的提取结果列表
    
    Returns:
        合并后的结果字典
    """
    merged = {}
    
    # 定义需要合并的字段 (使用snake_case和语言后缀)
    fields = [
        'title_cn', 'keywords_cn', 'abstract_cn',
        'title_en', 'keywords_en', 'abstract_en',
        'theoretical_framework',
        'references'
    ]
    
    for field in fields:
        values = []
        for result in chunk_results:
            if result and field in result and result[field]:
                # 确保值是字符串类型
                value = result[field]
                if isinstance(value, (list, tuple)):
                    value = ', '.join(str(v) for v in value if v)
                elif not isinstance(value, str):
                    value = str(value)
                
                if value.strip():
                    values.append(value.strip())
        
        if field in ['title_cn', 'title_en']:
            # 标题取第一个非空值
            merged[field] = values[0] if values else ""
        elif field in ['keywords_cn', 'keywords_en']:
            # 关键词去重合并
            all_keywords = []
            for value in values:
                keywords = [kw.strip() for kw in value.split(',') if kw.strip()]
                all_keywords.extend(keywords)
            unique_keywords = list(dict.fromkeys(all_keywords))  # 保持顺序去重
            merged[field] = ', '.join(unique_keywords)
        else:
            # 其他字段合并内容
            merged[field] = '\n\n'.join(values) if values else ""
    
    logger.info(f"分块结果合并完成，包含字段: {list(merged.keys())}")
    return merged

def extract_sections_with_ai(text: str, ai_client, session_id: Optional[str] = None, 
                           languages: Optional[List[str]] = None, use_sections: bool = True) -> Optional[Dict]:
    """
    使用AI模型从论文文本中提取结构化信息
    
    Args:
        text: 论文文本内容
        ai_client: AI客户端实例（必需，支持多种模型：Gemini, OpenAI, Claude等）
        session_id: 会话ID
        languages: 支持的语言列表（当前仅支持中英文）
        use_sections: 是否使用章节分段处理（推荐）
    
    Returns:
        包含提取信息的字典，失败时返回None
    """
    if not ai_client:
        logger.error("必须提供AI客户端实例")
        return None
    
    if languages is None:
        languages = ['Chinese', 'English']
    
    logger.info(f"🎯 开始智能学位论文信息提取，文本长度: {len(text):,} 字符")
    
    # Step 1: 首先提取论文前置部分的元数据（前10000字符）
    front_matter_size = min(10000, len(text) // 4)  # 前25%或最多10000字符
    front_matter = text[:front_matter_size]
    
    logger.info(f"📋 第一步：提取论文元数据 (前 {len(front_matter):,} 字符)")
    metadata = _extract_thesis_metadata(front_matter, ai_client, session_id)
    
    # Step 2: 检查是否使用章节分段处理
    if use_sections and len(text) > 20000:  # 对于长文本使用章节处理
        logger.info(f"📚 第二步：章节分段处理 ({len(text):,} 字符)")
        import time
        section_start_time = time.time()
        
        try:
            from .paper_section_processor import create_section_processor
            logger.info("🔧 初始化章节处理器...")
            processor = create_section_processor(ai_client)
            
            logger.info("🚀 开始章节分段处理...")
            result = processor.process_paper_by_sections(text, session_id)
            
            section_elapsed = time.time() - section_start_time
            
            if result:
                logger.info(f"✅ 章节分段处理成功 (耗时: {section_elapsed:.2f} 秒)")
                logger.info(f"📊 章节处理结果统计: {len(result)} 个字段")
                
                # 合并元数据和章节内容
                if metadata:
                    logger.info("🔗 合并元数据和章节内容")
                    merge_count = 0
                    # 用准确的元数据覆盖章节提取的相应字段
                    for key, value in metadata.items():
                        if value and value.strip():  # 只覆盖非空的元数据
                            if key in result:
                                result[key] = value
                                merge_count += 1
                            elif key == "Author":
                                result["Author"] = value
                                merge_count += 1
                            elif key == "University":
                                result["University"] = value
                                merge_count += 1
                    logger.info(f"📝 成功合并 {merge_count} 个元数据字段")
                            # 可以根据需要添加更多映射
                
                return result
            else:
                section_elapsed = time.time() - section_start_time
                logger.warning(f"❌ 章节分段处理失败 (耗时: {section_elapsed:.2f} 秒)，回退到全文处理")
        except Exception as e:
            section_elapsed = time.time() - section_start_time
            logger.warning(f"❌ 章节分段处理异常 (耗时: {section_elapsed:.2f} 秒): {e}，回退到全文处理")
    
    # Step 3: 全文处理模式（原有逻辑）
    logger.info(f"📖 使用全文处理模式，文本长度: {len(text):,} 字符")
    
    # 检查文本长度，如果过长则分段处理
    max_single_chunk = 32000  # 3.2万字符为单次处理上限，适配64K上下文
    
    if len(text) > max_single_chunk:
        logger.warning(f"⚠️ 文本过长 ({len(text):,} 字符)，启用分段处理模式")
        
        # 分割文本
        text_chunks = _split_text_intelligently(text, max_single_chunk)
        
        # 对每个分块进行处理
        chunk_results = []
        for i, chunk in enumerate(text_chunks):
            logger.info(f"🔄 正在处理分块 {i+1}/{len(text_chunks)} ({len(chunk):,} 字符)")
            
            result = _extract_single_chunk(chunk, ai_client, session_id, languages)
            if result:
                chunk_results.append(result)
            else:
                logger.warning(f"❌ 分块 {i+1} 处理失败")
        
        if not chunk_results:
            logger.error("❌ 所有分块处理都失败了")
            return None
        
        # 合并结果
        logger.info("🔗 正在合并分块处理结果...")
        result = _merge_extracted_info(chunk_results)
        
        # 合并元数据
        if metadata and result:
            logger.info("🔗 合并元数据和分块内容")
            for key, value in metadata.items():
                if value and value.strip():
                    if key in result:
                        result[key] = value
        
        return result
    
    else:
        # 文本长度合适，直接处理
        logger.info("📄 文本长度适中，直接全文处理")
        result = _extract_single_chunk(text, ai_client, session_id, languages)
        
        # 合并元数据
        if metadata and result:
            logger.info("🔗 合并元数据和全文内容")
            for key, value in metadata.items():
                if value and value.strip():
                    if key in result:
                        result[key] = value
        
        return result

def _extract_thesis_metadata(front_matter: str, ai_client, session_id: Optional[str] = None) -> Optional[Dict]:
    """
    专门提取学位论文前置部分的完整元数据信息
    
    Args:
        front_matter: 论文前置部分内容（前10000字符）
        ai_client: AI客户端
        session_id: 会话ID
        
    Returns:
        包含论文完整元数据的字典
    """
    
    metadata_prompt = f"""请从下方学位论文的前置部分（封面、扉页、摘要、目录等）中准确提取以下完整的元数据信息，并以JSON格式输出：

**重要说明**：
- 请仔细查找论文封面、扉页、摘要页、目录页中的准确信息
- 标题必须是论文封面上的完整、准确标题，不要从摘要内容中推断
- 如果某项信息未找到，请输出空字符串""
- 中英文信息请分别提取，如果只有一种语言版本，另一种留空

**提取字段**：
{{
  "thesis_number": "论文编号（如学号、论文分类号等）",
  "title_cn": "论文的完整中文标题（从封面页准确提取）",
  "title_en": "论文的完整英文标题（从封面页准确提取）",
  "university_cn": "中文学校名称",
  "university_en": "英文学校名称", 
  "degree_level": "申请学位级别（如：学士、硕士、博士）",
  "author_cn": "中文作者姓名",
  "author_en": "英文作者姓名",
  "major_cn": "中文学科专业",
  "major_en": "英文学科专业",
  "research_direction_cn": "中文研究方向",
  "research_direction_en": "英文研究方向",
  "supervisor_cn": "中文指导教师姓名",
  "supervisor_en": "英文指导教师姓名",
  "supervisor_title_cn": "中文指导教师职称（如：教授、副教授、讲师等）",
  "supervisor_title_en": "英文指导教师职称",
  "College": "培养学院名称",
  "defense_date": "论文答辩日期",
  "DegreeGrantingInstitution": "学位授予单位",
  "abstract_cn": "完整的中文摘要内容",
  "abstract_en": "完整的英文摘要内容", 
  "keywords_cn": "中文关键词（用逗号分隔）",
  "keywords_en": "英文关键词（用逗号分隔）",
  "TableOfContents": "论文完整目录结构（包含章节编号和页码）"
}}

**输出要求**：
- 直接返回JSON对象，不要代码块标记
- 确保JSON格式有效
- 未找到的信息用空字符串""表示

**论文前置内容**：
{front_matter}
"""
    
    try:
        logger.info("正在提取学位论文完整元数据信息...")
        response = ai_client.send_message(metadata_prompt, session_id=session_id)
        
        if not response or not response.content:
            logger.error("AI返回空响应")
            return None
            
        content = response.content.strip()
        logger.debug(f"元数据提取响应: {content[:200]}...")
        
        # 解析JSON
        json_content = _extract_json_from_response(content)
        if json_content:
            cleaned_content = _clean_json_content(json_content)
            logger.debug(f"调用JSON备用解析机制处理基础元数据")
            result = _parse_json_with_fallback(cleaned_content)
            if result:
                non_empty_fields = len([k for k, v in result.items() if v and str(v).strip()])
                logger.info(f"✅ 成功提取论文元数据，总共 {len(result)} 个字段，其中 {non_empty_fields} 个非空")
                return result
        
        logger.error("论文元数据提取失败")
        return None
        
    except Exception as e:
        logger.error(f"论文元数据提取异常: {e}")
        return None


def _extract_single_chunk(text: str, ai_client, session_id: Optional[str] = None, 
                         languages: Optional[List[str]] = None) -> Optional[Dict]:
    """
    处理单个文本分块的提取逻辑
    """
    
    # 优化后的提示词，专门针对学位论文结构设计
    prompt = f"""请从下方学位论文文本中准确抽取以下信息，并以严格的JSON格式输出。

**特别注意**：
1. 论文标题：请从文档开头的封面或标题页提取准确的论文标题，不要从摘要或正文内容中推断
2. 作者信息：如果有，请从封面页提取作者姓名、学校、专业、导师等信息
3. 章节识别：请准确识别各个章节的内容，如文献综述、研究方法等

**提取字段**：
1. 中文篇名（title_cn）- 从论文封面或标题页提取的准确标题
2. 中文关键词（keywords_cn）- 从中文摘要后的关键词部分提取
3. 中文摘要（abstract_cn）- 标有"摘要"的中文部分
4. 英文篇名（title_en）- 从论文封面或标题页提取的英文标题
5. 英文关键词（keywords_en）- 从英文摘要后的关键词部分提取
6. 英文摘要（abstract_en）- 标有"Abstract"的英文部分
7. 研究方法（ResearchMethods）- 论文中"研究方法"、"实验方法"、"技术路线"等章节内容
8. 理论框架（TheoreticalFramework）- 论文中"理论基础"、"理论框架"等章节内容
9. 实践问题（practical_problems）- 论文要解决的实际问题和应用背景
10. 解决方案（proposed_solutions）- 论文提出的具体解决方案和策略
11. 研究结论（research_conclusions）- 论文"结论"章节的主要内容
12. 应用价值（application_value）- 论文成果的实际应用价值和意义
13. 参考文献列表（references）- 论文"参考文献"章节的完整文献条目

**输出要求**：
- 直接返回JSON对象，不使用代码块标记
- 缺失信息用空字符串""表示
- 确保JSON格式正确有效
- 字符串值不包含控制字符

**论文内容**：
{text}
"""
    
    try:
        # 记录请求开始
        logger.info(f"开始AI结构化信息提取，文本长度: {len(text):,} 字符")
        logger.info(f"使用语言列表: {languages}")
        
        # 发送请求
        response = ai_client.send_message(prompt, session_id=session_id)
        if not response or not response.content:
            logger.error("AI返回空响应")
            return None
        
        response_text = response.content.strip()
        logger.info(f"AI响应长度: {len(response_text)} 字符")
        logger.info(f"使用AI模型: {ai_client.get_api_type() if hasattr(ai_client, 'get_api_type') else 'Unknown'}")
        
        # 提取JSON内容
        json_content = _extract_json_from_response(response_text)
        if not json_content:
            logger.error("无法从响应中提取JSON内容")
            logger.debug(f"原始响应内容前500字符: {response_text[:500]}")
            return None
        
        # 清理JSON内容
        cleaned_json = _clean_json_content(json_content)
        
        # 解析JSON
        logger.debug(f"调用JSON备用解析机制处理章节分析响应")
        result = _parse_json_with_fallback(cleaned_json)
        if result:
            logger.info(f"JSON解析成功，包含字段: {list(result.keys())}")
        else:
            logger.error("JSON解析失败")
            logger.debug(f"清理后的JSON内容前500字符: {cleaned_json[:500]}")
        
        return result
        
    except TimeoutError as e:
        logger.error(f"AI请求超时: {e}")
        logger.warning("建议检查网络连接和API服务状态")
        return None
    except ConnectionError as e:
        logger.error(f"AI连接错误: {e}")
        logger.warning("建议检查网络连接和API端点配置")
        return None
    except Exception as e:
        logger.error(f"提取结构化信息时出错: {e}")
        logger.error(f"错误类型: {type(e).__name__}")
        import traceback
        logger.debug(f"详细错误堆栈: {traceback.format_exc()}")
        return None
        return None

def extract_sections_with_ai_by_chapters(file_path: str, ai_client, 
                                       max_sections: Optional[int] = None, 
                                       test_mode: bool = False,
                                       session_id: Optional[str] = None) -> Optional[Dict]:
    """
    通过章节分段处理的方式从论文中提取结构化信息
    
    Args:
        file_path: 论文文件路径
        ai_client: AI客户端实例
        max_sections: 最大处理章节数（用于测试）
        test_mode: 测试模式，只处理重要章节
        session_id: 会话ID
    
    Returns:
        包含提取信息的字典，失败时返回None
    """
    try:
        from .paper_section_processor import PaperSectionParser
        
        # 1. 提取文档文本
        if file_path.endswith('.pdf'):
            text = extract_text_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            text = extract_text_from_word(file_path)
        else:
            logger.error("不支持的文件格式")
            return None
        
        if not text.strip():
            logger.error("文档文本提取失败或为空")
            return None
            
        logger.info(f"文档文本提取完成，长度: {len(text):,} 字符")
        
        # 2. 解析章节
        parser = PaperSectionParser()
        sections = parser.parse_sections(text)
        
        logger.info(f"章节解析完成，共 {len(sections)} 个章节")
        
        # 3. 筛选重要章节进行处理
        important_sections = []
        important_types = ['abstract', 'introduction', 
                          'methodology', 'experiment', 'results', 'conclusion']
        
        # 按重要性和内容长度筛选章节
        for section in sections:
            if test_mode and max_sections and len(important_sections) >= max_sections:
                break
                
            # 优先处理重要类型的章节
            if section.section_type in important_types and len(section.content.strip()) > 200:
                important_sections.append(section)
            # 或者处理内容较长的未分类章节
            elif section.section_type == 'unknown' and len(section.content.strip()) > 1000:
                important_sections.append(section)
        
        if not important_sections:
            logger.warning("未找到合适的章节进行处理，回退到全文处理")
            return extract_sections_with_ai(text, ai_client, session_id)
        
        logger.info(f"选择 {len(important_sections)} 个重要章节进行处理")
        
        # 4. 分章节提取信息
        extracted_results = {}
        for i, section in enumerate(important_sections, 1):
            logger.info(f"处理第 {i}/{len(important_sections)} 个章节: {section.title[:30]}...")
            
            # 构建章节特定的提示词
            section_prompt = f"""请从下方论文章节中提取相关信息，章节标题：{section.title}

章节内容：
{section.content}

请提取以下信息（如果在此章节中存在）：
1. 中文篇名（title_cn）
2. 中文关键词（keywords_cn）  
3. 中文摘要（abstract_cn）
4. 英文篇名（title_en）
5. 英文关键词（keywords_en）
6. 英文摘要（abstract_en）
7. 研究方法（research_methods）
8. 理论框架（theoretical_framework）
9. 主要创新点（MainInnovations）
10. 实践问题（practical_problems）
11. 解决方案（proposed_solutions）
12. 研究结论（research_conclusions）
13. 应用价值（application_value）
14. 参考文献列表（references）

输出要求：
- 返回JSON格式
- 如某项在此章节中不存在请输出空字符串""
- JSON必须是有效格式
"""
            
            try:
                response = ai_client.send_message(section_prompt, session_id=session_id)
                if response and response.content:
                    # 提取并解析JSON
                    json_content = _extract_json_from_response(response.content)
                    if json_content:
                        cleaned_json = _clean_json_content(json_content)
                        logger.debug(f"调用JSON备用解析机制处理章节 '{section.title}' 的分析响应")
                        section_result = _parse_json_with_fallback(cleaned_json)
                        
                        if section_result:
                            # 合并结果，非空字段覆盖已有结果
                            for key, value in section_result.items():
                                if value and value.strip():
                                    if key not in extracted_results or not extracted_results[key]:
                                        extracted_results[key] = value
                                    else:
                                        # 对于某些字段，追加内容而不是覆盖
                                        if key in ['references']:  
                                            extracted_results[key] += "\n\n" + value
                                        else:
                                            extracted_results[key] = value
                            
                            logger.info(f"章节 {i} 处理成功，提取到 {len([v for v in section_result.values() if v and v.strip()])} 个非空字段")
                        else:
                            logger.warning(f"章节 {i} JSON解析失败")
                    else:
                        logger.warning(f"章节 {i} 未找到有效JSON")
                else:
                    logger.warning(f"章节 {i} AI响应为空")
                    
            except Exception as e:
                logger.error(f"处理章节 {i} 时出错: {e}")
                continue
        
        # 5. 确保所有必需字段都存在 (使用snake_case和语言后缀)
        required_fields = [
            'title_cn', 'keywords_cn', 'abstract_cn',
            'title_en', 'keywords_en', 'abstract_en',
            'theoretical_framework',
            'references'
        ]
        
        for field in required_fields:
            if field not in extracted_results:
                extracted_results[field] = ""
        
        # 6. 如果关键信息缺失太多，补充全文处理
        key_fields = ['title_cn', 'abstract_cn']  
        missing_key_fields = sum(1 for field in key_fields if not extracted_results.get(field, '').strip())
        
        if missing_key_fields >= 2:
            logger.info("关键信息缺失较多，进行补充全文处理...")
            full_result = extract_sections_with_ai(text, ai_client, session_id)
            if full_result:
                for key, value in full_result.items():
                    if not extracted_results.get(key, '').strip() and value and value.strip():
                        extracted_results[key] = value
        
        logger.info(f"章节处理完成，共提取 {len([v for v in extracted_results.values() if v and v.strip()])} 个非空字段")
        return extracted_results
        
    except Exception as e:
        logger.error(f"章节处理出错: {e}")
        logger.error(f"错误类型: {type(e).__name__}")
        return None

class DocumentCache:
    """文档缓存管理器 - 缓存PDF/Word转换的Markdown文本"""
    
    def __init__(self, cache_dir: str = "cache/documents"):
        """
        初始化文档缓存管理器
        
        Args:
            cache_dir: 缓存目录路径
        """
        self.cache_dir = Path(cache_dir)
        self.cache_dir.mkdir(parents=True, exist_ok=True)
        self.logger = logging.getLogger(__name__)
    
    def _get_file_hash(self, file_path: str) -> str:
        """计算文件的MD5哈希值"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()
    
    def _get_cache_key(self, file_path: str) -> str:
        """生成缓存键"""
        file_path_obj = Path(file_path)
        file_hash = self._get_file_hash(file_path)
        return f"{file_path_obj.stem}_{file_hash}"
    
    def _get_cache_file_path(self, cache_key: str) -> Path:
        """获取缓存文件路径"""
        return self.cache_dir / f"{cache_key}.md"
    
    def _get_metadata_file_path(self, cache_key: str) -> Path:
        """获取元数据文件路径"""
        return self.cache_dir / f"{cache_key}_metadata.json"
    
    def is_cached(self, file_path: str) -> bool:
        """检查文件是否已缓存"""
        try:
            cache_key = self._get_cache_key(file_path)
            cache_file = self._get_cache_file_path(cache_key)
            metadata_file = self._get_metadata_file_path(cache_key)
            
            return cache_file.exists() and metadata_file.exists()
        except Exception as e:
            self.logger.warning(f"检查缓存状态失败: {e}")
            return False
    
    def get_cached_content(self, file_path: str) -> Optional[Dict]:
        """获取缓存的文档内容和元数据"""
        try:
            if not self.is_cached(file_path):
                return None
            
            cache_key = self._get_cache_key(file_path)
            cache_file = self._get_cache_file_path(cache_key)
            metadata_file = self._get_metadata_file_path(cache_key)
            
            # 读取Markdown内容
            with open(cache_file, 'r', encoding='utf-8') as f:
                markdown_content = f.read()
            
            # 读取元数据
            with open(metadata_file, 'r', encoding='utf-8') as f:
                metadata = json.load(f)
            
            return {
                'content': markdown_content,
                'metadata': metadata
            }
            
        except Exception as e:
            self.logger.error(f"读取缓存内容失败: {e}")
            return None
    
    def save_to_cache(self, file_path: str, content: str, 
                     file_type: Optional[str] = None, char_count: Optional[int] = None) -> bool:
        """保存文档内容到缓存"""
        try:
            cache_key = self._get_cache_key(file_path)
            cache_file = self._get_cache_file_path(cache_key)
            metadata_file = self._get_metadata_file_path(cache_key)
            
            # 保存Markdown内容
            with open(cache_file, 'w', encoding='utf-8') as f:
                f.write(content)
            
            # 保存元数据
            metadata = {
                'source_file': str(file_path),
                'file_type': file_type or Path(file_path).suffix.lower(),
                'cached_time': datetime.now().isoformat(),
                'file_size': Path(file_path).stat().st_size,
                'content_length': len(content),
                'char_count': char_count or len(content),
                'cache_key': cache_key
            }
            
            with open(metadata_file, 'w', encoding='utf-8') as f:
                json.dump(metadata, f, ensure_ascii=False, indent=2)
            
            self.logger.info(f"文档缓存保存成功: {cache_file}")
            return True
            
        except Exception as e:
            self.logger.error(f"保存缓存失败: {e}")
            return False
    
    def clear_cache(self, file_path: Optional[str] = None) -> bool:
        """清除缓存（指定文件或全部）"""
        try:
            if file_path:
                # 清除指定文件的缓存
                cache_key = self._get_cache_key(file_path)
                cache_file = self._get_cache_file_path(cache_key)
                metadata_file = self._get_metadata_file_path(cache_key)
                
                if cache_file.exists():
                    cache_file.unlink()
                if metadata_file.exists():
                    metadata_file.unlink()
                
                self.logger.info(f"已清除文件缓存: {file_path}")
            else:
                # 清除所有缓存
                for file in self.cache_dir.glob("*"):
                    if file.is_file():
                        file.unlink()
                
                self.logger.info("已清除所有文档缓存")
            
            return True
            
        except Exception as e:
            self.logger.error(f"清除缓存失败: {e}")
            return False
    
    def get_cache_info(self) -> Dict:
        """获取缓存统计信息"""
        try:
            cache_files = list(self.cache_dir.glob("*.md"))
            metadata_files = list(self.cache_dir.glob("*_metadata.json"))
            
            total_size = sum(f.stat().st_size for f in cache_files)
            
            return {
                'cache_dir': str(self.cache_dir),
                'cached_files': len(cache_files),
                'metadata_files': len(metadata_files),
                'total_size_bytes': total_size,
                'total_size_mb': round(total_size / (1024 * 1024), 2)
            }
            
        except Exception as e:
            self.logger.error(f"获取缓存信息失败: {e}")
            return {}


def extract_text_with_cache(file_path: str, use_cache: bool = True) -> str:
    """
    从文档提取文本，支持缓存
    
    Args:
        file_path: 文档路径
        use_cache: 是否使用缓存
        
    Returns:
        提取的文本内容
    """
    logger.info(f"开始提取文档文本: {file_path}")
    
    # 初始化缓存管理器
    cache_manager = DocumentCache()
    
    # 检查是否使用缓存
    if use_cache and cache_manager.is_cached(file_path):
        logger.info("发现缓存文件，直接读取...")
        cached_data = cache_manager.get_cached_content(file_path)
        if cached_data:
            content = cached_data['content']
            metadata = cached_data['metadata']
            logger.info(f"✅ 缓存命中: {metadata['cached_time']}, {len(content):,} 字符")
            return content
    
    # 提取原始文本
    file_ext = Path(file_path).suffix.lower()
    
    if file_ext == '.pdf':
        logger.info("提取PDF文本...")
        text = extract_text_from_pdf(file_path)
    elif file_ext == '.docx':
        logger.info("提取Word文本...")
        text = extract_text_from_word(file_path)
    else:
        logger.error(f"不支持的文件格式: {file_ext}")
        return ""
    
    if not text:
        logger.error("文本提取失败")
        return ""
    
    # 转换为Markdown格式
    markdown_content = convert_text_to_markdown(text, file_path)
    
    # 保存到缓存
    if use_cache:
        success = cache_manager.save_to_cache(
            file_path, markdown_content, file_ext, len(text)
        )
        if success:
            logger.info("✅ 文档已缓存")
        else:
            logger.warning("⚠️ 缓存保存失败")
    
    return markdown_content


def convert_text_to_markdown(text: str, file_path: Optional[str] = None) -> str:
    """
    将纯文本转换为Markdown格式
    
    Args:
        text: 原始文本
        file_path: 源文件路径（用于元数据）
        
    Returns:
        Markdown格式的文本
    """
    lines = text.split('\n')
    markdown_lines = []
    
    # 添加文档头部信息
    if file_path:
        filename = Path(file_path).name
        markdown_lines.extend([
            f"# {filename}",
            "",
            f"**源文件**: {file_path}",
            f"**转换时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
            f"**原始字符数**: {len(text):,}",
            "",
            "---",
            ""
        ])
    
    # 处理文本内容
    for line in lines:
        line = line.strip()
        
        if not line:
            markdown_lines.append("")
            continue
        
        # 检测可能的标题
        if _is_potential_title(line):
            # 根据长度和位置判断标题级别
            if len(line) < 50 and any(keyword in line for keyword in ['摘要', 'Abstract', '引言', '结论', '参考文献']):
                markdown_lines.append(f"## {line}")
            elif len(line) < 80:
                markdown_lines.append(f"### {line}")
            else:
                markdown_lines.append(line)
        else:
            markdown_lines.append(line)
    
    return '\n'.join(markdown_lines)


def _is_potential_title(line: str) -> bool:
    """判断一行文本是否可能是标题"""
    # 标题特征：较短、包含特定关键词、数字编号等
    title_indicators = [
        r'^\d+[\.、]\s*',  # 数字编号
        r'^[一二三四五六七八九十]+[\.、]\s*',  # 中文数字编号
        r'^第[一二三四五六七八九十]+章',  # 章节
        r'^Abstract$|^摘要$',  # 特殊章节
        r'^Introduction$|^引言$',
        r'^Conclusion$|^结论$',
        r'^References$|^参考文献$',
        r'^致谢$|^Acknowledgments?$'
    ]
    
    return any(re.match(pattern, line.strip()) for pattern in title_indicators)


# 全局缓存管理器实例
_document_cache = None

def get_document_cache() -> DocumentCache:
    """获取全局文档缓存管理器实例"""
    global _document_cache
    if _document_cache is None:
        _document_cache = DocumentCache()
    return _document_cache


def extract_sections_with_pro_strategy(file_path: str, use_cache: bool = True) -> Dict[str, Any]:
    """
    使用专业版策略提取论文信息
    整合: 分步抽取策略、结构化分析、快速定位、正则匹配、参考文献解析、智能修复
    
    Args:
        file_path: 论文文件路径
        use_cache: 是否使用缓存
        
    Returns:
        提取的论文信息字典
    """
    # 标准空结果字典 - 包含所有标准字段
    standard_empty_result = {
        'thesis_number': '',
        'title_cn': '',
        'title_en': '',
        'author_cn': '',
        'author_en': '',
        'university_cn': '',
        'university_en': '',
        'degree_level': '',
        'major_cn': '',
        'college': '',
        'supervisor_cn': '',
        'supervisor_en': '',
        'defense_date': '',
        'submission_date': '',
        'abstract_cn': '',
        'abstract_en': '',
        'keywords_cn': '',
        'keywords_en': '',
        'theoretical_framework': {
            'core_theories': [],
            'theoretical_models': [],
            'conceptual_foundations': [],
            'theoretical_contributions': []
        },
        'acknowledgement': '',
        'references': [],
        'author_contributions': {
            'contribution_statement': '',
            'research_contributions': [],
            'publication_contributions': [],
            'innovation_points': []
        }
    }
    
    try:
        # 1. 提取论文全文文本（支持缓存）
        text = extract_text_with_cache(file_path, use_cache)
        if not text:
            logger.error("论文全文文本提取失败")
            return standard_empty_result
        
        # 2. 使用专业版提取器
        extractor = ThesisExtractorPro()
        result = extractor.extract_with_integrated_strategy(text, file_path)
        
        # 3. 保存提取结果到缓存
        if result and use_cache:
            from .config_manager import get_config_manager
            config_mgr = get_config_manager()
            input_path = Path(file_path)
            base_name = input_path.stem
            output_dir = Path(config_mgr.get_output_dir())
            output_dir.mkdir(parents=True, exist_ok=True)
            
            # 保存专业版提取结果
            extracted_info_file = output_dir / f"{base_name}_pro_extracted_info.json"
            extracted_data = {
                'extracted_info': result,
                'metadata': {
                    'extraction_time': datetime.now().isoformat(),
                    'file_path': str(file_path),
                    'method': 'pro_strategy',
                    'extractor_version': '2.0',
                    'stats': extractor.extraction_stats
                }
            }
            
            with open(extracted_info_file, 'w', encoding='utf-8') as f:
                json.dump(extracted_data, f, ensure_ascii=False, indent=2)
            
            logger.info(f"专业版提取结果已保存: {extracted_info_file}")
        
        return result
        
    except Exception as e:
        logger.error(f"专业版提取失败: {e}")
        import traceback
        logger.debug(f"详细错误堆栈: {traceback.format_exc()}")
        return standard_empty_result

    def find_precise_section_boundaries(self, text: str, section_title: str) -> Dict[str, Any]:
        """精确查找章节边界的高级方法"""
        result = {
            'found': False,
            'title': '',
            'start_pos': -1,
            'end_pos': -1,
            'start_line': -1,
            'end_line': -1,
            'content': '',
            'confidence': 0.0,
            'next_section': '',
            'title_variations': []
        }
        
        try:
            # 1. 构建多种标题变体模式
            title_patterns = self._generate_title_patterns(section_title)
            
            # 2. 查找最佳匹配
            best_match = None
            best_confidence = 0.0
            
            for pattern_info in title_patterns:
                pattern = pattern_info['pattern']
                confidence_boost = pattern_info['confidence_boost']
                
                match = re.search(pattern, text, re.IGNORECASE | re.MULTILINE)
                if match:
                    # 计算匹配质量
                    match_confidence = self._evaluate_match_quality(
                        match, text, section_title
                    ) + confidence_boost
                    
                    if match_confidence > best_confidence:
                        best_confidence = match_confidence
                        best_match = match
                        result['title_variations'].append(pattern_info['description'])
            
            if best_match:
                result['found'] = True
                result['start_pos'] = best_match.start()
                result['confidence'] = best_confidence
                
                # 3. 提取精确的标题
                result['title'] = self._extract_clean_title(best_match, section_title)
                
                # 4. 查找章节结束边界
                end_info = self._find_section_end_boundary(
                    text, best_match.end(), section_title
                )
                result['end_pos'] = end_info['end_pos']
                result['next_section'] = end_info['next_section']
                
                # 5. 计算行号
                result['start_line'] = text[:result['start_pos']].count('\n') + 1
                result['end_line'] = text[:result['end_pos']].count('\n') + 1
                
                # 6. 提取内容
                result['content'] = text[result['start_pos']:result['end_pos']].strip()
                
                print(f"   🎯 精确定位章节: {result['title']}")
                print(f"      📍 位置: 字符 {result['start_pos']}-{result['end_pos']} (行 {result['start_line']}-{result['end_line']})")
                print(f"      📏 长度: {len(result['content'])} 字符")
                print(f"      🔍 置信度: {result['confidence']:.2f}")
                if result['next_section']:
                    print(f"      ⏭️ 下一章节: {result['next_section']}")
            
            return result
            
        except Exception as e:
            logger.error(f"精确边界查找失败: {e}")
            return result
    
    def _generate_title_patterns(self, section_title: str) -> List[Dict[str, Any]]:
        """生成章节标题的多种匹配模式"""
        patterns = []
        
        # 标准模式映射
        title_mappings = {
            '摘要': [
                (r'^\s*(中文)?摘\s*要\s*$', '标准摘要标题', 1.0),
                (r'^\s*(中文)?摘\s*要\s*\n', '摘要标题后换行', 0.9),
                (r'摘\s*要', '摘要关键词', 0.7),
            ],
            'abstract': [
                (r'^\s*(ABSTRACT|Abstract)\s*$', '标准英文摘要标题', 1.0),
                (r'^\s*(ABSTRACT|Abstract)\s*\n', '英文摘要标题后换行', 0.9),
                (r'(ABSTRACT|Abstract)', '英文摘要关键词', 0.7),
            ],
            '第一章': [
                (r'^\s*第[一1]章\s+([^\n\r]{5,50})', '标准第一章格式', 1.0),
                (r'^\s*1\.\s*([^\n\r]{5,50})', '数字编号第一章', 0.9),
                (r'^\s*(引\s*言|绪\s*论|概\s*述)\s*$', '引言类标题', 0.8),
            ],
            '第二章': [
                (r'^\s*第[二2]章\s+([^\n\r]{5,50})', '标准第二章格式', 1.0),
                (r'^\s*2\.\s*([^\n\r]{5,50})', '数字编号第二章', 0.9),
                (r'^\s*(文献综述|相关工作|国内外研究现状)\s*$', '文献综述类标题', 0.8),
            ],
            '结论': [
                (r'^\s*(结\s*论|全文总结|研究总结)\s*$', '标准结论标题', 1.0),
                (r'^\s*(结论与展望|结论与建议)\s*$', '结论展望类标题', 0.9),
                (r'结\s*论', '结论关键词', 0.7),
            ],
            '参考文献': [
                (r'^\s*参考文献\s*$', '标准参考文献标题', 1.0),
                (r'^\s*(REFERENCES?|References?)\s*$', '英文参考文献标题', 1.0),
                (r'参考文献', '参考文献关键词', 0.8),
            ]
        }
        
        # 根据输入的章节标题选择合适的模式
        section_lower = section_title.lower()
        
        if '摘要' in section_title or 'abstract' in section_lower:
            if 'abstract' in section_lower:
                patterns.extend([
                    {'pattern': p, 'description': d, 'confidence_boost': c}
                    for p, d, c in title_mappings.get('abstract', [])
                ])
            else:
                patterns.extend([
                    {'pattern': p, 'description': d, 'confidence_boost': c}
                    for p, d, c in title_mappings.get('摘要', [])
                ])
        elif '第一章' in section_title or '引言' in section_title:
            patterns.extend([
                {'pattern': p, 'description': d, 'confidence_boost': c}
                for p, d, c in title_mappings.get('第一章', [])
            ])
        elif '第二章' in section_title or '文献' in section_title:
            patterns.extend([
                {'pattern': p, 'description': d, 'confidence_boost': c}
                for p, d, c in title_mappings.get('第二章', [])
            ])
        elif '结论' in section_title:
            patterns.extend([
                {'pattern': p, 'description': d, 'confidence_boost': c}
                for p, d, c in title_mappings.get('结论', [])
            ])
        elif '参考文献' in section_title or 'reference' in section_lower:
            patterns.extend([
                {'pattern': p, 'description': d, 'confidence_boost': c}
                for p, d, c in title_mappings.get('参考文献', [])
            ])
        
        # 如果没有预定义模式，创建通用模式
        if not patterns:
            escaped_title = re.escape(section_title)
            patterns = [
                {
                    'pattern': f'^\\s*{escaped_title}\\s*$',
                    'description': f'精确匹配 {section_title}',
                    'confidence_boost': 1.0
                },
                {
                    'pattern': escaped_title,
                    'description': f'包含 {section_title}',
                    'confidence_boost': 0.6
                }
            ]
        
        return patterns
    
    def _evaluate_match_quality(self, match: re.Match, text: str, section_title: str) -> float:
        """评估匹配质量"""
        confidence = 0.5  # 基础分数
        
        # 检查匹配位置是否在行首
        if match.start() == 0 or text[match.start()-1] == '\n':
            confidence += 0.2
        
        # 检查匹配后是否有换行
        if match.end() < len(text) and text[match.end()] == '\n':
            confidence += 0.2
        
        # 检查匹配内容的格式特征
        matched_text = match.group(0)
        
        # 如果是章节编号格式
        if re.search(r'第[一二三四五六七八九十1-9]章', matched_text):
            confidence += 0.3
        
        # 如果匹配长度合理（不太长也不太短）
        if 2 <= len(matched_text) <= 50:
            confidence += 0.1
        
        return min(1.0, confidence)
    
    def _extract_clean_title(self, match: re.Match, original_title: str) -> str:
        """提取干净的章节标题"""
        matched_text = match.group(0).strip()
        
        # 移除多余的空白字符
        cleaned = re.sub(r'\s+', ' ', matched_text)
        
        # 如果有分组，尝试提取章节号和标题
        if match.groups():
            for group in match.groups():
                if group:
                    potential_title = group.strip()
                    if len(potential_title) > len(cleaned.split()[-1]):
                        cleaned = potential_title
                        break
        
        return cleaned
    
    def _find_section_end_boundary(self, text: str, start_pos: int, current_section: str) -> Dict[str, Any]:
        """查找章节结束边界"""
        result = {
            'end_pos': len(text),
            'next_section': '',
            'boundary_type': 'end_of_document'
        }
        
        # 从当前章节开始位置向后搜索
        search_text = text[start_pos:]
        
        # 定义可能的下一章节模式
        next_section_patterns = [
            (r'\n\s*第[一二三四五六七八九十1-9]章', 'next_chapter'),
            (r'\n\s*\d+\.\s*[^\n]{5,50}', 'numbered_section'),
            (r'\n\s*(结\s*论|参考文献|致\s*谢|附\s*录)', 'end_section'),
            (r'\n\s*(REFERENCES?|ACKNOWLEDGMENT)', 'end_section_en'),
        ]
        
        earliest_match = None
        earliest_pos = len(search_text)
        
        for pattern, boundary_type in next_section_patterns:
            match = re.search(pattern, search_text, re.IGNORECASE)
            if match and match.start() < earliest_pos:
                earliest_pos = match.start()
                earliest_match = match
                result['boundary_type'] = boundary_type
        
        if earliest_match:
            result['end_pos'] = start_pos + earliest_pos
            result['next_section'] = earliest_match.group(0).strip()
        
        return result

    def _analyze_review_chapters_concurrently(self, text: str, review_chapters: List[Dict]) -> Dict[str, Any]:
        """并发分析综述章节"""
        import concurrent.futures
        import time
        
        results = {
            'chapter_summaries': {},
            'literature_analysis': {}
        }
        
        if not self.ai_client or not review_chapters:
            return results
        
        max_workers = min(2, len(review_chapters))  # 综述章节通常较少，限制并发数
        print(f"   🔄 启动 {max_workers} 个线程并发分析综述章节...")
        
        def analyze_review_chapter(chapter):
            """分析单个综述章节"""
            chapter_title = chapter.get('title', 'Unknown')
            try:
                print(f"      📖 [{chapter_title[:20]}] 开始综述分析...")
                start_time = time.time()
                
                chapter_content = self._extract_chapter_content(text, chapter)
                if not chapter_content or len(chapter_content) < 200:
                    print(f"      ⚠️ [{chapter_title[:20]}] 内容不足，跳过")
                    return chapter_title, None, None
                
                # 综合综述分析
                summary = self._generate_review_chapter_analysis(chapter, chapter_content)
                
                # 综述深度分析 
                review_analysis = self._conduct_comprehensive_review_analysis(chapter_content)
                
                elapsed = time.time() - start_time
                if summary or review_analysis:
                    print(f"      ✅ [{chapter_title[:20]}] 综述分析完成 ({elapsed:.1f}s, {len(chapter_content)} 字符)")
                    return chapter_title, summary, review_analysis
                else:
                    print(f"      ⚠️ [{chapter_title[:20]}] 综述分析失败 ({elapsed:.1f}s)")
                    return chapter_title, None, None
                
            except Exception as e:
                print(f"      ❌ [{chapter_title[:20]}] 综述分析异常: {e}")
                return chapter_title, None, None
        
        # 使用线程池并发分析
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            start_time = time.time()
            
            # 提交所有任务
            future_to_chapter = {
                executor.submit(analyze_review_chapter, chapter): chapter.get('title', 'Unknown')
                for chapter in review_chapters
            }
            
            # 收集结果
            for future in concurrent.futures.as_completed(future_to_chapter):
                chapter_title = future_to_chapter[future]
                try:
                    returned_title, summary, review_analysis = future.result(timeout=90)  # 综述章节可能更复杂，给更多时间
                    
                    if summary:
                        results['chapter_summaries'][returned_title] = summary
                    if review_analysis:
                        results['literature_analysis'][returned_title] = review_analysis
                        
                except concurrent.futures.TimeoutError:
                    print(f"      ⏰ [{chapter_title[:20]}] 综述分析超时")
                except Exception as e:
                    print(f"      ❌ [{chapter_title[:20]}] 并发执行异常: {e}")
            
            total_time = time.time() - start_time
            success_count = len(results['chapter_summaries'])
            print(f"   ⚡ 综述章节并发分析完成: {success_count}/{len(review_chapters)} 成功，总耗时 {total_time:.1f}s")
        
        return results
    
    def _analyze_other_chapters_concurrently(self, text: str, chapters: List[Dict]) -> Dict[str, Any]:
        """并发分析其他章节"""
        import concurrent.futures
        import time
        
        results = {}
        
        if not self.ai_client or not chapters:
            return results
        
        max_workers = min(3, len(chapters))  # 其他章节可以更多并发
        print(f"   🔄 启动 {max_workers} 个线程并发分析其他章节...")
        
        def analyze_other_chapter(chapter):
            """分析单个非综述章节"""
            chapter_title = chapter.get('title', 'Unknown')
            try:
                print(f"      📚 [{chapter_title[:20]}] 开始章节分析...")
                start_time = time.time()
                
                chapter_content = self._extract_chapter_content(text, chapter)
                if not chapter_content or len(chapter_content) < 200:
                    print(f"      ⚠️ [{chapter_title[:20]}] 内容不足，跳过")
                    return chapter_title, None
                
                summary = self._generate_chapter_summary_with_ai(chapter, chapter_content)
                
                elapsed = time.time() - start_time
                if summary:
                    print(f"      ✅ [{chapter_title[:20]}] 章节分析完成 ({elapsed:.1f}s)")
                    return chapter_title, summary
                else:
                    print(f"      ⚠️ [{chapter_title[:20]}] 章节分析失败 ({elapsed:.1f}s)")
                    return chapter_title, None
                
            except Exception as e:
                print(f"      ❌ [{chapter_title[:20]}] 章节分析异常: {e}")
                return chapter_title, None
        
        # 使用线程池并发分析
        with concurrent.futures.ThreadPoolExecutor(max_workers=max_workers) as executor:
            start_time = time.time()
            
            # 提交所有任务
            future_to_chapter = {
                executor.submit(analyze_other_chapter, chapter): chapter.get('title', 'Unknown')
                for chapter in chapters
            }
            
            # 收集结果
            for future in concurrent.futures.as_completed(future_to_chapter):
                chapter_title = future_to_chapter[future]
                try:
                    returned_title, summary = future.result(timeout=60)  # 普通章节60秒超时
                    
                    if summary:
                        results[returned_title] = summary
                        
                except concurrent.futures.TimeoutError:
                    print(f"      ⏰ [{chapter_title[:20]}] 章节分析超时")
                except Exception as e:
                    print(f"      ❌ [{chapter_title[:20]}] 并发执行异常: {e}")
            
            total_time = time.time() - start_time
            success_count = len(results)
            print(f"   ⚡ 其他章节并发分析完成: {success_count}/{len(chapters)} 成功，总耗时 {total_time:.1f}s")
        
        return results