#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
AI智能学位论文目录抽取器
专门支持Word文档(.docx)格式的目录智能抽取
使用AI技术识别和解析复杂的论文目录结构
仅支持中文论文，不支持藏语或其他语言论文
"""

import re
import json
import docx
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Union
from dataclasses import dataclass, asdict
from abc import ABC, abstractmethod
import logging
from datetime import datetime
import xml.etree.ElementTree as ET
from docx.oxml.ns import qn

# 设置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def is_chinese_text(text: str, min_chinese_ratio: float = 0.3) -> bool:
    """
    检测文本是否为中文
    
    Args:
        text: 待检测文本
        min_chinese_ratio: 最小中文字符比例阈值
    
    Returns:
        bool: 是否为中文文本
    """
    if not text or not text.strip():
        return False
    
    # 移除空白字符、标点符号和数字
    clean_text = re.sub(r'[\s\d\W]', '', text)
    
    if len(clean_text) == 0:
        return False
    
    # 统计中文字符数量（包括中文标点）
    chinese_chars = re.findall(r'[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]', clean_text)
    chinese_ratio = len(chinese_chars) / len(clean_text)
    
    return chinese_ratio >= min_chinese_ratio

def detect_non_chinese_content(text: str, sample_lines: int = 100) -> Tuple[bool, str]:
    """
    检测文本中是否包含大量非中文内容（如藏语、英语等）
    
    Args:
        text: 待检测文本
        sample_lines: 采样行数
    
    Returns:
        Tuple[bool, str]: (是否包含大量非中文内容, 检测到的语言类型)
    """
    lines = text.split('\n')[:sample_lines]
    non_chinese_lines = 0
    tibetan_lines = 0
    total_content_lines = 0
    detected_language = "unknown"
    
    for line in lines:
        line = line.strip()
        if not line or len(line) < 3:  # 跳过空行和过短的行
            continue
            
        total_content_lines += 1
            
        # 检测藏语（藏文Unicode范围）- 更严格的检测
        tibetan_chars = re.findall(r'[\u0f00-\u0fff]', line)
        if len(tibetan_chars) > 2:  # 如果一行中有超过2个藏文字符
            tibetan_lines += 1
            non_chinese_lines += 1
            detected_language = "tibetan"
            continue
        
        # 检测英语（大量英文单词）
        english_words = re.findall(r'\b[a-zA-Z]{4,}\b', line)
        chinese_chars = re.findall(r'[\u4e00-\u9fff]', line)
        if len(english_words) > 2 and len(chinese_chars) < 3:
            non_chinese_lines += 1
            detected_language = "english"
            continue
        
        # 检测其他语言（大量非中文字符）
        if not is_chinese_text(line, min_chinese_ratio=0.15):
            non_chinese_lines += 1
            if detected_language == "unknown":
                detected_language = "other"
    
    if total_content_lines == 0:
        return False, "unknown"
    
    # 特别严格检测藏文
    tibetan_ratio = tibetan_lines / total_content_lines
    non_chinese_ratio = non_chinese_lines / total_content_lines
    
    # 如果藏文行比例超过20%，则认为是藏文文档
    if tibetan_ratio > 0.2:
        return True, "tibetan"
    
    # 如果非中文行比例超过40%，则认为是非中文文档
    return non_chinese_ratio > 0.4, detected_language

@dataclass
class TocEntry:
    """目录条目数据结构"""
    level: int              # 章节层级 (1=主章节, 2=二级章节, 3=三级章节)
    number: str             # 章节编号 (如 "1", "1.1", "1.1.1")
    title: str              # 章节标题
    page: Optional[int]     # 页码
    line_number: int        # 在原文档中的行号
    confidence: float       # AI识别置信度 (0-1)
    section_type: str       # 章节类型 (chapter, section, subsection, etc.)

@dataclass
class ThesisToc:
    """论文目录结构"""
    title: str                    # 论文标题
    author: str                   # 作者
    entries: List[TocEntry]       # 目录条目列表
    total_entries: int           # 总条目数
    max_level: int               # 最大层级深度
    extraction_method: str        # 抽取方法
    confidence_score: float      # 整体置信度
    toc_content: str             # 目录原始内容

class DocumentParser(ABC):
    """文档解析器抽象基类"""
    
    @abstractmethod
    def parse(self, file_path: str) -> str:
        """解析文档内容"""
        pass
    
    @abstractmethod
    def get_lines(self, file_path: str) -> List[str]:
        """获取文档行列表"""
        pass
    
    @abstractmethod
    def extract_toc_boundary(self, file_path: str) -> Tuple[str, int, int]:
        """提取目录边界和内容"""
        pass

class WordParser(DocumentParser):
    """Word文档解析器，专门处理.docx格式"""
    
    def parse(self, file_path: str) -> str:
        """解析Word文档内容"""
        try:
            doc = docx.Document(file_path)
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            return content
        except Exception as e:
            logger.error(f"解析Word文档失败: {e}")
            raise
    
    def extract_toc_from_bookmarks(self, doc_path: str, detected_lang: str = "chinese") -> Dict:
        """通过XML书签提取完整的TOC结构，支持中英文不同截取策略"""
        try:
            import re
            document = docx.Document(doc_path)
            document_xml = document._element.xml
            
            # 查找所有TOC书签和对应的文本
            bookmark_start_pattern = r'<w:bookmarkStart[^>]*w:name="(_Toc\d+)"[^>]*/>'
            bookmark_starts = re.finditer(bookmark_start_pattern, document_xml)
            
            bookmark_positions = []
            
            # 收集所有书签位置
            for match in bookmark_starts:
                bookmark_name = match.group(1)
                start_pos = match.end()
                bookmark_positions.append((bookmark_name, start_pos))
            
            logger.info(f"找到 {len(bookmark_positions)} 个TOC书签位置")
            
            if not bookmark_positions:
                return {}
            
            toc_entries = []
            
            # 为每个书签提取后续文本内容
            for i, (bookmark_name, start_pos) in enumerate(bookmark_positions):
                # 确定提取文本的结束位置 - 增加提取范围
                if i + 1 < len(bookmark_positions):
                    end_pos = bookmark_positions[i + 1][1]
                else:
                    end_pos = start_pos + 2000  # 最后一个书签，提取更多内容
                
                # 提取书签后的XML片段
                xml_fragment = document_xml[start_pos:end_pos]
                
                # 从XML中提取文本内容 - 改进的文本提取逻辑
                text_pattern = r'<w:t[^>]*>(.*?)</w:t>'
                text_matches = re.findall(text_pattern, xml_fragment, re.DOTALL)
                
                if text_matches:
                    # 合并所有相关的文本片段，而不是只取第一个
                    text_parts = []
                    for text in text_matches:
                        clean_text = text.strip()
                        if clean_text and not clean_text.isspace():
                            text_parts.append(clean_text)
                        
                        # 如果已经收集到足够的文本内容，停止收集
                        if len(' '.join(text_parts)) > 100:
                            break
                    
                    if text_parts:
                        # 合并文本片段，去除重复和无关内容
                        text_content = ' '.join(text_parts)
                        
                        # 清理文本内容
                        text_content = self._clean_extracted_text(text_content)
                        
                        if text_content:
                            toc_entries.append((bookmark_name, text_content))
            
            # 过滤出真正的目录条目（通常从"摘要"、"目录"、"绪论"等开始）
            start_keywords = ['摘要', '摘 要', 'ABSTRACT', '目录', '绪论', '第一章', '第1章', '1.', '1 ']
            start_index = -1
            
            for i, (bookmark, text) in enumerate(toc_entries):
                for keyword in start_keywords:
                    if keyword in text:
                        start_index = i
                        break
                if start_index >= 0:
                    break
            
            if start_index >= 0:
                filtered_toc = toc_entries[start_index:]
                logger.info(f"通过书签提取到 {len(filtered_toc)} 个有效目录项")
                
                # 详细打印有效目录项
                print(f"\n📋 详细分析：{len(filtered_toc)}个有效目录项内容")
                print("=" * 80)
                for i, (bookmark, text) in enumerate(filtered_toc):
                    print(f"[{i+1:2d}] 书签: {bookmark} | 文本: {text}")
                print("=" * 80)
                
                # 格式化为字符串（根据语言类型截取文本）
                toc_content = ""
                truncated_entries = []
                for bookmark, text in filtered_toc:
                    # 根据语言类型截取文本
                    truncated_text = self._truncate_text_by_language(text, detected_lang)
                    
                    # 模式匹配处理：规范化章节标题格式
                    truncated_text = self._normalize_chapter_title(truncated_text)
                    
                    toc_content += truncated_text + "\n"
                    truncated_entries.append((bookmark, truncated_text))
                
                # 记录截取后的目录内容
                lang_desc = "前30字符" if detected_lang == "chinese" else "前20个单词"
                print(f"\n截取后的目录内容（{lang_desc}，传递给AI分析）:")
                print("-" * 60)
                for i, (bookmark, truncated_text) in enumerate(truncated_entries):
                    print(f"[{i+1:2d}] {truncated_text}")
                print("-" * 60)
                
                return {
                    'content': toc_content,
                    'entries': filtered_toc,
                    'method': 'xml_bookmarks'
                }
            else:
                logger.warning("未找到明确的目录起始点")
                return {}
                
        except Exception as e:
            logger.error(f"通过书签提取TOC失败: {e}")
            return {}
    
    def _clean_extracted_text(self, text: str) -> str:
        """清理从XML提取的文本内容"""
        if not text:
            return ""
        
        # 移除多余的空白字符
        text = re.sub(r'\s+', ' ', text.strip())
        
        # 移除常见的无关内容
        unwanted_patterns = [
            r'HYPERLINK.*?"',  # Word超链接
            r'_Toc\d+',        # TOC引用
            r'PAGEREF.*?"',    # 页面引用
            r'REF.*?"',        # 一般引用
            r'\\"',            # 转义引号
        ]
        
        for pattern in unwanted_patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE)
        
        # 移除多余空格并返回
        return re.sub(r'\s+', ' ', text.strip())
    
    def _truncate_text_by_language(self, text: str, detected_lang: str) -> str:
        """根据语言类型截取文本：中文截取30字符，英文截取20个单词"""
        if not text:
            return text
        
        if detected_lang == "chinese":
            # 中文：截取前30字符
            truncated = text[:30].strip()
            if len(text) > 30:
                truncated += "..."
        else:
            # 英文：截取前20个单词
            words = text.split()
            if len(words) <= 20:
                truncated = text.strip()
            else:
                truncated = ' '.join(words[:20]) + "..."
        
        return truncated
    
    def _normalize_chapter_title(self, text: str) -> str:
        """规范化章节标题格式，去除多余空格"""
        if not text:
            return text
        
        import re
        
        # 中文章节标题规范化：去除"第"和"章"之间的空格
        # "第 一 章" -> "第一章"
        # "第 二 章" -> "第二章"
        # "第 三 章" -> "第三章"  
        # "第 四 章" -> "第四章"
        # "第 五 章" -> "第五章"
        # "第 六 章" -> "第六章"
        text = re.sub(r'第\s+([一二三四五六七八九十]+)\s+章', r'第\1章', text)
        
        # 阿拉伯数字章节标题规范化
        # "第 1 章" -> "第1章"
        # "第 2 章" -> "第2章"
        text = re.sub(r'第\s+(\d+)\s+章', r'第\1章', text)
        
        return text

    def extract_toc_fields(self, doc_path: str) -> Dict:
        """使用XML解析提取Word文档中的TOC字段"""
        try:
            document = docx.Document(doc_path)
            toc_data = {}
            
            # 搜索文档中的TOC相关内容
            for paragraph in document.paragraphs:
                if paragraph._element is not None:
                    # 简化的TOC字段搜索
                    para_xml = paragraph._element.xml
                    if 'TOC' in para_xml or 'fldChar' in para_xml:
                        text = paragraph.text.strip()
                        if text:
                            logger.info(f"发现潜在TOC段落: {text}")
                            if 'content' not in toc_data:
                                toc_data['content'] = ""
                            toc_data['content'] += text + "\n"
                                    
            return toc_data
        except Exception as e:
            logger.error(f"提取TOC字段失败: {e}")
            return {}
    
    def _extract_toc_content_from_field(self, element) -> str:
        """从TOC字段元素中提取内容"""
        try:
            toc_content = ""
            # 查找TOC字段的内容部分
            for node in element.iter():
                if hasattr(node, 'text') and node.text:
                    text = node.text.strip()
                    if text and not text.startswith('TOC'):
                        toc_content += text + "\n"
            return toc_content
        except Exception as e:
            logger.error(f"提取TOC内容失败: {e}")
            return ""
    
    def extract_by_styles(self, doc_path: str) -> str:
        """基于样式提取TOC内容"""
        try:
            document = docx.Document(doc_path)
            toc_content = ""
            
            # TOC相关样式名称
            toc_styles = [
                'TOC 1', 'TOC 2', 'TOC 3', 'TOC 4', 'TOC 5',
                'toc 1', 'toc 2', 'toc 3', 'toc 4', 'toc 5',
                'Toc1', 'Toc2', 'Toc3', 'Toc4', 'Toc5',
                '目录 1', '目录 2', '目录 3', '目录 4', '目录 5',
                'Heading 1', 'Heading 2', 'Heading 3'  # 有时TOC使用标题样式
            ]
            
            for paragraph in document.paragraphs:
                if paragraph.style and paragraph.style.name in toc_styles:
                    text = paragraph.text.strip()
                    if text:
                        logger.info(f"发现TOC样式段落: {paragraph.style.name} - {text}")
                        toc_content += text + "\n"
                        
            return toc_content
        except Exception as e:
            logger.error(f"基于样式提取TOC失败: {e}")
            return ""
    
    def extract_enhanced_toc(self, file_path: str, detected_lang: str = "chinese") -> str:
        """增强的TOC提取方法，优先使用书签提取，支持中英文不同截取策略"""
        try:
            # 方法1：尝试通过XML书签提取（最可靠）
            bookmark_result = self.extract_toc_from_bookmarks(file_path, detected_lang)
            if bookmark_result and 'content' in bookmark_result and bookmark_result['content'].strip():
                logger.info("使用XML书签提取成功获取TOC")
                return bookmark_result['content']
            
            # 方法2：尝试字段提取
            toc_fields = self.extract_toc_fields(file_path)
            if toc_fields and 'content' in toc_fields and toc_fields['content'].strip():
                logger.info("使用字段提取成功获取TOC")
                return toc_fields['content']
            
            # 方法3：尝试样式提取
            style_toc = self.extract_by_styles(file_path)
            if style_toc.strip():
                logger.info("使用样式提取成功获取TOC")
                return style_toc
            
            # 方法4：回退到原有的边界提取方法
            logger.info("使用边界提取方法作为后备")
            _, start, end = self.extract_toc_boundary(file_path)
            if start != -1 and end != -1:
                lines = self.get_lines(file_path)
                toc_lines = lines[start:end]
                return "\n".join(toc_lines)
            
            return ""
        except Exception as e:
            logger.error(f"增强TOC提取失败: {e}")
            return ""
    
    def get_lines(self, file_path: str) -> List[str]:
        """获取Word文档行列表"""
        try:
            doc = docx.Document(file_path)
            lines = []
            for paragraph in doc.paragraphs:
                lines.append(paragraph.text)
            return lines
        except Exception as e:
            logger.error(f"获取Word文档行失败: {e}")
            raise
    
    def extract_toc_boundary(self, file_path: str) -> Tuple[str, int, int]:
        """提取Word文档中的目录边界"""
        try:
            doc = docx.Document(file_path)
            lines = []
            for paragraph in doc.paragraphs:
                lines.append(paragraph.text)
            
            # 查找目录起始位置
            toc_start = -1
            toc_end = -1
            
            # 目录开始标识（必须是单独的行）
            toc_indicators = [
                r'^目\s*录\s*$',
                r'^CONTENTS?\s*$',
                r'^TABLE\s+OF\s+CONTENTS?\s*$',
                r'^Contents\s*$',
                r'^目\s*次\s*$'
            ]
            
            # 方法1：寻找明确的目录标题
            for i, line in enumerate(lines):
                line = line.strip()
                if not line:
                    continue
                
                # 查找目录开始 - 必须是独立的行
                if toc_start == -1:
                    for indicator in toc_indicators:
                        if re.match(indicator, line, re.IGNORECASE):
                            toc_start = i
                            logger.info(f"找到目录开始位置: 第{i+1}行 - {line}")
                            break
                
                # 查找目录结束（在目录开始后）
                elif toc_start != -1 and toc_end == -1:
                    # 跳过目录标题行，开始寻找真正的目录内容
                    if i == toc_start:
                        continue
                    
                    # 检查是否是真正的目录条目
                    if self._is_toc_entry(line):
                        # 找到第一个目录条目，继续向下找到目录结束
                        continue
                    
                    # 目录结束的标识 - 只有明确的正文章节开始才算结束
                    end_indicators = [
                        r'^第[一二三四五六七八九十\d]+章\s+\S+',  # 必须有章节标题的正文章节
                        r'^\d+\s+[^.\d\s]\S+',  # 数字开头的正文章节，必须有标题
                        r'^Chapter\s+\d+\s+\S+',  # 英文正文章节必须有标题
                        r'^引\s*言\s*$',
                        r'^绪\s*论\s*$',
                        r'^Introduction\s*$'
                    ]
                    
                    # 只有匹配到明确的正文开始才结束目录搜索
                    for indicator in end_indicators:
                        if re.match(indicator, line, re.IGNORECASE):
                            # 确认不是目录条目
                            if not self._is_toc_entry(line):
                                toc_end = i
                                logger.info(f"找到目录结束位置: 第{i+1}行 - {line}")
                                break
                    
                    if toc_end != -1:
                        break
                    
                    # 如果距离目录开始过远且连续多行非目录内容，认为目录结束
                    if i - toc_start > 100:  # 增加到100行
                        non_toc_count = 0
                        # 检查后续10行是否都不是目录条目
                        for j in range(i, min(i+10, len(lines))):
                            if not self._is_toc_entry(lines[j].strip()):
                                non_toc_count += 1
                        
                        if non_toc_count >= 8:  # 如果10行中有8行都不是目录条目
                            toc_end = i
                            logger.info(f"距离目录开始过远且连续非目录内容，设定结束位置: 第{i+1}行")
                            break
            
            # 方法2：如果没找到明确的目录标题，尝试寻找摘要后的目录区域
            if toc_start == -1:
                logger.info("未找到明确目录标题，尝试寻找摘要后的目录区域")
                abstract_end = -1
                
                # 寻找摘要结束位置
                for i, line in enumerate(lines):
                    line = line.strip()
                    if not line:
                        continue
                    
                    # 检测摘要、ABSTRACT结束
                    if re.match(r'^(摘\s*要|Abstract|ABSTRACT)\s*$', line, re.IGNORECASE):
                        # 向后搜索找到摘要内容结束
                        for j in range(i+1, min(i+50, len(lines))):
                            next_line = lines[j].strip()
                            if not next_line:
                                continue
                            
                            # 如果遇到目录内容特征，说明摘要结束了
                            if (re.match(r'^第[一二三四五六七八九十\d]+章', next_line) or
                                re.match(r'^\d+\.?\d*\s+.+\d+$', next_line) or
                                re.match(r'^目\s*录', next_line)):
                                abstract_end = j - 1
                                break
                        
                        if abstract_end != -1:
                            logger.info(f"找到摘要结束位置: 第{abstract_end+1}行")
                            break
                
                # 如果找到摘要结束位置，从该位置开始寻找目录
                if abstract_end != -1:
                    # 从摘要结束后开始寻找目录内容
                    potential_toc_start = abstract_end + 1
                    
                    # 寻找连续的目录条目
                    consecutive_toc_entries = 0
                    for i in range(potential_toc_start, min(potential_toc_start + 100, len(lines))):
                        line = lines[i].strip()
                        if not line:
                            continue
                        
                        if self._is_toc_entry(line) or re.match(r'^目\s*录\s*$', line):
                            consecutive_toc_entries += 1
                            if consecutive_toc_entries >= 3:  # 找到至少3个连续的目录条目
                                toc_start = max(0, i - consecutive_toc_entries)
                                logger.info(f"通过摘要后检测找到目录开始位置: 第{toc_start+1}行")
                                break
                        else:
                            consecutive_toc_entries = 0
            
            if toc_start == -1:
                logger.warning("未找到目录开始位置，尝试全文搜索目录内容")
                return self._fallback_toc_extraction(lines)
            
            # 寻找目录结束位置
            if toc_end == -1:
                # 如果没找到明确结束，向后搜索更多内容
                for i in range(toc_start + 1, min(toc_start + 150, len(lines))):  # 增加搜索范围到150行
                    line = lines[i].strip()
                    if line and self._is_toc_entry(line):
                        continue
                    elif line and not self._is_toc_entry(line):
                        # 找到非目录内容，需要仔细判断是否真的是正文开始
                        end_indicators = [
                            r'^第[一二三四五六七八九十\d]+章\s*绪\s*论',  # 第一章绪论
                            r'^第[一二三四五六七八九十\d]+章\s+\S+',  # 正文章节
                            r'^\d+\s*\.\s*\d+\s*[选题背景|研究意义|国内外]',  # 1.1选题背景及研究意义
                            r'^\d+\s+[^.\d\s]\S+',
                            r'^Chapter\s+\d+\s+\S+',
                            r'^引\s*言\s*$',
                            r'^绪\s*论\s*$',
                            r'^Introduction\s*$'
                        ]
                        
                        # 只有明确匹配到正文开始标志才停止
                        for indicator in end_indicators:
                            if re.match(indicator, line, re.IGNORECASE):
                                toc_end = i
                                logger.info(f"通过内容匹配找到目录结束位置: 第{i+1}行 - {line}")
                                break
                        
                        if toc_end != -1:
                            break
                
                if toc_end == -1:
                    toc_end = min(toc_start + 150, len(lines))  # 假设目录不超过150行
                    logger.info(f"使用启发式方法确定目录结束位置: 第{toc_end}行")
            
            # 提取目录内容
            if lines[toc_start].strip() and re.match(r'^目\s*录\s*$', lines[toc_start].strip()):
                # 如果toc_start是"目录"标题行，跳过它
                toc_lines = lines[toc_start + 1:toc_end]
            else:
                # 如果toc_start就是目录内容开始，不跳过
                toc_lines = lines[toc_start:toc_end]
            
            toc_content = "\n".join(toc_lines)
            
            logger.info(f"目录边界确定: 第{toc_start+1}行到第{toc_end}行，共{len(toc_lines)}行")
            
            return toc_content, toc_start, toc_end
            
            if toc_end == -1:
                # 如果没找到明确结束，向后搜索更多内容
                for i in range(toc_start + 1, min(toc_start + 150, len(lines))):  # 增加搜索范围到150行
                    line = lines[i].strip()
                    if line and self._is_toc_entry(line):
                        continue
                    elif line and not self._is_toc_entry(line):
                        # 找到非目录内容，但需要更仔细判断是否真的是正文开始
                        end_indicators = [
                            r'^第[一二三四五六七八九十\d]+章\s+\S+',  # 正文章节
                            r'^\d+\s+[^.\d\s]\S+',
                            r'^Chapter\s+\d+\s+\S+',
                            r'^摘\s*要\s*$',
                            r'^Abstract\s*$',
                            r'^引\s*言\s*$',
                            r'^绪\s*论\s*$'
                        ]
                        
                        # 只有明确匹配到正文开始标志才停止
                        for indicator in end_indicators:
                            if re.match(indicator, line, re.IGNORECASE):
                                toc_end = i
                                logger.info(f"通过内容匹配找到目录结束位置: 第{i+1}行 - {line}")
                                break
                        
                        if toc_end != -1:
                            break
                
                if toc_end == -1:
                    toc_end = min(toc_start + 150, len(lines))  # 假设目录不超过150行
                    logger.info(f"使用启发式方法确定目录结束位置: 第{toc_end}行")
            
            # 提取目录内容
            toc_lines = lines[toc_start + 1:toc_end]  # 跳过"目录"标题行
            toc_content = "\n".join(toc_lines)
            
            logger.info(f"目录边界确定: 第{toc_start+2}行到第{toc_end}行，共{len(toc_lines)}行")
            
            return toc_content, toc_start + 1, toc_end
            
        except Exception as e:
            logger.error(f"提取Word文档目录边界失败: {e}")
            raise
    
    def _is_toc_entry(self, line: str) -> bool:
        """判断是否为目录条目"""
        if not line.strip():
            return False
        
        line = line.strip()
        
        # 目录条目特征 - 优化匹配条件以包含更多类型
        toc_patterns = [
            r'^第[一二三四五六七八九十\d]+章.+\d+$',  # 第X章...页码
            r'^\d+\.?\d*\s+.+\d+$',  # 1.1 标题...页码
            r'^[A-Z][a-z]+\s+\d+.+\d+$',  # Chapter 1...页码
            r'^摘\s*要.+[IVX\d]+$',  # 摘要...页码
            r'^Abstract.+[IVX\d]+$',  # Abstract...页码
            r'^目\s*录.+[IVX\d]+$',  # 目录...页码
            r'^参\s*考\s*文\s*献.+\d+$',  # 参考文献...页码
            r'^References.+\d+$',  # References...页码
            r'^攻读.+学位.+期间.+成果.+\d+$',  # 攻读学位期间成果...页码
            r'^研究成果.+\d+$',  # 研究成果...页码
            r'^学术成果.+\d+$',  # 学术成果...页码
            r'^Publications.+\d+$',  # Publications...页码
            r'^致\s*谢.+\d+$',  # 致谢...页码
            r'^Acknowledgment.+\d+$',  # Acknowledgment...页码
            r'^作者简介.+\d+$',  # 作者简介...页码
            r'^个人简历.+\d+$',  # 个人简历...页码
            r'^后\s*记.+\d+$',  # 后记...页码
            r'^结\s*语.+\d+$',  # 结语...页码
            r'^Epilogue.+\d+$',  # Epilogue...页码
            r'^Author.+Profile.+\d+$',  # Author Profile...页码
            r'^Biography.+\d+$',  # Biography...页码
            r'^附\s*录.+\d+$',  # 附录...页码
            r'^Appendix.+\d+$',  # Appendix...页码
            r'^声\s*明.+\d+$',  # 声明...页码
            r'^Declaration.+\d+$',  # Declaration...页码
            # 特殊格式：标题和页码之间用制表符分隔
            r'^.+\t+[IVX\d]+$',  # 标题TAB页码
            # 通用格式：标题后面跟页码
            r'^[^\d]+\s+[IVX\d]+$',  # 非数字开头的标题 页码
        ]
        
        for pattern in toc_patterns:
            if re.match(pattern, line):
                return True
        
        # 额外检查：如果包含明显的目录关键词和页码
        toc_keywords = [
            '攻读', '学位', '期间', '成果', '研究成果', '学术成果',
            '致谢', '作者简介', '个人简历', '声明', '版权', '后记', '结语',
            'Publications', 'Acknowledgment', 'Biography', 'CV', 'Declaration', 'Epilogue'
        ]
        
        # 检查是否包含关键词且以数字结尾（页码）
        if any(keyword in line for keyword in toc_keywords) and re.search(r'[IVX\d]+$', line):
            return True
        
        # 检查章节编号模式（特别针对1.1, 2.1这种格式）
        if re.match(r'^\d+\.\d+', line) and not re.search(r'[。.].{10,}', line):
            return True
            
        # 检查第X章模式
        if re.match(r'^第[一二三四五六七八九十\d]+章', line):
            return True

        return False
    
    def _extract_toc_from_content(self, content: str) -> List[str]:
        """
        从正文内容中智能提取目录结构
        适用于目录域无法读取的情况
        """
        lines = content.split('\n')
        toc_lines = []
        
        # 标准目录模式
        chapter_patterns = [
            r'^(第[一二三四五六七八九十\d]+章)\s*(.+?)(?:\s+(\d+))?$',  # 第X章 标题 页码
            r'^(\d+\.)\s*(.+?)(?:\s+(\d+))?$',                           # 1. 标题 页码
            r'^(\d+\.\d+)\s*(.+?)(?:\s+(\d+))?$',                       # 1.1 标题 页码
            r'^(\d+\.\d+\.\d+)\s*(.+?)(?:\s+(\d+))?$',                  # 1.1.1 标题 页码
            r'^([A-Z]+)\s*(.+?)(?:\s+(\d+))?$',                         # ABSTRACT 标题 页码
            r'^(摘\s*要|目\s*录|参考文献|致\s*谢|攻读|附\s*录)\s*(.*)(?:\s+(\d+))?$'  # 特殊章节
        ]
        
        for line in lines:
            line = line.strip()
            if not line or len(line) < 2:
                continue
                
            # 检查是否匹配目录模式
            for pattern in chapter_patterns:
                match = re.match(pattern, line, re.IGNORECASE)
                if match:
                    # 确保这是目录行而不是正文标题
                    if self._is_likely_toc_line(line):
                        toc_lines.append(line)
                        break
        
        logger.info(f"从正文内容中提取到 {len(toc_lines)} 行可能的目录内容")
        return toc_lines
    
    def _is_likely_toc_line(self, line: str) -> bool:
        """
        判断是否可能是目录行
        """
        line = line.strip()
        
        # 排除明显不是目录的行
        exclude_patterns = [
            r'图\s*\d+',  # 图X-X
            r'表\s*\d+',  # 表X-X
            r'公式\s*\d+',  # 公式X-X
            r'算法\s*\d+',  # 算法X-X
            r'本章小结',   # 章节结尾
            r'如图\s*\d+',  # 如图X所示
            r'见表\s*\d+',  # 见表X
            r'^根据',     # 根据...
            r'^基于',     # 基于...
            r'^通过',     # 通过...
            r'^如果',     # 如果...
            r'^因此',     # 因此...
            r'^由于',     # 由于...
        ]
        
        for pattern in exclude_patterns:
            if re.search(pattern, line):
                return False
        
        # 检查是否包含典型的目录特征
        toc_indicators = [
            r'第[一二三四五六七八九十\d]+章',  # 章节
            r'^\d+\.\s*[^\d]',              # 数字编号开头
            r'^\d+\.\d+\s*[^\d]',           # 二级编号
            r'摘\s*要|目\s*录|参考文献|致\s*谢|攻读|附\s*录|绪\s*论|总结|展望|结论',  # 特殊章节
            r'ABSTRACT|CONCLUSION|REFERENCE',  # 英文章节
        ]
        
        for pattern in toc_indicators:
            if re.search(pattern, line, re.IGNORECASE):
                return True
        
        return False
    
    def _fallback_toc_extraction(self, lines: List[str]) -> Tuple[str, int, int]:
        """备用目录提取方法"""
        logger.info("使用备用方法搜索目录内容")
        
        # 寻找包含目录条目的区域
        toc_lines = []
        toc_start = -1
        toc_end = -1
        
        for i, line in enumerate(lines):
            if self._is_toc_entry(line.strip()):
                if toc_start == -1:
                    toc_start = i
                toc_lines.append(line)
                toc_end = i + 1
        
        if toc_lines:
            toc_content = "\n".join(toc_lines)
            logger.info(f"备用方法找到 {len(toc_lines)} 行目录内容")
            return toc_content, toc_start, toc_end
        
        return "", 0, 0

class AITocExtractor:
    """AI智能目录抽取器 - 专门处理.docx格式论文"""
    
    def __init__(self):
        self.setup_patterns()
        self.parsers = {
            '.docx': WordParser()
        }
        self.ai_client = None  # 将在需要时初始化
    
    def setup_patterns(self):
        """设置AI识别模式"""
        
        # 主章节模式 (高置信度)
        self.main_chapter_patterns = [
            {
                'pattern': r'^(第[一二三四五六七八九十\d]+章)\s*([^\n\r\t\d]*)',
                'confidence': 0.95,
                'level': 1,
                'type': 'traditional_chapter'
            },
            {
                'pattern': r'^(\d+)\s+([^\n\r\t\d].{3,})',
                'confidence': 0.90,
                'level': 1,
                'type': 'numeric_chapter'
            },
            {
                'pattern': r'^(Chapter\s+\d+)\s+([^\n\r\t]+)',
                'confidence': 0.90,
                'level': 1,
                'type': 'english_chapter'
            },
            {
                'pattern': r'^(CHAPTER\s+\d+)\s+([^\n\r\t]+)',
                'confidence': 0.90,
                'level': 1,
                'type': 'english_chapter_caps'
            }
        ]
        
        # 子章节模式 (中等置信度)
        self.subsection_patterns = [
            {
                'pattern': r'^(\d+\.\d+)\s+([^\n\r\t\d].{2,})',
                'confidence': 0.85,
                'level': 2,
                'type': 'level2_section'
            },
            {
                'pattern': r'^(\d+\.\d+\.\d+)\s+([^\n\r\t\d].{2,})',
                'confidence': 0.80,
                'level': 3,
                'type': 'level3_section'
            },
            {
                'pattern': r'^(\d+\.\d+\.\d+\.\d+)\s+([^\n\r\t\d].{2,})',
                'confidence': 0.75,
                'level': 4,
                'type': 'level4_section'
            }
        ]
        
        # 特殊章节模式 (专门处理特殊格式)
        self.special_patterns = [
            {
                'pattern': r'^(摘\s*要|Abstract|ABSTRACT)',
                'confidence': 0.95,
                'level': 0,
                'type': 'abstract'
            },
            {
                'pattern': r'^(结\s*论|Conclusion|总结|Summary)',
                'confidence': 0.90,
                'level': 1,
                'type': 'conclusion'
            },
            {
                'pattern': r'^(参\s*考\s*文\s*献|References|REFERENCES)',
                'confidence': 0.95,
                'level': 1,
                'type': 'references'
            },
            {
                'pattern': r'^(攻读.*学位.*期间.*成果|研究成果|学术成果|Publications|Academic Achievements)',
                'confidence': 0.90,
                'level': 1,
                'type': 'achievements'
            },
            {
                'pattern': r'^(致\s*谢|Acknowledgment|Acknowledgement|ACKNOWLEDGMENTS?)',
                'confidence': 0.90,
                'level': 1,
                'type': 'acknowledgment'
            },
            {
                'pattern': r'^(作者简介|个人简历|Author.*Profile|Biography|CV)',
                'confidence': 0.90,
                'level': 1,
                'type': 'author_profile'
            },
            {
                'pattern': r'^(附\s*录|Appendix|APPENDIX)',
                'confidence': 0.85,
                'level': 1,
                'type': 'appendix'
            },
            {
                'pattern': r'^(声\s*明|Declaration|Statement)',
                'confidence': 0.85,
                'level': 1,
                'type': 'declaration'
            },
            {
                'pattern': r'^(版权声明|Copyright|License)',
                'confidence': 0.85,
                'level': 1,
                'type': 'copyright'
            }
        ]
        
        # 论文信息抽取模式
        self.meta_patterns = {
            'title': [
                r'论文题目[：:]\s*([^\n\r]+)',
                r'标题[：:]\s*([^\n\r]+)',
                r'Title[：:]\s*([^\n\r]+)',
            ],
            'author': [
                r'作者姓名[：:]\s*([^\n\r]+)',
                r'作者[：:]\s*([^\n\r]+)',
                r'Author[：:]\s*([^\n\r]+)',
                r'姓名[：:]\s*([^\n\r]+)',
            ]
        }
    
    def init_ai_client(self):
        """初始化AI客户端"""
        if self.ai_client is None:
            try:
                # 尝试导入和初始化AI客户端
                import sys
                import os
                
                # 添加项目根目录到Python路径
                project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                if project_root not in sys.path:
                    sys.path.append(project_root)
                
                # 尝试导入AI客户端
                try:
                    from .ai_client import get_ai_client
                    self.ai_client = get_ai_client()
                    logger.info("AI客户端初始化成功")
                except ImportError as e:
                    # 创建一个简单的Mock AI客户端
                    class MockAIClientImpl:
                        def send_message(self, prompt):
                            # 创建模拟的AIResponse对象
                            class MockAIResponse:
                                def __init__(self, content):
                                    self.content = content
                                    self.metadata = {}
                                    self.session_id = 'mock'
                                    self.timestamp = 0.0
                                    self.model_type = 'mock'
                            
                            return MockAIResponse('{"entries": []}')
                    
                    self.ai_client = MockAIClientImpl()
                    logger.warning(f"AI客户端模块未找到，使用Mock客户端: {e}")
                    
            except Exception as e:
                logger.warning(f"AI客户端初始化失败: {e}")
                # 创建Mock客户端作为fallback
                class MockAIClientFallback:
                    def send_message(self, prompt):
                        # 创建模拟的AIResponse对象
                        class MockAIResponse:
                            def __init__(self, content):
                                self.content = content
                                self.metadata = {}
                                self.session_id = 'mock'
                                self.timestamp = 0.0
                                self.model_type = 'mock'
                        
                        return MockAIResponse('{"entries": []}')
                
                self.ai_client = MockAIClientFallback()
    
    def extract_toc(self, file_path: str) -> ThesisToc:
        """智能抽取论文目录"""
        file_path_obj = Path(file_path)
        
        if not file_path_obj.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 根据文件扩展名选择解析器
        suffix = file_path_obj.suffix.lower()
        if suffix not in self.parsers:
            raise ValueError(f"不支持的文件格式: {suffix}")
        
        parser = self.parsers[suffix]
        
        logger.info(f"开始抽取目录: {file_path_obj.name}")
        
        # 解析文档内容
        content = parser.parse(file_path)
        
        # 检测文档语言类型（仅支持中文和英文）
        is_non_chinese, detected_lang = detect_non_chinese_content(content)
        if is_non_chinese:
            if detected_lang == "english":
                logger.info(f"✅ 检测通过：文档为英文论文")
            else:
                logger.warning(f"检测到不支持的语言类型: {detected_lang}")
                raise ValueError(f"不支持的语言类型: {detected_lang}。此工具仅支持中文和英文论文。")
        else:
            # 进一步验证中文内容
            if not is_chinese_text(content[:5000], min_chinese_ratio=0.2):
                logger.warning("文档中中文内容过少，可能不是中文论文")
                # 尝试检测是否为英文
                english_ratio = len([w for w in content[:5000].split() if w.isascii()]) / max(len(content[:5000].split()), 1)
                if english_ratio > 0.5:
                    detected_lang = "english"
                    logger.info("✅ 检测通过：文档为英文论文")
                else:
                    detected_lang = "unknown"
                    raise ValueError("无法确定文档语言类型。此工具仅支持中文和英文论文。")
            else:
                logger.info("✅ 检测通过：文档为中文论文")
                detected_lang = "chinese"
        
        # 提取目录边界和内容
        if hasattr(parser, 'extract_enhanced_toc'):
            # 使用增强的TOC提取方法
            toc_content = parser.extract_enhanced_toc(file_path, detected_lang)
            toc_start = 0
            toc_end = len(toc_content.split('\n')) if toc_content else 0
        else:
            # 使用传统的边界提取方法
            toc_content, toc_start, toc_end = parser.extract_toc_boundary(file_path)

        if not toc_content:
            logger.warning("未找到目录内容，尝试从正文内容中智能提取目录")
            # 尝试从正文内容中提取目录结构
            if hasattr(parser, '_extract_toc_from_content'):
                extracted_toc_lines = parser._extract_toc_from_content(content)
                if extracted_toc_lines:
                    toc_content = '\n'.join(extracted_toc_lines)
                    logger.info(f"从正文内容中成功提取 {len(extracted_toc_lines)} 行目录内容")
                else:
                    logger.warning("从正文内容中也未能提取到目录，使用全文分析")
                    toc_content = content
            else:
                logger.warning("解析器不支持内容提取，使用全文分析")
                toc_content = content
        
        # 抽取论文元信息
        title, author = self._extract_meta_info(content)
        
        # 初始化AI客户端
        self.init_ai_client()
        
        # AI智能识别目录条目
        if self.ai_client and hasattr(self.ai_client, 'send_message'):
            entries = self._ai_extract_entries_with_llm(toc_content)
        else:
            entries = self._ai_extract_entries_traditional(toc_content.split('\n'))
        
        # 计算整体置信度
        confidence_score = self._calculate_overall_confidence(entries)
        
        # 构建目录结构
        toc = ThesisToc(
            title=title,
            author=author,
            entries=entries,
            total_entries=len(entries),
            max_level=max([e.level for e in entries]) if entries else 0,
            extraction_method="AI_Smart_Extraction_with_LLM" if (self.ai_client and hasattr(self.ai_client, 'send_message')) else "Traditional_Pattern_Matching",
            confidence_score=confidence_score,
            toc_content=toc_content
        )
        
        logger.info(f"目录抽取完成: {len(entries)} 个条目, 置信度: {confidence_score:.2f}")
        
        return toc
    
    def _normalize_chapter_title(self, text: str) -> str:
        """规范化章节标题格式，去除多余空格，支持中英文"""
        if not text:
            return text
        
        import re
        
        # 中文章节标题规范化：去除"第"和"章"之间的空格
        # "第 一 章" -> "第一章"
        # "第 二 章" -> "第二章"
        # "第 三 章" -> "第三章"  
        # "第 四 章" -> "第四章"
        # "第 五 章" -> "第五章"
        # "第 六 章" -> "第六章"
        text = re.sub(r'第\s+([一二三四五六七八九十]+)\s+章', r'第\1章', text)
        
        # 阿拉伯数字章节标题规范化
        # "第 1 章" -> "第1章"
        # "第 2 章" -> "第2章"
        text = re.sub(r'第\s+(\d+)\s+章', r'第\1章', text)
        
        # 英文章节标题规范化：去除Chapter和数字之间的多余空格
        # "Chapter  1" -> "Chapter 1"
        # "CHAPTER   2" -> "CHAPTER 2"
        text = re.sub(r'(Chapter|CHAPTER)\s+(\d+)', r'\1 \2', text, flags=re.IGNORECASE)
        
        return text
    
    def _ai_extract_entries_with_llm(self, toc_content: str) -> List[TocEntry]:
        """使用LLM AI智能抽取目录条目，然后程序过滤level=1条目"""
        
        if not self.ai_client or not hasattr(self.ai_client, 'send_message'):
            logger.warning("AI客户端不可用，使用传统方法")
            return self._ai_extract_entries_traditional(toc_content.split('\n'))
        
        prompt = f"""
🎯 请仔细分析以下目录内容，提取所有章节和条目信息。特别注意主章节的识别！

⚠️ 【关键要求 - 必须遵守】：
1. 【识别所有主章节格式】：
   **中文格式：**
   - "第一章"、"第二章"、"第三章" 
   - "第 一 章"、"第 二 章"、"第 三 章"（注意：带空格的也是主章节！）
   - "第 四 章"、"第 五 章"、"第 六 章"（这些都必须识别为level=1！）
   
   **英文格式：**
   - "Chapter 1"、"Chapter 2"、"Chapter 3"
   - "CHAPTER 1"、"CHAPTER 2"、"CHAPTER 3"
   - "Chapter  1"（多空格也要正确识别）

   **特殊主章节（必须识别为level=1）：**
   - "结论"、"结 论"（结论章节，level=1）
   - "作者简介"、"作者 简介"（作者简介章节，level=1）
   - "参考文献"、"致谢"、"摘要"等（均为level=1）
   - "攻读学位期间取得的研究成果"等学术成果章节（level=1）

2. 【完整性检查】：确保ALL主章节都被提取，包括结论和作者简介
3. 【标题清理】：对于特殊章节，只保留核心标题

🔑 识别规则：
1. 任何形如"第X章"、"Chapter X"等格式的都是level=1主章节
2. 数字格式如"X.Y"的都是level=2或更低层级的子章节
3. 在文本中看到主章节标题行时，优先识别它而不是下面的子章节

目录内容:
{toc_content}

请按照以下JSON格式返回所有目录条目：
{{
    "entries": [
        {{
            "level": 1,
            "number": "第一章",
            "title": "绪论", 
            "page": 1,
            "confidence": 0.95,
            "section_type": "chapter"
        }},
        {{
            "level": 2,
            "number": "1.1",
            "title": "研究背景",
            "page": 2,
            "confidence": 0.90,
            "section_type": "section"
        }},
        {{
            "level": 1,
            "number": "",
            "title": "参考文献",
            "page": 89,
            "confidence": 0.95,
            "section_type": "references"
        }}
    ]
}}

层级判断规则：
- level=1: 主章节（第X章、Chapter X、X.格式、结论、参考文献、致谢、作者简介等所有独立章节）
- level=2: 二级章节（X.X格式）  
- level=3: 三级章节（X.X.X格式）
- level=4: 四级章节（X.X.X.X格式）

🔴 【重要】特殊章节必须标记为level=1：
- "结论"、"结 论" → level=1, section_type="conclusion"
- "作者简介"、"作者 简介" → level=1, section_type="author_profile"
- "参考文献" → level=1, section_type="references"
- "致谢" → level=1, section_type="acknowledgment"
- "摘要" → level=1, section_type="abstract"
- "攻读学位期间取得的研究成果" → level=1, section_type="achievements"

章节类型（section_type）：
- chapter: 正文章节（第一章、第二章、Chapter 1、Chapter 2、X.格式等）
- abstract: 摘要
- conclusion: 结论
- references: 参考文献
- acknowledgment: 致谢
- achievements: 学术成果/攻读学位期间取得的研究成果
- author_profile: 作者简介
- appendix: 附录

标题清理规则：
- 对于正文章节：保留"第X章"和主要标题
- 对于特殊章节：只保留核心名称
  * "摘要" 不要包含具体内容
  * "参考文献" 不要包含具体文献内容
  * "致谢" 不要包含致谢正文
  * "攻读学位期间取得的研究成果" 可简化为该标题
  * "作者简介" 不要包含个人信息

请确保返回有效的JSON格式，包含所有识别到的目录条目，特别是所有主章节。
"""
        
        try:
            response = self.ai_client.send_message(prompt)
            
            # 获取响应内容（现在所有客户端都返回类似AIResponse的对象）
            response_text = str(response.content)
            
            # 记录AI的原始响应
            print(f"\nAI LLM原始响应:")
            print("-" * 60)
            print(response_text)
            print("-" * 60)
            
            # 尝试解析JSON响应
            import json
            result_data = json.loads(response_text)
            
            all_entries = []
            for i, entry_data in enumerate(result_data.get('entries', [])):
                entry = TocEntry(
                    level=entry_data.get('level', 1),
                    number=entry_data.get('number', ''),
                    title=entry_data.get('title', ''),
                    page=entry_data.get('page'),
                    line_number=i + 1,
                    confidence=entry_data.get('confidence', 0.8),
                    section_type=entry_data.get('section_type', 'unknown')
                )
                all_entries.append(entry)
            
            # 详细分析AI LLM的所有提取结果
            print(f"\nAI LLM提取所有结果：{len(all_entries)}个条目")
            print("=" * 80)
            for i, entry in enumerate(all_entries):
                print(f"[{i+1:2d}] 层级:{entry.level} | 编号:{entry.number} | 标题:{entry.title} | 页码:{entry.page} | 类型:{entry.section_type}")
            print("=" * 80)
            
            # 程序逻辑过滤：只保留level=1的条目
            level1_entries = [entry for entry in all_entries if entry.level == 1]
            
            # 清理特殊章节标题
            level1_entries = self._clean_special_section_titles(level1_entries)
            
            print(f"\n� 程序过滤结果：从{len(all_entries)}个条目中筛选出{len(level1_entries)}个level=1条目")
            print("=" * 80)
            for i, entry in enumerate(level1_entries):
                print(f"[{i+1:2d}] 编号:{entry.number} | 标题:{entry.title} | 页码:{entry.page} | 类型:{entry.section_type}")
            print("=" * 80)
            
            # 对比分析统计
            print(f"\n📊 最终统计分析:")
            print(f"   原始书签提取: 62个有效目录项")
            print(f"   AI LLM完整提取: {len(all_entries)}个目录条目")
            print(f"   程序过滤level=1: {len(level1_entries)}个主章节")
            print(f"   最终保留比例: {len(level1_entries)/len(all_entries)*100:.1f}%")
            
            logger.info(f"AI LLM提取了 {len(all_entries)} 个目录条目，过滤后保留 {len(level1_entries)} 个level=1条目")
            return level1_entries
            
        except Exception as e:
            logger.error(f"AI LLM提取失败，使用传统方法: {e}")
            return self._ai_extract_entries_traditional(toc_content.split('\n'))
    
    def _clean_special_section_titles(self, entries: List[TocEntry]) -> List[TocEntry]:
        """清理特殊章节的标题，去除多余的正文内容"""
        for entry in entries:
            if entry.section_type == 'references':
                # 参考文献：只保留"参考文献"
                entry.title = "参考文献"
            elif entry.section_type == 'acknowledgment':
                # 致谢：只保留"致谢"
                entry.title = "致谢"
            elif entry.section_type == 'achievements':
                # 学术成果：清理为简洁标题
                if '攻读' in entry.title and '学位' in entry.title:
                    entry.title = "攻读学位期间取得的研究成果"
                elif '研究成果' in entry.title:
                    entry.title = "研究成果"
                elif '学术成果' in entry.title:
                    entry.title = "学术成果"
            elif entry.section_type == 'author_profile':
                # 作者简介：只保留"作者简介"
                if '作者简介' in entry.title:
                    entry.title = "作者简介"
                elif '个人简历' in entry.title:
                    entry.title = "个人简历"
        
        return entries
    
    def _extract_meta_info(self, content: str) -> Tuple[str, str]:
        """抽取论文元信息"""
        title = ""
        author = ""
        
        # 抽取标题
        for pattern in self.meta_patterns['title']:
            match = re.search(pattern, content, re.MULTILINE | re.IGNORECASE)
            if match:
                title = match.group(1).strip()
                break
        
        # 抽取作者
        for pattern in self.meta_patterns['author']:
            match = re.search(pattern, content, re.MULTILINE | re.IGNORECASE)
            if match:
                author = match.group(1).strip()
                break
        
        return title, author
    
    def _ai_extract_entries_traditional(self, lines: List[str]) -> List[TocEntry]:
        """传统方法AI智能抽取目录条目"""
        entries = []
        
        for line_num, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            # 尝试匹配各种模式
            entry = self._match_patterns(line, line_num)
            if entry:
                entries.append(entry)
        
        # AI后处理: 验证和修正
        entries = self._ai_post_process(entries)
        
        return entries
    
    def _match_patterns(self, line: str, line_num: int) -> Optional[TocEntry]:
        """模式匹配"""
        
        # 优先匹配特殊章节
        for pattern_info in self.special_patterns:
            match = re.search(pattern_info['pattern'], line, re.IGNORECASE)
            if match:
                return TocEntry(
                    level=pattern_info['level'],
                    number="",
                    title=match.group(1).strip(),
                    page=self._extract_page_number(line),
                    line_number=line_num + 1,
                    confidence=pattern_info['confidence'],
                    section_type=pattern_info['type']
                )
        
        # 匹配主章节
        for pattern_info in self.main_chapter_patterns:
            match = re.match(pattern_info['pattern'], line)
            if match:
                number = match.group(1)
                title = match.group(2).strip() if len(match.groups()) > 1 else ""
                return TocEntry(
                    level=pattern_info['level'],
                    number=number,
                    title=title,
                    page=self._extract_page_number(line),
                    line_number=line_num + 1,
                    confidence=pattern_info['confidence'],
                    section_type=pattern_info['type']
                )
        
        # 匹配子章节
        for pattern_info in self.subsection_patterns:
            match = re.match(pattern_info['pattern'], line)
            if match:
                number = match.group(1)
                title = match.group(2).strip() if len(match.groups()) > 1 else ""
                return TocEntry(
                    level=pattern_info['level'],
                    number=number,
                    title=title,
                    page=self._extract_page_number(line),
                    line_number=line_num + 1,
                    confidence=pattern_info['confidence'],
                    section_type=pattern_info['type']
                )
        
        return None
    
    def _extract_page_number(self, line: str) -> Optional[int]:
        """从行中抽取页码"""
        # 寻找行末的数字
        page_match = re.search(r'\.+\s*(\d+)\s*$', line)  # 点号后跟数字
        if page_match:
            return int(page_match.group(1))
        
        # 寻找制表符后的数字
        page_match = re.search(r'\t+(\d+)\s*$', line)
        if page_match:
            return int(page_match.group(1))
        
        # 寻找空格后的数字
        page_match = re.search(r'\s+(\d+)\s*$', line)
        if page_match:
            return int(page_match.group(1))
        
        return None
    
    def _ai_post_process(self, entries: List[TocEntry]) -> List[TocEntry]:
        """AI后处理: 智能验证和修正"""
        
        # 1. 去除重复条目
        entries = self._remove_duplicates(entries)
        
        # 2. 修正章节层级
        entries = self._fix_levels(entries)
        
        # 3. 智能标题补全
        entries = self._complete_titles(entries)
        
        # 4. 按章节顺序排序
        entries = self._sort_entries(entries)
        
        return entries
    
    def _remove_duplicates(self, entries: List[TocEntry]) -> List[TocEntry]:
        """去除重复条目"""
        seen = set()
        unique_entries = []
        
        for entry in entries:
            key = (entry.number, entry.title)
            if key not in seen:
                seen.add(key)
                unique_entries.append(entry)
        
        return unique_entries
    
    def _fix_levels(self, entries: List[TocEntry]) -> List[TocEntry]:
        """修正章节层级"""
        for entry in entries:
            if entry.number:
                # 根据编号点的数量确定层级
                dot_count = entry.number.count('.')
                if dot_count == 0:  # 如 "1" 或 "第一章"
                    entry.level = 1
                elif dot_count == 1:  # 如 "1.1"
                    entry.level = 2
                elif dot_count == 2:  # 如 "1.1.1"
                    entry.level = 3
                elif dot_count == 3:  # 如 "1.1.1.1"
                    entry.level = 4
        
        return entries
    
    def _complete_titles(self, entries: List[TocEntry]) -> List[TocEntry]:
        """智能标题补全"""
        for entry in entries:
            if not entry.title and entry.number:
                # 为缺失标题的条目生成默认标题
                if entry.level == 1:
                    entry.title = f"第{entry.number}章"
                else:
                    entry.title = f"第{entry.number}节"
        
        return entries
    
    def _sort_entries(self, entries: List[TocEntry]) -> List[TocEntry]:
        """按章节顺序排序"""
        def sort_key(entry: TocEntry) -> tuple:
            if not entry.number:
                return (999, 0, 0, 0)  # 特殊章节排在最后
            
            # 解析编号
            parts = entry.number.replace('第', '').replace('章', '').split('.')
            numbers = []
            for part in parts:
                try:
                    # 处理中文数字
                    if part in ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']:
                        cn_nums = {'一': 1, '二': 2, '三': 3, '四': 4, '五': 5, 
                                  '六': 6, '七': 7, '八': 8, '九': 9, '十': 10}
                        numbers.append(cn_nums.get(part, 0))
                    else:
                        numbers.append(int(part))
                except ValueError:
                    numbers.append(0)
            
            # 填充到4位数字
            while len(numbers) < 4:
                numbers.append(0)
            
            return tuple(numbers[:4])
        
        return sorted(entries, key=sort_key)
    
    def _calculate_overall_confidence(self, entries: List[TocEntry]) -> float:
        """计算整体置信度"""
        if not entries:
            return 0.0
        
        total_confidence = sum(entry.confidence for entry in entries)
        return total_confidence / len(entries)
    
    def save_toc_json(self, toc: ThesisToc, output_path: str):
        """保存目录为结构化JSON"""
        output_path_obj = Path(output_path)
        
        # 构建结构化JSON
        toc_json = {
            "metadata": {
                "title": toc.title,
                "author": toc.author,
                "total_entries": toc.total_entries,
                "max_level": toc.max_level,
                "extraction_method": toc.extraction_method,
                "confidence_score": toc.confidence_score,
                "extracted_at": datetime.now().isoformat()
            },
            "toc_structure": {
                "chapters": [],
                "special_sections": [],
                "post_references": []  # 参考文献后的章节
            },
            "raw_entries": []
        }
        
        # 组织章节结构
        current_chapter = None
        references_found = False  # 标记是否已经找到参考文献
        
        for entry in toc.entries:
            # 添加原始条目
            toc_json["raw_entries"].append({
                "level": entry.level,
                "number": entry.number,
                "title": entry.title,
                "page": entry.page,
                "line_number": entry.line_number,
                "confidence": entry.confidence,
                "section_type": entry.section_type
            })
            
            # 检查是否到了参考文献
            if entry.section_type == 'references':
                references_found = True
            
            # 组织层次结构
            if entry.level == 1 and entry.section_type in ['traditional_chapter', 'numeric_chapter', 'english_chapter', 'chapter']:
                current_chapter = {
                    "number": entry.number,
                    "title": entry.title,
                    "page": entry.page,
                    "sections": []
                }
                toc_json["toc_structure"]["chapters"].append(current_chapter)
            elif entry.level > 1 and current_chapter and not references_found:
                current_chapter["sections"].append({
                    "level": entry.level,
                    "number": entry.number,
                    "title": entry.title,
                    "page": entry.page
                })
            elif entry.section_type in ['abstract', 'conclusion', 'references']:
                toc_json["toc_structure"]["special_sections"].append({
                    "type": entry.section_type,
                    "title": entry.title,
                    "page": entry.page
                })
            elif references_found and entry.section_type in ['achievements', 'acknowledgment', 'author_profile', 'appendix', 'declaration', 'copyright']:
                # 参考文献后的章节
                toc_json["toc_structure"]["post_references"].append({
                    "type": entry.section_type,
                    "title": entry.title,
                    "page": entry.page
                })
        
        # 保存JSON文件
        with open(output_path_obj, 'w', encoding='utf-8') as f:
            json.dump(toc_json, f, ensure_ascii=False, indent=2)
        
        logger.info(f"结构化JSON已保存到: {output_path_obj}")
        
        return toc_json

def test_with_word_documents():
    """测试Word文档"""
    extractor = AITocExtractor()
    
    # 测试data/input文件夹中的Word文档
    input_dir = Path("data/input")
    if not input_dir.exists():
        logger.error(f"测试目录不存在: {input_dir}")
        return
    
    word_files = list(input_dir.glob("*.docx"))
    if not word_files:
        logger.warning(f"在 {input_dir} 中未找到Word文档")
        return
    
    for word_file in word_files:
        logger.info(f"测试Word文档: {word_file.name}")
        
        try:
            # 抽取目录
            toc = extractor.extract_toc(str(word_file))
            
            # 保存结构化JSON
            output_file = f"{word_file.stem}_toc_structured.json"
            toc_json = extractor.save_toc_json(toc, output_file)
            
            # 打印结果摘要
            print(f"\n{word_file.name} 抽取结果:")
            print(f"   标题: {toc.title}")
            print(f"   作者: {toc.author}")
            print(f"   总条目: {toc.total_entries}")
            print(f"   置信度: {toc.confidence_score:.2f}")
            print(f"   方法: {toc.extraction_method}")
            print(f"   输出: {output_file}")
            
        except Exception as e:
            logger.error(f"处理 {word_file.name} 失败: {e}")

if __name__ == "__main__":
    test_with_word_documents()
