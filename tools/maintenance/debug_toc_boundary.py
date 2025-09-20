"""
专门调试马克思主义哲学论文的目录边界问题
"""
import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), 'src'))

import docx

def debug_toc_boundary():
    """调试目录边界问题"""
    doc_path = r"c:\MyProjects\thesis_Inno_Eval\data\input\1_马克思主义哲学86406_010101_81890101_LW.docx"
    
    print("🔍 调试马克思主义哲学论文目录边界")
    print("="*80)
    
    try:
        doc = docx.Document(doc_path)
        
        print("📄 检查第25-35行内容 (查找后记):")
        print("-" * 60)
        
        for i in range(24, 35):  # 第25-35行
            if i < len(doc.paragraphs):
                text = doc.paragraphs[i].text.strip()
                if text:
                    print(f"第{i+1:3d}行: {text}")
                    if "后" in text and "记" in text:
                        print(f"     ▶ 发现后记相关内容!")
        
        # 直接应用目录提取逻辑
        print(f"\n🔧 应用目录提取逻辑:")
        print("-" * 60)
        
        # 模拟AI目录提取器的逻辑
        toc_start = -1
        toc_end = -1
        
        for i, paragraph in enumerate(doc.paragraphs):
            line = paragraph.text.strip()
            
            # 查找目录开始
            if toc_start == -1 and "目  录" in line:
                toc_start = i
                print(f" 目录开始: 第{i+1}行 - {line}")
            
            # 查找目录结束
            elif toc_start != -1 and toc_end == -1:
                # 检查是否匹配结束标志
                if line and not line.startswith("目"):
                    # 应用新的结束检测逻辑
                    import re
                    end_indicators = [
                        r'^第[一二三四五六七八九十\d]+章\s+\S+',
                        r'^\d+\s+[^.\d\s]\S+',
                        r'^Chapter\s+\d+\s+\S+',
                        r'^引\s*言\s*$',
                        r'^绪\s*论\s*$',
                        r'^Introduction\s*$'
                    ]
                    
                    should_end = False
                    for indicator in end_indicators:
                        if re.match(indicator, line, re.IGNORECASE):
                            # 确认不是目录条目
                            if not is_toc_entry_debug(line):
                                should_end = True
                                print(f"⚠️  匹配结束标志: 第{i+1}行 - {line}")
                                break
                    
                    if should_end:
                        toc_end = i
                        print(f" 目录结束: 第{i+1}行 - {line}")
                        break
        
        if toc_start != -1:
            if toc_end == -1:
                toc_end = min(toc_start + 50, len(doc.paragraphs))
                print(f" 未找到明确结束，设置为: 第{toc_end}行")
            
            print(f"\n📋 目录内容 (第{toc_start+1}行到第{toc_end}行):")
            print("-" * 80)
            
            for i in range(toc_start, toc_end):
                if i < len(doc.paragraphs):
                    text = doc.paragraphs[i].text.strip()
                    if text:
                        is_toc = is_toc_entry_debug(text)
                        marker = "" if is_toc else "❌"
                        print(f"第{i+1:3d}行: {marker} {text}")
                        if "后" in text and "记" in text:
                            print(f"     ▶ 后记相关条目!")
        
    except Exception as e:
        print(f"❌ 调试失败: {str(e)}")
        import traceback
        traceback.print_exc()

def is_toc_entry_debug(line: str) -> bool:
    """调试版本的目录条目判断"""
    if not line.strip():
        return False
    
    line = line.strip()
    
    import re
    # 目录条目特征 - 包含后记的识别
    toc_patterns = [
        r'^第[一二三四五六七八九十\d]+章.+\d+$',  # 第X章...页码
        r'^\d+\.?\d*\s+.+\d+$',  # 1.1 标题...页码
        r'^摘\s*要.+[IVX\d]+$',  # 摘要...页码
        r'^Abstract.+[IVX\d]+$',  # Abstract...页码
        r'^参\s*考\s*文\s*献.+\d+$',  # 参考文献...页码
        r'^个人简历.+\d+$',  # 个人简历...页码
        r'^后\s*记.+\d+$',  # 后记...页码 - 专门添加
        r'^结\s*束\s*语.+\d+$',  # 结束语...页码
        r'^.+\s+\d+$',  # 标题 页码（最宽泛的匹配）
    ]
    
    for pattern in toc_patterns:
        if re.match(pattern, line):
            return True
    
    return False

if __name__ == "__main__":
    debug_toc_boundary()

