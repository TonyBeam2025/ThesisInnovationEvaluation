#!/usr/bin/env python3
"""
调试Word文档内容和TOC字段
"""

import docx
import os

def debug_word_document(file_path):
    """调试Word文档的详细内容"""
    print(f"📄 分析文件: {file_path}")
    print("=" * 60)
    
    if not os.path.exists(file_path):
        print(f"❌ 文件不存在")
        return
    
    try:
        doc = docx.Document(file_path)
        
        print(f"📊 文档统计:")
        print(f"  段落数量: {len(doc.paragraphs)}")
        print(f"  表格数量: {len(doc.tables)}")
        print(f"  节数量: {len(doc.sections)}")
        
        print(f"\n📝 段落内容分析:")
        total_chars = 0
        non_empty_paragraphs = 0
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text
            total_chars += len(text)
            
            if text.strip():
                non_empty_paragraphs += 1
                print(f"  段落 {i+1}: {repr(text[:100])} {'...' if len(text) > 100 else ''}")
                
                # 检查段落的样式信息
                if paragraph.style:
                    print(f"    样式: {paragraph.style.name}")
                
                # 检查段落的XML
                if hasattr(paragraph, '_element'):
                    xml_text = paragraph._element.xml
                    if 'TOC' in xml_text or 'fldChar' in xml_text:
                        print(f"    🔍 包含TOC字段信息")
                        print(f"    XML片段: {xml_text[:200]}...")
        
        print(f"\n📊 内容统计:")
        print(f"  总字符数: {total_chars}")
        print(f"  非空段落数: {non_empty_paragraphs}")
        print(f"  平均段落长度: {total_chars/len(doc.paragraphs) if doc.paragraphs else 0:.1f}")
        
        # 检查文档属性
        if hasattr(doc, 'core_properties'):
            props = doc.core_properties
            print(f"\n📋 文档属性:")
            print(f"  标题: {props.title}")
            print(f"  作者: {props.author}")
            print(f"  主题: {props.subject}")
            print(f"  创建时间: {props.created}")
            print(f"  修改时间: {props.modified}")
        
    except Exception as e:
        print(f"❌ 分析失败: {e}")
        import traceback
        traceback.print_exc()

def test_multiple_files():
    """测试多个文件"""
    test_files = [
        "data/input/计算机应用技术_test2.docx",
        "data/input/计算机应用技术_test1.docx",
        "data/input/1_计算机应用技术_17211204005-苏慧婧-基于MLP和SepCNN模型的藏文文本分类研究与实现-计算机应用技术-群诺.docx"
    ]
    
    for file_path in test_files:
        debug_word_document(file_path)
        print("\n" + "="*80 + "\n")

if __name__ == "__main__":
    test_multiple_files()
