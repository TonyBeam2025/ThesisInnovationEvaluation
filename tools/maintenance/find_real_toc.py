#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
精确分析Word文档中的真实TOC内容
"""

from docx import Document
import re
import xml.etree.ElementTree as ET

def find_real_toc_content(docx_path):
    """找到真实的目录内容"""
    print(f"📄 分析真实TOC内容: {docx_path}")
    print("=" * 80)
    
    try:
        doc = Document(docx_path)
        
        toc_candidates = []
        
        for i, paragraph in enumerate(doc.paragraphs, 1):
            text = paragraph.text.strip()
            
            # 跳过空段落
            if not text:
                continue
            
            # 查找可能的目录条目模式
            toc_patterns = [
                # 中文章节标题 + 页码
                r'^(.+?)\s*\.{2,}\s*(\d+)$',  # 标题...页码
                r'^(.+?)\s+(\d+)$',           # 标题 页码  
                r'^(\d+[\.\)]\s*.+?)\s*\.{2,}\s*(\d+)$',  # 1. 标题...页码
                r'^(第[一二三四五六七八九十\d]+章\s*.+?)\s*\.{2,}\s*(\d+)$',  # 第X章 标题...页码
                r'^(\d+[\.\)]\d*[\.\)]*\s*.+?)\s+(\d+)$',  # 1.1 标题 页码
                # 英文章节标题
                r'^([A-Z][a-zA-Z\s]+)\s*\.{2,}\s*(\d+)$',  # ABSTRACT...1
                # 超链接格式
                r'^(.+?)\s*\[(\d+)\]$',       # 标题 [页码]
            ]
            
            # 检查样式是否像目录
            style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else 'normal'
            is_toc_style = any(keyword in style_name for keyword in [
                'toc', 'table', 'content', '目录', 'heading'
            ])
            
            # 检查文本是否符合目录模式
            for pattern in toc_patterns:
                match = re.match(pattern, text)
                if match:
                    title = match.group(1).strip()
                    page = match.group(2).strip()
                    
                    # 过滤明显不是目录的内容
                    if len(title) > 100:  # 标题太长
                        continue
                    if not page.isdigit() or int(page) > 1000:  # 页码不合理
                        continue
                    
                    toc_candidates.append({
                        'paragraph_index': i,
                        'title': title,
                        'page': page,
                        'raw_text': text,
                        'style': paragraph.style.name if paragraph.style and paragraph.style.name else 'Normal',
                        'is_toc_style': is_toc_style,
                        'pattern': pattern
                    })
                    
                    print(f"🔍 可能的目录条目 {i}: {text}")
                    print(f"   标题: {title}")
                    print(f"   页码: {page}")
                    print(f"   样式: {paragraph.style.name if paragraph.style and paragraph.style.name else 'Normal'}")
                    print(f"   TOC样式: {is_toc_style}")
                    print()
                    break
        
        # 查找连续的目录区域
        if toc_candidates:
            print(f"\n📊 找到 {len(toc_candidates)} 个可能的目录条目")
            
            # 按段落索引分组，查找连续区域
            continuous_groups = []
            current_group = [toc_candidates[0]]
            
            for i in range(1, len(toc_candidates)):
                curr = toc_candidates[i]
                prev = toc_candidates[i-1]
                
                # 如果段落索引相近（允许中间有少量空段落）
                if curr['paragraph_index'] - prev['paragraph_index'] <= 5:
                    current_group.append(curr)
                else:
                    # 保存当前组，开始新组
                    if len(current_group) >= 3:  # 至少3个条目才认为是目录
                        continuous_groups.append(current_group)
                    current_group = [curr]
            
            # 添加最后一组
            if len(current_group) >= 3:
                continuous_groups.append(current_group)
            
            print(f"\n📋 找到 {len(continuous_groups)} 个连续的目录区域:")
            
            for group_idx, group in enumerate(continuous_groups):
                print(f"\n目录区域 {group_idx + 1} (共{len(group)}个条目):")
                print(f"  段落范围: {group[0]['paragraph_index']} - {group[-1]['paragraph_index']}")
                
                for entry in group[:5]:  # 只显示前5个条目
                    print(f"    {entry['title']} -> 第{entry['page']}页")
                
                if len(group) > 5:
                    print(f"    ... 还有 {len(group) - 5} 个条目")
        
        else:
            print("❌ 未找到符合目录模式的内容")
            
            # 尝试查找一些常见的目录开始标记
            print("\n🔍 查找可能的目录开始标记:")
            for i, paragraph in enumerate(doc.paragraphs[:50], 1):  # 只检查前50段
                text = paragraph.text.strip().lower()
                if any(keyword in text for keyword in ['目录', 'contents', 'table of contents', '目　录']):
                    print(f"   段落 {i}: {paragraph.text}")
        
    except Exception as e:
        print(f"❌ 分析失败: {e}")
        import traceback
        traceback.print_exc()

def main():
    # 分析文档
    test_files = [
        "data/input/1_计算机应用技术_17211204005-苏慧婧-基于MLP和SepCNN模型的藏文文本分类研究与实现-计算机应用技术-群诺.docx",
        "data/input/计算机应用技术_test1.docx"
    ]
    
    for file_path in test_files:
        find_real_toc_content(file_path)
        print("\n" + "="*80 + "\n")

if __name__ == "__main__":
    main()
