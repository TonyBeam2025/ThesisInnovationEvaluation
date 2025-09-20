#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
深度分析文档XML结构，查找隐藏的目录信息
"""

from docx import Document
from docx.oxml.ns import qn
import xml.etree.ElementTree as ET
import re

def analyze_xml_structure(doc_path):
    """分析Word文档的完整XML结构"""
    print(f"🔍 深度分析Word文档XML结构: {doc_path}")
    print("=" * 80)
    
    doc = Document(doc_path)
    
    # 1. 分析文档部分的XML
    print("\n📋 文档部分XML分析:")
    print("-" * 40)
    
    # 获取文档XML
    document_xml = doc._element.xml
    
    # 查找TOC相关的XML元素
    toc_patterns = [
        r'<w:fldChar[^>]*w:fldCharType="begin"[^>]*>.*?TOC.*?<w:fldChar[^>]*w:fldCharType="end"[^>]*>',
        r'<w:hyperlink[^>]*w:anchor="[^"]*"[^>]*>.*?</w:hyperlink>',
        r'<w:instrText>.*?TOC.*?</w:instrText>',
        r'<w:bookmarkStart[^>]*w:name="_Toc[^"]*"[^>]*/>',
        r'<w:bookmarkEnd[^>]*w:id="[^"]*"[^>]*/>',
    ]
    
    for i, pattern in enumerate(toc_patterns, 1):
        matches = re.findall(pattern, document_xml, re.DOTALL | re.IGNORECASE)
        print(f"\n模式 {i} 匹配结果 ({len(matches)} 个):")
        for j, match in enumerate(matches[:5]):  # 只显示前5个
            print(f"  {j+1}. {match[:200]}...")
    
    # 2. 分析超链接
    print("\n🔗 超链接分析:")
    print("-" * 40)
    
    hyperlink_pattern = r'<w:hyperlink[^>]*w:anchor="([^"]*)"[^>]*>(.*?)</w:hyperlink>'
    hyperlinks = re.findall(hyperlink_pattern, document_xml, re.DOTALL)
    
    print(f"找到 {len(hyperlinks)} 个超链接:")
    for i, (anchor, content) in enumerate(hyperlinks[:10]):  # 显示前10个
        # 提取文本内容
        text_content = re.sub(r'<[^>]+>', '', content).strip()
        print(f"  {i+1}. 锚点: {anchor}")
        print(f"     文本: {text_content}")
        print()
    
    # 3. 分析书签
    print("\n📑 书签分析:")
    print("-" * 40)
    
    bookmark_start_pattern = r'<w:bookmarkStart[^>]*w:name="([^"]*)"[^>]*/>'
    bookmarks = re.findall(bookmark_start_pattern, document_xml)
    
    toc_bookmarks = [bm for bm in bookmarks if '_Toc' in bm]
    print(f"找到 {len(toc_bookmarks)} 个TOC书签:")
    for i, bookmark in enumerate(toc_bookmarks[:20]):  # 显示前20个
        print(f"  {i+1}. {bookmark}")
    
    # 4. 分析字段代码
    print("\n🔧 字段代码分析:")
    print("-" * 40)
    
    field_pattern = r'<w:instrText[^>]*>(.*?)</w:instrText>'
    fields = re.findall(field_pattern, document_xml, re.DOTALL)
    
    toc_fields = [field for field in fields if 'TOC' in field.upper()]
    print(f"找到 {len(toc_fields)} 个TOC字段:")
    for i, field in enumerate(toc_fields):
        print(f"  {i+1}. {field.strip()}")
    
    # 5. 查找标题样式
    print("\n📝 标题样式分析:")
    print("-" * 40)
    
    heading_pattern = r'<w:pStyle[^>]*w:val="(Heading[^"]*)"[^>]*/>'
    headings = re.findall(heading_pattern, document_xml)
    
    unique_headings = sorted(set(headings))
    print(f"找到 {len(unique_headings)} 种标题样式:")
    for heading in unique_headings:
        print(f"  - {heading}")
    
    # 6. 特殊分析：段落18的XML内容
    print("\n🕵️ 段落18 XML详细分析:")
    print("-" * 40)
    
    if len(doc.paragraphs) > 18:
        para18 = doc.paragraphs[18]
        para18_xml = para18._element.xml
        
        print(f"段落18 XML长度: {len(para18_xml)} 字符")
        print("XML内容片段:")
        print(para18_xml[:1000])
        print("...")
        print(para18_xml[-500:])
        
        # 查找特殊模式
        special_patterns = [
            r'<w:fldChar[^>]*>',
            r'<w:instrText[^>]*>.*?</w:instrText>',
            r'<w:fldSimple[^>]*>',
            r'<w:hyperlink[^>]*>',
            r'<w:bookmarkStart[^>]*>',
        ]
        
        print("\n段落18中的特殊XML元素:")
        for pattern in special_patterns:
            matches = re.findall(pattern, para18_xml, re.DOTALL)
            if matches:
                print(f"  - {pattern}: {len(matches)} 个匹配")
                for match in matches[:3]:
                    print(f"    {match}")

def extract_toc_from_bookmarks():
    """通过书签提取目录结构"""
    doc_path = "data/input/1_计算机应用技术_17211204005-苏慧婧-基于MLP和SepCNN模型的藏文文本分类研究与实现-计算机应用技术-群诺.docx"
    
    print("\n📚 通过书签提取目录结构:")
    print("-" * 40)
    
    doc = Document(doc_path)
    document_xml = doc._element.xml
    
    # 查找所有TOC书签和对应的文本
    bookmark_pattern = r'<w:bookmarkStart[^>]*w:name="(_Toc\d+)"[^>]*/>(.*?)<w:bookmarkEnd[^>]*>'
    bookmark_matches = re.findall(bookmark_pattern, document_xml, re.DOTALL)
    
    toc_entries = []
    for bookmark_name, content in bookmark_matches:
        # 提取文本内容
        text_content = re.sub(r'<[^>]+>', '', content).strip()
        if text_content:
            toc_entries.append((bookmark_name, text_content))
    
    print(f"通过书签找到 {len(toc_entries)} 个目录项:")
    for i, (bookmark, text) in enumerate(toc_entries[:20]):
        print(f"  {i+1}. {bookmark}: {text}")
    
    return toc_entries

if __name__ == "__main__":
    doc_path = "data/input/1_计算机应用技术_17211204005-苏慧婧-基于MLP和SepCNN模型的藏文文本分类研究与实现-计算机应用技术-群诺.docx"
    
    # 深度分析XML结构
    analyze_xml_structure(doc_path)
    
    # 通过书签提取目录
    extract_toc_from_bookmarks()
