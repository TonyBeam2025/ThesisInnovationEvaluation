#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
分析Word文档的实际章节结构
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

import docx
import re

def analyze_word_structure():
    """分析Word文档结构"""
    print("📖 分析Word文档结构...")
    
    file_path = r".\data\input\跨模态图像融合技术在医疗影像分析中的研究.docx"
    
    if not os.path.exists(file_path):
        print(f"❌ 文件不存在: {file_path}")
        return
    
    try:
        # 读取Word文档
        doc = docx.Document(file_path)
        
        print(f"\n📄 文档段落总数: {len(doc.paragraphs)}")
        
        # 分析段落样式
        styles = {}
        for para in doc.paragraphs:
            style_name = para.style.name if para.style else "Normal"
            if style_name not in styles:
                styles[style_name] = 0
            styles[style_name] += 1
        
        print(f"\n🎨 段落样式分布:")
        for style, count in sorted(styles.items(), key=lambda x: x[1], reverse=True):
            print(f"   {style}: {count}")
        
        # 查找可能的章节标题
        print(f"\n🔍 查找可能的章节标题...")
        potential_headings = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:
                # 检查是否是标题样式
                style_name = para.style.name if para.style else "Normal"
                
                # 检查是否匹配章节模式
                chapter_patterns = [
                    r'^第[一二三四五六七八九十\d]+章',
                    r'^第\d+章',
                    r'^Chapter\s+\d+',
                    r'^绪论$',
                    r'^引言$',
                    r'^文献综述$',
                    r'^相关工作$',
                    r'^国内外研究现状$',
                    r'^结\s*论$',
                    r'^总\s*结$',
                    r'^参考文献$',
                    r'^致\s*谢$',
                    r'^谢\s*辞$',
                    r'^附\s*录$',
                ]
                
                is_chapter = any(re.match(pattern, text, re.IGNORECASE) for pattern in chapter_patterns)
                
                # 检查字体属性
                is_bold = False
                font_size = None
                try:
                    if para.runs:
                        first_run = para.runs[0]
                        is_bold = first_run.bold if first_run.bold is not None else False
                        font_size = first_run.font.size
                except:
                    pass
                
                if (is_chapter or 
                    (style_name and "Heading" in style_name) or 
                    len(text) < 50 and (is_bold or font_size)):
                    
                    potential_headings.append({
                        'index': i,
                        'text': text,
                        'style': style_name,
                        'is_bold': is_bold,
                        'font_size': font_size,
                        'is_chapter_pattern': is_chapter
                    })
        
        print(f"\n📋 潜在章节标题 ({len(potential_headings)} 个):")
        for heading in potential_headings:
            print(f"   [{heading['index']}] {heading['text']}")
            print(f"       样式: {heading['style']}, 粗体: {heading['is_bold']}, 字号: {heading['font_size']}, 章节模式: {heading['is_chapter_pattern']}")
        
        # 查找参考文献部分
        print(f"\n📚 查找参考文献部分...")
        ref_found = False
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if re.match(r'^(参考文献|REFERENCES?)$', text, re.IGNORECASE):
                print(f"   找到参考文献标题: [{i}] {text}")
                ref_found = True
                
                # 查看后续几个段落
                print(f"   后续内容:")
                for j in range(i+1, min(i+6, len(doc.paragraphs))):
                    next_text = doc.paragraphs[j].text.strip()
                    if next_text:
                        print(f"     [{j}] {next_text[:100]}...")
                break
        
        if not ref_found:
            print("   ❌ 未找到参考文献标题")
        
        # 查找致谢部分
        print(f"\n🙏 查找致谢部分...")
        ack_found = False
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if re.match(r'^(致\s*谢|谢\s*辞|ACKNOWLEDGEMENTS?)$', text, re.IGNORECASE):
                print(f"   找到致谢标题: [{i}] {text}")
                ack_found = True
                
                # 查看后续几个段落
                print(f"   后续内容:")
                for j in range(i+1, min(i+6, len(doc.paragraphs))):
                    next_text = doc.paragraphs[j].text.strip()
                    if next_text:
                        print(f"     [{j}] {next_text[:100]}...")
                break
        
        if not ack_found:
            print("   ❌ 未找到致谢标题")
        
    except Exception as e:
        print(f"❌ 分析过程中出现错误: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    analyze_word_structure()
