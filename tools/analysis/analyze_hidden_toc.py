#!/usr/bin/env python3
"""
检查Word文档中的隐藏目录域和空段落
专门查找段落17-19之间的目录内容
"""

import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
import logging

def analyze_hidden_toc_fields(file_path):
    """分析隐藏的TOC字段"""
    print(f"🔍 分析文件: {file_path}")
    print("="*60)
    
    try:
        doc = docx.Document(file_path)
        
        # 重点检查段落17-19附近
        target_range = range(15, 25)  # 扩大范围以确保覆盖
        
        for i, paragraph in enumerate(doc.paragraphs):
            if i in target_range:
                print(f"\n段落 {i} 详细分析:")
                print(f"  文本: '{paragraph.text}'")
                print(f"  样式: {paragraph.style.name if paragraph.style else 'None'}")
                print(f"  段落XML长度: {len(paragraph._element.xml)}")
                
                # 检查段落的XML内容
                para_xml = paragraph._element.xml
                print(f"  包含字段代码: {'fldChar' in para_xml}")
                print(f"  包含TOC: {'TOC' in para_xml}")
                print(f"  包含超链接: {'hyperlink' in para_xml}")
                
                # 如果是空段落但XML很长，说明可能包含隐藏内容
                if not paragraph.text.strip() and len(para_xml) > 100:
                    print(f"  ⚠️  空段落但XML内容丰富，可能包含隐藏的TOC字段")
                    print(f"  XML片段: {para_xml[:200]}...")
                    
                    # 尝试解析XML中的TOC内容
                    if 'TOC' in para_xml:
                        print("  🎯 发现TOC字段!")
                        toc_content = extract_toc_from_xml(para_xml)
                        if toc_content:
                            print(f"  TOC内容: {toc_content}")
                
                # 检查runs中的字段
                for j, run in enumerate(paragraph.runs):
                    run_xml = run._element.xml
                    if 'fldChar' in run_xml or 'TOC' in run_xml:
                        print(f"    Run {j}: 包含字段信息")
                        print(f"    Run文本: '{run.text}'")
                        print(f"    Run XML: {run_xml[:150]}...")
        
        print("\n" + "="*60)
        print("🔍 搜索文档中所有的TOC字段...")
        
        # 在整个文档中搜索TOC字段
        all_toc_content = []
        for i, paragraph in enumerate(doc.paragraphs):
            para_xml = paragraph._element.xml
            if 'TOC' in para_xml or 'fldChar' in para_xml:
                print(f"\n段落 {i} 包含字段:")
                print(f"  文本: '{paragraph.text}'")
                
                # 尝试提取TOC内容
                toc_content = extract_toc_from_xml(para_xml)
                if toc_content:
                    all_toc_content.extend(toc_content)
                    print(f"  提取的TOC内容: {toc_content}")
        
        if all_toc_content:
            print(f"\n📋 找到的所有TOC条目:")
            for i, entry in enumerate(all_toc_content, 1):
                print(f"{i:2d}. {entry}")
        else:
            print("\n❌ 未找到TOC字段内容")
            
    except Exception as e:
        print(f"❌ 分析失败: {e}")
        import traceback
        traceback.print_exc()

def extract_toc_from_xml(xml_content):
    """从XML内容中提取TOC条目"""
    import re
    
    toc_entries = []
    
    try:
        # 查找可能的TOC条目模式
        patterns = [
            r'<w:t[^>]*>([^<]+)</w:t>',  # 文本内容
            r'<w:hyperlink[^>]*>.*?<w:t[^>]*>([^<]+)</w:t>.*?</w:hyperlink>',  # 超链接文本
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, xml_content, re.DOTALL)
            for match in matches:
                text = match.strip()
                if text and len(text) > 3:
                    # 检查是否像目录条目
                    if (any(keyword in text for keyword in ['章', '节', '摘要', '绪论', '参考文献', '致谢']) or
                        re.match(r'\d+', text) or
                        '.' in text):
                        toc_entries.append(text)
    
    except Exception as e:
        print(f"XML解析错误: {e}")
    
    return toc_entries

def analyze_document_structure(file_path):
    """分析文档的整体结构"""
    print(f"\n📊 文档结构分析:")
    print("-"*40)
    
    try:
        doc = docx.Document(file_path)
        
        # 统计不同类型的段落
        style_counts = {}
        empty_paragraphs = []
        field_paragraphs = []
        
        for i, paragraph in enumerate(doc.paragraphs):
            style_name = paragraph.style.name if paragraph.style else 'None'
            style_counts[style_name] = style_counts.get(style_name, 0) + 1
            
            if not paragraph.text.strip():
                empty_paragraphs.append(i)
            
            para_xml = paragraph._element.xml
            if 'fldChar' in para_xml or 'TOC' in para_xml:
                field_paragraphs.append(i)
        
        print(f"总段落数: {len(doc.paragraphs)}")
        print(f"空段落: {len(empty_paragraphs)} 个")
        print(f"包含字段的段落: {len(field_paragraphs)} 个")
        
        print(f"\n样式统计:")
        for style, count in sorted(style_counts.items()):
            print(f"  {style}: {count} 个")
        
        print(f"\n空段落位置: {empty_paragraphs[:10]}")  # 只显示前10个
        print(f"字段段落位置: {field_paragraphs}")
        
    except Exception as e:
        print(f"结构分析失败: {e}")

if __name__ == "__main__":
    # 测试有内容的文档
    test_file = "data/input/1_计算机应用技术_17211204005-苏慧婧-基于MLP和SepCNN模型的藏文文本分类研究与实现-计算机应用技术-群诺.docx"
    
    analyze_hidden_toc_fields(test_file)
    analyze_document_structure(test_file)
