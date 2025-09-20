#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
检查Word文档中的目录域和字段
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

import docx
from docx.oxml.ns import qn
from docx.oxml import parse_xml
import zipfile
import xml.etree.ElementTree as ET

def analyze_docx_structure(file_path):
    """分析docx文件的内部结构"""
    print(f"📄 分析Word文档结构: {os.path.basename(file_path)}")
    print("=" * 80)
    
    try:
        # 方法1: 使用python-docx读取
        print("\n🔍 方法1: 使用python-docx读取段落")
        print("-" * 50)
        doc = docx.Document(file_path)
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                print(f"第{i+1:3d}行: {text}")
        
        print(f"\n📊 总段落数: {len(doc.paragraphs)}")
        
        # 方法2: 检查是否有目录域
        print("\n🔍 方法2: 检查目录域(TOC Fields)")
        print("-" * 50)
        
        toc_found = False
        for i, paragraph in enumerate(doc.paragraphs):
            # 检查段落的XML内容
            para_xml = paragraph._element.xml
            if 'TOC' in para_xml or 'fldChar' in para_xml:
                print(f"第{i+1}行发现TOC域: {paragraph.text}")
                toc_found = True
        
        if not toc_found:
            print("❌ 未发现TOC域")
        
        # 方法3: 直接读取XML内容
        print("\n🔍 方法3: 解析ZIP结构")
        print("-" * 50)
        
        with zipfile.ZipFile(file_path, 'r') as zip_file:
            # 读取document.xml
            if 'word/document.xml' in zip_file.namelist():
                document_xml = zip_file.read('word/document.xml')
                root = ET.fromstring(document_xml)
                
                # 搜索TOC相关内容
                namespaces = {
                    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                }
                
                # 查找fldChar元素（域字符）
                fld_chars = root.findall('.//w:fldChar', namespaces)
                if fld_chars:
                    print(f" 发现 {len(fld_chars)} 个域字符")
                
                # 查找instrText元素（指令文本）
                instr_texts = root.findall('.//w:instrText', namespaces)
                for instr in instr_texts:
                    if instr.text and 'TOC' in instr.text:
                        print(f" 发现TOC指令: {instr.text}")
                
                # 查找所有文本内容
                print("\n📝 所有文本内容:")
                text_elements = root.findall('.//w:t', namespaces)
                for i, text_elem in enumerate(text_elements):
                    if text_elem.text and text_elem.text.strip():
                        print(f"文本{i+1:3d}: {text_elem.text.strip()}")
        
    except Exception as e:
        print(f"❌ 分析失败: {str(e)}")
        import traceback
        traceback.print_exc()

def test_toc_field_detection():
    """测试目录域检测"""
    test_files = [
        r"c:\MyProjects\thesis_Inno_Eval\data\input\计算机应用技术_test1.docx",
        r"c:\MyProjects\thesis_Inno_Eval\data\input\计算机应用技术_test2.docx",
        r"c:\MyProjects\thesis_Inno_Eval\data\input\1_计算机应用技术_17211204005-苏慧婧-基于MLP和SepCNN模型的藏文文本分类研究与实现-计算机应用技术-群诺.docx"
    ]
    
    for file_path in test_files:
        if os.path.exists(file_path):
            analyze_docx_structure(file_path)
            print("\n" + "="*100 + "\n")
        else:
            print(f"❌ 文件不存在: {file_path}")

if __name__ == "__main__":
    test_toc_field_detection()

