#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
检查Word文档的开头部分内容
"""

from docx import Document

def analyze_document_beginning(docx_path, num_paragraphs=50):
    """分析文档开头的内容"""
    print(f"📄 分析文档开头内容: {docx_path}")
    print("=" * 80)
    
    try:
        doc = Document(docx_path)
        
        print(f"📊 文档统计:")
        print(f"  总段落数: {len(doc.paragraphs)}")
        
        print(f"\n📝 前{num_paragraphs}个段落内容:")
        print("-" * 60)
        
        for i, paragraph in enumerate(doc.paragraphs[:num_paragraphs], 1):
            text = paragraph.text.strip()
            
            if text:  # 只显示非空段落
                style_name = paragraph.style.name if paragraph.style and paragraph.style.name else 'Normal'
                print(f"段落 {i:2d} [{style_name}]: {text[:100]}{'...' if len(text) > 100 else ''}")
                
                # 检查是否可能是目录相关
                if any(keyword in text.lower() for keyword in ['目录', 'contents', 'table', '目　录']):
                    print(f"         ⭐ 可能是目录标题")
                    
                # 检查是否是章节标题
                if any(pattern in text for pattern in ['第', '章', 'ABSTRACT', '摘要', '绪论', '总结', '参考文献']):
                    print(f"         📖 可能是章节标题")
            else:
                print(f"段落 {i:2d}: [空段落]")
        
    except Exception as e:
        print(f"❌ 分析失败: {e}")

def main():
    # 分析文档开头
    test_files = [
        "data/input/1_计算机应用技术_17211204005-苏慧婧-基于MLP和SepCNN模型的藏文文本分类研究与实现-计算机应用技术-群诺.docx",
        "data/input/计算机应用技术_test1.docx"
    ]
    
    for file_path in test_files:
        analyze_document_beginning(file_path)
        print("\n" + "="*80 + "\n")

if __name__ == "__main__":
    main()
