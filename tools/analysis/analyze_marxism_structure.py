"""
更详细检查马克思主义哲学论文的目录结构
"""
import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), 'src'))

import docx
import re

def detailed_toc_analysis():
    """详细分析目录结构"""
    doc_path = r"c:\MyProjects\thesis_Inno_Eval\data\input\1_马克思主义哲学86406_010101_81890101_LW.docx"
    
    print("🔍 详细分析马克思主义哲学论文目录结构")
    print("="*80)
    
    try:
        doc = docx.Document(doc_path)
        
        print("📄 逐行检查文档内容...")
        
        toc_area = False
        chapter_pattern = re.compile(r'第\d+章|参考文献|个人简历|后\s*记')
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            # 显示前50行来理解文档结构
            if i < 50:
                if text:
                    print(f"第{i+1:3d}行: {text}")
                    
                    # 检查是否是目录相关
                    if "目  录" in text:
                        toc_area = True
                        print(f"     ▶ 目录开始标志")
                    elif toc_area and ("摘  要" in text or "Abstract" in text):
                        print(f"     ▶ 目录结束标志")
                        break
                    elif toc_area and (re.search(r'第\d+章', text) or "参考文献" in text or "个人简历" in text or "后记" in text):
                        print(f"     ▶ 可能的目录条目")
        
        print(f"\n🔍 搜索所有可能的章节标题模式:")
        print("-" * 60)
        
        chapter_lines = []
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            # 查找章节标题模式
            if re.search(r'第\d+章\s+', text) or \
               text.startswith('参考文献') or \
               '个人简历' in text or \
               re.search(r'后\s*记', text):
                chapter_lines.append((i+1, text))
                print(f"第{i+1:3d}行: {text}")
        
        # 查找目录中的页码信息
        print(f"\n📖 查找带页码的目录条目:")
        print("-" * 60)
        
        page_pattern = re.compile(r'.+\s+(\d+)$')
        for i, paragraph in enumerate(doc.paragraphs[3:35]):  # 目录应该在前30行内
            text = paragraph.text.strip()
            if text and page_pattern.match(text):
                print(f"第{i+4:3d}行: {text}")
        
    except Exception as e:
        print(f"❌ 分析失败: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    detailed_toc_analysis()
