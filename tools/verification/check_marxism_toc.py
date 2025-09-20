"""
检查马克思主义哲学论文的完整目录内容
"""
import sys
import os
sys.path.append(os.path.join(os.path.dirname(__file__), 'src'))

import docx

def check_full_toc():
    """检查完整的目录内容"""
    doc_path = r"c:\MyProjects\thesis_Inno_Eval\data\input\1_马克思主义哲学86406_010101_81890101_LW.docx"
    
    print("🔍 检查马克思主义哲学论文的完整目录内容")
    print("="*80)
    
    try:
        doc = docx.Document(doc_path)
        
        print("📄 搜索目录相关内容...")
        
        # 查找目录区域
        toc_start = -1
        toc_end = -1
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            # 查找目录开始
            if toc_start == -1 and "目  录" in text:
                toc_start = i
                print(f" 目录开始: 第{i+1}行 - {text}")
            
            # 查找目录结束（摘要开始）
            elif toc_start != -1 and "摘  要" in text:
                toc_end = i
                print(f" 目录结束: 第{i+1}行 - {text}")
                break
        
        if toc_start != -1 and toc_end != -1:
            print(f"\n📋 目录内容 (第{toc_start+1}行到第{toc_end}行):")
            print("-" * 80)
            
            for i in range(toc_start, toc_end):
                text = doc.paragraphs[i].text.strip()
                if text:
                    print(f"第{i+1:3d}行: {text}")
        
        # 专门搜索"后记"相关内容
        print(f"\n🔍 搜索'后记'相关内容:")
        print("-" * 50)
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text and ("后记" in text or "后　记" in text):
                print(f"第{i+1:3d}行: {text}")
        
        # 检查页码48附近的内容
        print(f"\n📖 检查页码48附近内容 (查找可能的后记章节标题):")
        print("-" * 60)
        
        # 简单估算，如果每页约10-15行段落
        estimated_line_48 = 48 * 12  # 粗略估算
        start_check = max(0, estimated_line_48 - 50)
        end_check = min(len(doc.paragraphs), estimated_line_48 + 50)
        
        for i in range(start_check, end_check):
            text = doc.paragraphs[i].text.strip()
            if text and len(text) < 50:  # 可能是章节标题
                if any(keyword in text for keyword in ["后记", "后　记", "结语", "结　语", "总结"]):
                    print(f"第{i+1:3d}行: {text}")
        
    except Exception as e:
        print(f"❌ 检查失败: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    check_full_toc()

